from __future__ import annotations

import csv
import hashlib
import json
import math
import os
import random
import re
import shutil
import subprocess
import sys
import tempfile
import textwrap
import uuid
from collections import Counter
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Optional

try:
    import fitz  # type: ignore
except ImportError:
    fitz = None

try:
    import nltk  # type: ignore
    from nltk import sent_tokenize, word_tokenize  # type: ignore
    from nltk.corpus import stopwords  # type: ignore
    from nltk.stem import WordNetLemmatizer  # type: ignore
except ImportError:
    nltk = None
    sent_tokenize = None
    word_tokenize = None
    stopwords = None
    WordNetLemmatizer = None

try:
    import spacy  # type: ignore
except ImportError:
    spacy = None

try:
    from docx import Document  # type: ignore
except ImportError:
    Document = None

try:
    from PIL import Image, ImageDraw, ImageEnhance, ImageFilter, ImageFont, ImageOps, ImageStat  # type: ignore
except ImportError:
    Image = None
    ImageDraw = None
    ImageEnhance = None
    ImageFont = None
    ImageFilter = None
    ImageOps = None
    ImageStat = None


MAX_TEXT_CHARS = 25_000_000
CHUNK_SIZE = 5_000
MAX_FILE_SIZE = 50 * 1024 * 1024
PANELS_PER_COMIC_PAGE = 4
PANEL_PAGE_PATTERN = (3, 4, 3, 4, 4, 3, 4)
MIN_RENDER_PANEL_HEIGHT = 300
MAX_ANALYSIS_SENTENCES = 800
MAX_GENERATED_SCENES = 96
EDUCATIONAL_FLOW_LABELS = (
    "Concept Intro",
    "Why It Matters",
    "Real Example",
    "Quick Tip",
    "Connect Ideas",
    "Remember This",
    "Check Yourself",
    "Final Idea",
)
PANEL_MODE_SEQUENCE = (
    "question",
    "concept",
    "analogy",
    "why",
    "tip",
    "connect",
    "recap",
    "concept",
)

TECH_TERM_TEACHING = {
    "user interface design": "planning the screens, buttons, messages, and flows users interact with",
    "interface design": "planning how users move through screens, buttons, messages, and actions",
    "software testing": "checking software with planned examples to catch mistakes early",
    "software design": "the blueprint you make before coding, so screens, data, and logic fit together",
    "architecture": "the big structure of a system, like rooms and hallways in a building",
    "model": "a simple version of the real thing that helps you understand it faster",
    "requirement": "a clear need the software must satisfy for its users",
    "requirements": "the needs the software must satisfy for its users",
    "furps": "a software quality checklist: functionality, usability, reliability, performance, and supportability",
    "functionality": "what the system can actually do for the user",
    "usability": "how easy and comfortable the system feels to use",
    "reliability": "how often the system works correctly without failing",
    "performance": "how fast and smoothly the system responds",
    "supportability": "how easy the system is to fix, update, and maintain",
    "maintainability": "how easy the software is to improve later without breaking things",
    "algorithm": "a step-by-step recipe for solving a problem",
    "database": "an organized place where an app stores and finds information",
    "api": "a doorway that lets two pieces of software talk to each other",
    "interface": "the meeting point where people or systems interact",
    "module": "a smaller piece of a system that handles one responsibility",
    "modularity": "splitting a system into smaller parts so each part is easier to understand and test",
    "coupling": "how strongly one part of a program depends on another part",
    "cohesion": "how well one module stays focused on one clear job",
    "class": "a reusable template for creating related objects in code",
    "object": "a bundle of data and behavior that represents one thing in a program",
    "testing": "checking software with planned examples to catch mistakes early",
    "review": "a careful check of design or code before it reaches users",
    "defect": "a mistake in software that should be found and fixed",
    "abstraction": "hiding extra detail so you can focus on the main idea",
    "encapsulation": "keeping data and behavior together so a part of the program is easier to protect",
    "inheritance": "letting one class reuse features from another class",
    "polymorphism": "letting different objects respond to the same action in their own way",
}

JARGON_REPLACEMENTS = (
    (r"\butili[sz]e\b", "use"),
    (r"\bfacilitate\b", "help"),
    (r"\bprior to\b", "before"),
    (r"\bcommence\b", "start"),
    (r"\bcomprises?\b", "includes"),
    (r"\bsubsequently\b", "then"),
    (r"\btherefore\b", "so"),
    (r"\bconsequently\b", "as a result"),
    (r"\bapproximately\b", "about"),
    (r"\bobtain\b", "get"),
    (r"\bdemonstrates?\b", "shows"),
    (r"\brepresents?\b", "shows"),
    (r"\bensures?\b", "helps make sure"),
    (r"\bmaintainability\b", "how easy it is to update later"),
    (r"\breliability\b", "how often it works without failing"),
    (r"\busability\b", "how easy it is to use"),
    (r"\bperformance\b", "how fast and smooth it is"),
    (r"\bfunctionality\b", "what it can do"),
    (r"\bsupportability\b", "how easy it is to fix and support"),
)

ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}
ALLOWED_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}
MAX_IMAGE_UPLOADS = 12
MAX_IMAGE_FILE_SIZE = 15 * 1024 * 1024
MAX_IMAGE_RENDER_DIMENSION = 2200
FALLBACK_STOPWORDS = {
    "a", "an", "and", "are", "as", "at", "be", "by", "for", "from", "has",
    "have", "in", "into", "is", "it", "its", "of", "on", "or", "that", "the",
    "their", "this", "to", "was", "were", "with",
}

SUBJECT_KEYWORDS = {
    "science": {"science", "energy", "plant", "cell", "atom", "force", "oxygen", "experiment"},
    "history": {"history", "empire", "battle", "revolution", "leader", "kingdom", "war", "timeline"},
    "mathematics": {"equation", "graph", "number", "angle", "algebra", "formula", "value", "slope"},
    "literature": {"theme", "plot", "character", "setting", "story", "conflict", "symbol", "narrative"},
}
OCR_TOPIC_KEYWORDS = {
    "Physics": {"force", "motion", "energy", "gravity", "velocity", "acceleration", "wave", "light", "electricity", "magnet", "quantum", "physics"},
    "Chemistry": {"atom", "molecule", "reaction", "compound", "element", "acid", "base", "bond", "periodic", "chemistry", "solution", "electron"},
    "Biology": {"cell", "organism", "biology", "gene", "dna", "plant", "animal", "tissue", "photosynthesis", "microbe", "microscope", "evolution"},
    "Science": {"science", "experiment", "observation", "hypothesis", "microscope", "molecular", "biology", "chemistry", "physics", "research"},
    "Maths": {"equation", "algebra", "geometry", "calculus", "number", "graph", "ratio", "fraction", "theorem", "probability", "statistics", "math"},
    "History": {"history", "empire", "war", "revolution", "king", "queen", "dynasty", "civilization", "independence", "ancient", "timeline"},
    "Geography": {"geography", "map", "climate", "river", "mountain", "plateau", "continent", "ocean", "latitude", "longitude", "population"},
    "Literature": {"literature", "poem", "novel", "author", "character", "plot", "theme", "metaphor", "story", "drama", "narrator"},
    "Computer Science": {"computer", "algorithm", "program", "software", "hardware", "database", "network", "code", "binary", "data", "api"},
    "Economics": {"economics", "market", "demand", "supply", "inflation", "price", "trade", "consumer", "producer", "money", "gdp"},
}
OCR_TOPIC_SUBJECT_MAP = {
    "Physics": "science",
    "Chemistry": "science",
    "Biology": "science",
    "Science": "science",
    "Maths": "mathematics",
    "History": "history",
    "Geography": "history",
    "Literature": "literature",
    "Computer Science": "science",
    "Economics": "history",
}
OCR_TEXT_CORRECTIONS = (
    (r"\bde\s+pect\b", "depict"),
    (r"\bbiclogy\b", "biology"),
    (r"\bbiolcgy\b", "biology"),
    (r"\bchernistry\b", "chemistry"),
    (r"\bchem\s+istry\b", "chemistry"),
    (r"\bphys\s+ics\b", "physics"),
    (r"\bmicro\s+scopes?\b", "microscopes"),
    (r"\bmole\s+cular\b", "molecular"),
)
SUBJECT_STYLE_DEFAULTS = {
    "science": "anime_school",
    "history": "superhero_comic",
    "mathematics": "educational_kids",
    "literature": "cute_cartoon",
}
TITLE_SUBJECT_HINTS = {
    "science": {"stranger things", "marvel", "gran turismo", "lab", "science", "space"},
    "history": {"biography", "amma", "freedom", "empire", "war"},
    "mathematics": {"math", "algebra", "geometry", "equation", "numbers"},
    "literature": {
        "alice", "gatsby", "wuthering", "mockingbird", "lord of the rings", "panchatanthra",
        "panchatantra", "jataka", "jungle book", "grandma stories", "stilton", "wimpy",
        "journey to the west", "kalevala", "bond", "johnny english",
    },
}
TITLE_STYLE_HINTS = {
    "superhero_comic": {"marvel", "bond", "hero"},
    "pixar_colorful": {"stranger things", "gran turismo", "science", "space"},
    "cute_cartoon": {"alice", "jungle book", "jataka", "panchatanthra", "panchatantra", "grandma", "wimpy", "stilton"},
    "classic": {"gatsby", "wuthering", "mockingbird", "kalevala", "lord of the rings", "journey to the west"},
}
GENERIC_ENTITY_STOPWORDS = {
    "chapter", "contents", "introduction", "copyright", "bibliography", "about",
    "author", "book", "books", "part", "page", "pages", "figure", "table",
    "series", "television", "netflix", "productions", "entertainment", "volume",
    "download", "ebooks", "classic", "literature", "drama", "overview",
    "what", "name", "dedication", "however", "meanwhile", "therefore",
}

STYLE_LIBRARY = {
    "pow": {
        "name": "Pow Burst",
        "page": "#fdf3c0",
        "panel_bg": "#fff8d6",
        "panel_top": "#ffcc00",
        "panel_bot": "#fff3a0",
        "accent": "#e8230a",
        "accent_alt": "#ff6b00",
        "accent3": "#1a0aff",
        "bubble": "#ffffff",
        "caption_bg": "#1a0aff",
        "caption": "#ffffff",
        "header_text": "#ffffff",
        "title_bg": "#e8230a",
        "title_stripe": "#ffcc00",
        "sky_top": "#87ceeb",
        "sky_bot": "#d4f0ff",
        "ground": "#8B5E3C",
    },
    "campus": {
        "name": "Campus Chronicle",
        "page": "#eef2ff",
        "panel_bg": "#f5f8ff",
        "panel_top": "#6c8fff",
        "panel_bot": "#c8d8ff",
        "accent": "#2240cc",
        "accent_alt": "#009977",
        "accent3": "#cc4400",
        "bubble": "#ffffff",
        "caption_bg": "#2240cc",
        "caption": "#ffffff",
        "header_text": "#ffffff",
        "title_bg": "#2240cc",
        "title_stripe": "#6c8fff",
        "sky_top": "#c8d8ff",
        "sky_bot": "#eef2ff",
        "ground": "#4a7c59",
    },
    "lab": {
        "name": "Lab Adventure",
        "page": "#d8fff2",
        "panel_bg": "#edfff8",
        "panel_top": "#00cc88",
        "panel_bot": "#b2ffe0",
        "accent": "#006644",
        "accent_alt": "#0055aa",
        "accent3": "#aa0066",
        "bubble": "#ffffff",
        "caption_bg": "#006644",
        "caption": "#ffffff",
        "header_text": "#ffffff",
        "title_bg": "#006644",
        "title_stripe": "#00cc88",
        "sky_top": "#b2ffe0",
        "sky_bot": "#d8fff2",
        "ground": "#2d6a4f",
    },
    "classic": {
        "name": "Classic Strip",
        "page": "#f5e8d0",
        "panel_bg": "#fdf5e8",
        "panel_top": "#c8a870",
        "panel_bot": "#f5e8d0",
        "accent": "#1a1a1a",
        "accent_alt": "#cc3300",
        "accent3": "#004488",
        "bubble": "#fffef0",
        "caption_bg": "#1a1a1a",
        "caption": "#ffffff",
        "header_text": "#ffffff",
        "title_bg": "#1a1a1a",
        "title_stripe": "#cc3300",
        "sky_top": "#a8c8e8",
        "sky_bot": "#d8ecf8",
        "ground": "#7a5c3a",
    },
    "anime": {
        "name": "Anime Spark",
        "page": "#fff4fb",
        "panel_bg": "#fff9ff",
        "panel_top": "#ff5fa2",
        "panel_bot": "#f4e9ff",
        "accent": "#6c5ce7",
        "accent_alt": "#00b4d8",
        "accent3": "#ff8fab",
        "bubble": "#ffffff",
        "caption_bg": "#6c5ce7",
        "caption": "#ffffff",
        "header_text": "#ffffff",
        "title_bg": "#6c5ce7",
        "title_stripe": "#ff5fa2",
        "sky_top": "#dfe8ff",
        "sky_bot": "#fff4fb",
        "ground": "#8d6e63",
    },
    "pixar": {
        "name": "Pixar Glow",
        "page": "#fff7ef",
        "panel_bg": "#fffdf8",
        "panel_top": "#ff8744",
        "panel_bot": "#ffe2bf",
        "accent": "#ff8744",
        "accent_alt": "#20b2aa",
        "accent3": "#5e60ce",
        "bubble": "#fffdf7",
        "caption_bg": "#ff8744",
        "caption": "#ffffff",
        "header_text": "#ffffff",
        "title_bg": "#ff8744",
        "title_stripe": "#ffd166",
        "sky_top": "#dff7ff",
        "sky_bot": "#fff7ef",
        "ground": "#c78d5f",
    },
    "superhero": {
        "name": "Hero Impact",
        "page": "#fff1ef",
        "panel_bg": "#fff7f5",
        "panel_top": "#ea233f",
        "panel_bot": "#ffd0bf",
        "accent": "#ea233f",
        "accent_alt": "#0f52ba",
        "accent3": "#ffd447",
        "bubble": "#ffffff",
        "caption_bg": "#111111",
        "caption": "#ffffff",
        "header_text": "#ffffff",
        "title_bg": "#111111",
        "title_stripe": "#ea233f",
        "sky_top": "#2b3a67",
        "sky_bot": "#fff1ef",
        "ground": "#4f5d75",
    },
    "cartoon": {
        "name": "Cartoon Pop",
        "page": "#fff8dd",
        "panel_bg": "#fffdf1",
        "panel_top": "#ffb703",
        "panel_bot": "#ffe8a3",
        "accent": "#fb5607",
        "accent_alt": "#2ec4b6",
        "accent3": "#7b2cbf",
        "bubble": "#fffef9",
        "caption_bg": "#fb5607",
        "caption": "#ffffff",
        "header_text": "#ffffff",
        "title_bg": "#fb5607",
        "title_stripe": "#ffb703",
        "sky_top": "#d8f3dc",
        "sky_bot": "#fff8dd",
        "ground": "#b98b5d",
    },
    "manga": {
        "name": "Manga Ink",
        "page": "#f1f1f1",
        "panel_bg": "#fbfbfb",
        "panel_top": "#232323",
        "panel_bot": "#d0d0d0",
        "accent": "#232323",
        "accent_alt": "#666666",
        "accent3": "#bdbdbd",
        "bubble": "#ffffff",
        "caption_bg": "#232323",
        "caption": "#ffffff",
        "header_text": "#ffffff",
        "title_bg": "#111111",
        "title_stripe": "#d0d0d0",
        "sky_top": "#f5f5f5",
        "sky_bot": "#d8d8d8",
        "ground": "#8d8d8d",
    },
    "fantasy": {
        "name": "Fantasy Arc",
        "page": "#f8f2ff",
        "panel_bg": "#fdf8ff",
        "panel_top": "#7d4dff",
        "panel_bot": "#e6d9ff",
        "accent": "#7d4dff",
        "accent_alt": "#2dd4bf",
        "accent3": "#facc15",
        "bubble": "#fffdf9",
        "caption_bg": "#4f2ca8",
        "caption": "#ffffff",
        "header_text": "#ffffff",
        "title_bg": "#4f2ca8",
        "title_stripe": "#7d4dff",
        "sky_top": "#d7f5f0",
        "sky_bot": "#f8f2ff",
        "ground": "#6f5b4b",
    },
}

CAST_LIBRARY = {
    "science": [
        ("Nova", "Curious Learner", "student"),
        ("Dr. Spark", "Science Guide", "scientist"),
        ("Lab Buddy", "Experiment Specialist", "scientist"),
        ("Narrator", "Story Voice", "narrator"),
    ],
    "history": [
        ("Mina", "Timeline Explorer", "student"),
        ("Archivist Arjun", "History Guide", "teacher"),
        ("Witness", "Event Expert", "narrator"),
        ("Narrator", "Story Voice", "narrator"),
    ],
    "mathematics": [
        ("Riya", "Problem Solver", "student"),
        ("Coach Sigma", "Math Guide", "teacher"),
        ("Graph Bot", "Formula Expert", "scientist"),
        ("Narrator", "Story Voice", "narrator"),
    ],
    "literature": [
        ("Asha", "Reader", "student"),
        ("Professor Quill", "Literature Guide", "teacher"),
        ("Story Muse", "Theme Expert", "narrator"),
        ("Narrator", "Story Voice", "narrator"),
    ],
}

BACKGROUND_LIBRARY = {
    "science": ["laboratory", "greenhouse", "observation deck", "revision corner"],
    "history": ["archive hall", "timeline wall", "old city square", "classroom map board"],
    "mathematics": ["graph studio", "equation board", "pattern lab", "revision desk"],
    "literature": ["library corner", "story stage", "reading garden", "theme board"],
}

# Character visual configs: hair_color, skin_tone, costume_color, costume_alt
CHARACTER_STYLES = {
    "student":   ("#ff8c42", "#ffe0bd", "#4a90d9", "#2d6cdf"),
    "teacher":   ("#3d2b1f", "#f5cba7", "#2d6cdf", "#1a1a1a"),
    "scientist": ("#1a1a2e", "#ffe0bd", "#1fa38a", "#0d7a68"),
    "narrator":  ("#6c3483", "#f0e6ff", "#8e44ad", "#6c3483"),
}

# Expressions: (eye_style, mouth_style)
EXPRESSIONS = {
    0: ("happy",   "smile"),   # intro
    1: ("excited", "wide"),    # key idea
    2: ("think",   "hmm"),     # connection
    3: ("happy",   "smile"),   # panel 4
    4: ("shock",   "o"),       # panel 5
    5: ("happy",   "grin"),    # outro
}

EFFECT_WORDS = ["POW!", "ZAP!", "WOW!", "AHA!", "BOOM!", "YES!"]
STYLE_RESOLUTION_MAP = {
    "pow": ("superhero", "superhero_comic"),
    "campus": ("anime", "anime_school"),
    "lab": ("pixar", "pixar_colorful"),
    "classic": ("cartoon", "cute_cartoon"),
    "anime": ("anime", "anime_school"),
    "anime_school": ("anime", "anime_school"),
    "pixar": ("pixar", "pixar_colorful"),
    "pixar_colorful": ("pixar", "pixar_colorful"),
    "superhero": ("superhero", "superhero_comic"),
    "superhero_comic": ("superhero", "superhero_comic"),
    "cartoon": ("cartoon", "cute_cartoon"),
    "cartoon_kids": ("cartoon", "cute_cartoon"),
    "cute_cartoon": ("cartoon", "cute_cartoon"),
    "educational_kids": ("cartoon", "educational_kids"),
    "manga": ("manga", "manga_story"),
    "manga_story": ("manga", "manga_story"),
    "fantasy": ("fantasy", "fantasy_adventure"),
    "fantasy_adventure": ("fantasy", "fantasy_adventure"),
}
DATASET_STYLE_LABELS = {
    "anime_school": "Anime School",
    "cute_cartoon": "Cartoon Kids",
    "superhero_comic": "Superhero Comic",
    "pixar_colorful": "Pixar Colorful",
    "educational_kids": "Educational Kids",
    "manga_story": "Manga Story",
    "fantasy_adventure": "Fantasy Adventure",
}
THEME_PROFILE_FALLBACKS = {
    "anime_school": "anime",
    "pixar_colorful": "pixar",
    "superhero_comic": "superhero",
    "cute_cartoon": "cartoon_kids",
    "educational_kids": "cartoon_kids",
    "manga_story": "manga",
    "fantasy_adventure": "fantasy",
}
SUBJECT_CHARACTER_HINTS = {
    "science": {"science", "lab", "robot", "astronaut", "tech", "stem", "biology", "physics", "experiment"},
    "history": {"history", "historical", "india", "royal", "freedom", "storytelling", "drama", "ancient"},
    "mathematics": {"math", "equation", "number", "graph", "calculator", "smart", "tech", "prodigy"},
    "literature": {"literature", "story", "book", "books", "library", "reading", "arts", "magical"},
}
SUBJECT_BACKGROUND_PRIORITY = {
    "science": ["BG_002", "BG_013", "BG_016", "BG_015", "BG_007", "BG_018", "BG_001"],
    "history": ["BG_010", "BG_017", "BG_003", "BG_001", "BG_014"],
    "mathematics": ["BG_004", "BG_020", "BG_001", "BG_013", "BG_019"],
    "literature": ["BG_003", "BG_011", "BG_001", "BG_012", "BG_015"],
}
EMOTION_TO_SFX = {
    "happy": ["YES!", "WOW!"],
    "excited": ["AHA!", "ZAP!", "POW!"],
    "thinking": ["AHA!"],
    "proud": ["YES!", "POW!"],
    "surprised": ["WOW!", "BOOM!"],
    "shocked": ["BOOM!", "WOW!"],
    "confused": ["OOPS!", "AHA!"],
}
HAIR_COLOR_MAP = {
    "black hair": "#1f1f26",
    "brown hair": "#5c3a21",
    "blond hair": "#e6c04d",
    "blonde hair": "#e6c04d",
    "white hair": "#eceff1",
    "grey hair": "#a5adb5",
    "gray hair": "#a5adb5",
    "red hair": "#b84b2d",
    "auburn hair": "#8a4a2f",
    "green hair": "#3a8f58",
    "blue hair": "#2f5faa",
    "purple hair": "#7b53b8",
    "orange hair": "#d9792b",
    "no hair": "#1f1f26",
    "bald": "#1f1f26",
}
EYE_COLOR_MAP = {
    "blue eyes": "#4b8fe8",
    "brown eyes": "#6f4b33",
    "green eyes": "#2f9e67",
    "hazel eyes": "#8d6e3d",
    "grey eyes": "#8e99a8",
    "gray eyes": "#8e99a8",
    "amber eyes": "#c9892f",
    "red eyes": "#c94b4b",
    "yellow eyes": "#d8b339",
    "purple eyes": "#7b53b8",
    "black eyes": "#202124",
}
ARCHIVE_PUBLISHER_COLORS = {
    "dc": ("#2864dc", "#62a3ff"),
    "marvel": ("#d53434", "#ff9f43"),
}


class ValidationError(ValueError):
    """Raised when the request does not satisfy project constraints."""


class DependencyError(RuntimeError):
    """Raised when optional runtime dependencies required for a path are unavailable."""


ProgressCallback = Callable[[str, int, str, Optional[dict[str, Any]]], None]


@dataclass
class Concept:
    label: str
    score: float
    kind: str


@dataclass
class Relation:
    subject: str
    verb: str
    object: str


@dataclass
class Scene:
    title: str
    speaker: str
    role: str
    dialogue: str
    caption: str
    background: str
    focus: str


@dataclass
class SourceSection:
    label: str
    text: str
    source: str
    index: int


@dataclass
class ImageComicAnalysis:
    filename: str
    width: int
    height: int
    orientation: str
    scene: str
    objects: list[str]
    people: str
    expression: str
    action: str
    background: str
    mood: str
    palette: list[str]
    brightness: float
    saturation: float
    contrast: float
    edge_density: float


def _split_text_smart(text: str, chunk_size: int = CHUNK_SIZE) -> list[str]:
    paragraphs = re.split(r"\n{2,}", text)
    chunks, current = [], ""
    for para in paragraphs:
        if len(current) + len(para) + 2 <= chunk_size:
            current = f"{current}\n\n{para}".strip()
        else:
            if current:
                chunks.append(current)
            if len(para) > chunk_size:
                sentences = re.split(r"(?<=[.!?])\s+", para)
                bucket = ""
                for sentence in sentences:
                    if len(bucket) + len(sentence) + 1 <= chunk_size:
                        bucket = f"{bucket} {sentence}".strip()
                    else:
                        if bucket:
                            chunks.append(bucket)
                        bucket = sentence
                if bucket:
                    chunks.append(bucket)
            else:
                current = para
    if current:
        chunks.append(current)
    return chunks or [text[:chunk_size]]


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _lerp_color(c1: tuple, c2: tuple, t: float) -> tuple[int, int, int]:
    return (
        int(c1[0] + (c2[0] - c1[0]) * t),
        int(c1[1] + (c2[1] - c1[1]) * t),
        int(c1[2] + (c2[2] - c1[2]) * t),
    )


def _rgb_to_hex(color: tuple[int, int, int]) -> str:
    return "#{:02x}{:02x}{:02x}".format(*color)


class CominoteEngine:
    def __init__(self, base_dir: Path) -> None:
        self.base_dir = base_dir
        self.output_dir = self.base_dir / "generated"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._spacy_model = None
        self._last_ocr_provider = ""
        self._load_dataset_assets()
        self._load_theme_datasets()

    # ─── Public API ─────────────────────────────────────────────────────────

    def generate(
        self,
        *,
        title: str,
        subject: str,
        style: str,
        cast_mode: str = "auto",
        theme_slug: str = "",
        text: str,
        uploaded_file,
        user_id: str = "",
        user_name: str = "",
        progress_callback: ProgressCallback | None = None,
    ) -> dict:
        cast_mode = cast_mode if cast_mode in {"auto", "cominote", "archive", "mixed"} else "auto"

        self._notify(progress_callback, "collecting", 10, "Collecting your notes and uploaded file.")
        cleaned_text, source_label, theme_meta, source_sections = self._collect_input(
            text=text,
            uploaded_file=uploaded_file,
            theme_slug=theme_slug,
        )
        if theme_meta and (not title.strip() or title.strip() == "Cominote Project"):
            title = f"{theme_meta['title']} Comic"
        if theme_meta and subject not in SUBJECT_KEYWORDS:
            subject = theme_meta["subject"]
        if theme_meta and not style:
            style = theme_meta["recommended_style"]

        subject = subject if subject in SUBJECT_KEYWORDS else "science"
        style = style or "pow"
        chunks = _split_text_smart(cleaned_text)

        self._notify(
            progress_callback, "reading", 18,
            f"Reading {len(chunks)} chunk(s) while preserving the original order.",
            {"chunk_count": len(chunks)},
        )
        concepts, relations, sentences = self._aggregate_chunks(chunks, subject, progress_callback=progress_callback)
        if len(sentences) < 2:
            raise ValidationError("Please provide at least two meaningful sentences so the comic can tell a story.")

        scenes = self._scenes_from_sections(
            title=title,
            subject=subject,
            concepts=concepts,
            relations=relations,
            sentences=sentences,
            sections=source_sections,
            theme_title=(theme_meta or {}).get("title", ""),
        )
        page_plan = self._paginate_story_indices(len(scenes))
        page_count = len(page_plan)

        self._notify(
            progress_callback,
            "rendering",
            82,
            f"Rendering {page_count} comic page(s) from {len(scenes)} ordered story panel(s).",
            {
                "chunk_count": len(chunks),
                "panel_count": len(scenes),
                "page_count": page_count,
            },
        )

        comic_id = uuid.uuid4().hex[:12]
        image_path = self.output_dir / f"{comic_id}.png"
        visual_plan = self._build_visual_plan(title, subject, style, scenes, cast_mode, theme_meta=theme_meta)
        dataset_style = visual_plan["dataset_style"]
        page_paths = self._render(
            image_path=image_path,
            title=title,
            subject=subject,
            style=visual_plan["palette_style"],
            scenes=scenes,
            visual_plan=visual_plan,
        )
        page_count = len(page_paths) or page_count

        payload = {
            "comic_id": comic_id,
            "title": title,
            "subject": subject.title(),
            "style": self._style_display_name(dataset_style),
            "source_label": source_label,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "summary": f"{title} was transformed into a {page_count}-page, {len(scenes)}-panel comic focused on {concepts[0].label if concepts else 'the main lesson'}.",
            "panel_count": len(scenes),
            "page_count": page_count,
            "concepts": [asdict(concept) for concept in concepts],
            "relations": [asdict(relation) for relation in relations],
            "scenes": [asdict(scene) for scene in scenes],
            "downloads": {
                "pdf": f"/api/download/{comic_id}?format=pdf",
                "png": f"/api/download/{comic_id}?format=png",
                "jpeg": f"/api/download/{comic_id}?format=jpeg",
            },
            "image_url": f"/api/comics/{comic_id}/image",
            "meta": {
                "subject_slug": subject,
                "style_slug": visual_plan["palette_style"],
                "dataset_style_slug": dataset_style,
                "cast_mode": cast_mode,
                "cast_sources": visual_plan["cast_sources"],
                "theme_slug": (theme_meta or {}).get("slug", ""),
                "theme_title": (theme_meta or {}).get("title", ""),
                "theme_profile": visual_plan["theme_profile"]["slug"],
                "theme_profile_label": visual_plan["theme_profile"]["label"],
                "font_label": visual_plan["theme_profile"].get("font_label", ""),
                "layout_label": visual_plan["theme_profile"].get("layout_label", ""),
                "theme_render": visual_plan.get("theme_render", {}),
                "featured_character_names": [
                    str(item.get("name") or "").strip()
                    for item in (visual_plan.get("theme_render", {}).get("featured_cast") or [])
                    if isinstance(item, dict) and str(item.get("name") or "").strip()
                ],
                "cast_names": [panel.get("character", {}).get("name", "") for panel in visual_plan["panels"]],
                "background_names": [panel.get("background", {}).get("name", "") for panel in visual_plan["panels"]],
                "user_id": user_id,
                "user_name": user_name,
                "chunk_count": len(chunks),
                "page_count": page_count,
                "panels_per_page": max((len(page) for page in page_plan), default=PANELS_PER_COMIC_PAGE),
                "source_section_count": len(source_sections),
                "source_page_count": sum(1 for section in source_sections if section.source == "PDF"),
                "page_image_files": [
                    str(path.relative_to(self.output_dir))
                    for path in page_paths
                    if path.exists() and path.is_relative_to(self.output_dir)
                ],
                "page_image_urls": [
                    f"/api/comics/{comic_id}/pages/{index + 1}"
                    for index, _path in enumerate(page_paths)
                ],
            },
        }

        self._write_metadata(comic_id, payload)
        self._notify(
            progress_callback,
            "completed",
            100,
            "Your comic is ready.",
            {"chunk_count": len(chunks), "panel_count": len(scenes), "page_count": page_count},
        )
        return payload

    def generate_from_images(
        self,
        *,
        title: str,
        theme_slug: str,
        image_files: list[Path],
        text: str = "",
        user_id: str = "",
        user_name: str = "",
        progress_callback: ProgressCallback | None = None,
    ) -> dict:
        if Image is None or ImageOps is None:
            raise DependencyError("Text image conversion needs Pillow installed in the Python environment.")
        if not image_files:
            raise ValidationError("Please upload at least one JPG or PNG text image.")
        if len(image_files) > MAX_IMAGE_UPLOADS:
            raise ValidationError(f"Please upload {MAX_IMAGE_UPLOADS} text images or fewer.")

        self._last_ocr_provider = ""
        resolved_theme_slug = self._resolve_theme_alias(theme_slug)
        theme_meta = self._theme_datasets_by_slug.get(resolved_theme_slug)
        if theme_meta is None:
            raise ValidationError("Please choose a supported text image comic theme before converting.")

        title = (title or "").strip() or f"{theme_meta['title']} Study Comic"
        guidance = self._normalise_whitespace(text)

        self._notify(
            progress_callback,
            "collecting",
            10,
            "Reading image...",
            {"image_count": len(image_files)},
        )

        ocr_sections = self._extract_ocr_sections_from_images(image_files, progress_callback)
        raw_text = "\n\n".join(section["text"] for section in ocr_sections if section.get("text"))
        cleaned_text = self._clean_ocr_text(raw_text)
        if not cleaned_text:
            raise ValidationError(
                "This image has no readable study content. Please upload notes, textbook screenshots, or educational text images."
            )
        if not self._looks_like_study_content(cleaned_text):
            raise ValidationError(
                "This image has no readable study content. Please upload notes, textbook screenshots, or educational text images."
            )

        detected_topic, subject = self._detect_ocr_topic(cleaned_text)
        learning_points = self._ocr_learning_points(cleaned_text, detected_topic)
        comic_text = self._build_ocr_comic_source_text(
            cleaned_text=cleaned_text,
            learning_points=learning_points,
            detected_topic=detected_topic,
            guidance=guidance,
        )

        self._notify(
            progress_callback,
            "reading",
            42,
            f"Extracted study text and detected {detected_topic} as the topic.",
            {
                "image_count": len(image_files),
                "ocr_character_count": len(cleaned_text),
                "detected_topic": detected_topic,
            },
        )

        def comic_progress(stage: str, progress: int, message: str, meta: dict[str, Any] | None = None) -> None:
            shifted = 44 + int(max(0, min(progress, 100)) * 0.52)
            if stage == "rendering":
                shifted = max(82, shifted)
            image_message = message
            if stage in {"collecting", "reading"}:
                image_message = "Generating comic from extracted study text..."
            self._notify(progress_callback, stage, min(98, shifted), image_message, meta)

        result = self.generate(
            title=title,
            subject=subject,
            style=str(theme_meta.get("recommended_style") or ""),
            cast_mode="auto",
            theme_slug=resolved_theme_slug,
            text=comic_text,
            uploaded_file=None,
            user_id=user_id,
            user_name=user_name,
            progress_callback=comic_progress,
        )

        result["subject"] = detected_topic
        result["source_label"] = (
            f"{len(image_files)} Text Image OCR" if len(image_files) == 1 else f"{len(image_files)} Text Images OCR"
        )
        result["summary"] = (
            f"{title} was generated from OCR-extracted {detected_topic} study content using the "
            f"{theme_meta['title']} art style."
        )
        result.setdefault("meta", {}).update(
            {
                "input_mode": "image_ocr",
                "image_count": len(image_files),
                "ocr_provider": self._last_ocr_provider,
                "ocr_character_count": len(cleaned_text),
                "ocr_detected_topic": detected_topic,
                "ocr_subject_slug": subject,
                "ocr_learning_points": learning_points,
                "ocr_sections": [
                    {
                        "filename": section.get("filename", ""),
                        "line_count": section.get("line_count", 0),
                        "confidence": section.get("confidence", 0),
                    }
                    for section in ocr_sections
                ],
            }
        )
        self._write_metadata(result["comic_id"], result)
        self._notify(
            progress_callback,
            "completed",
            100,
            "Your OCR-based comic is ready.",
            {
                "image_count": len(image_files),
                "panel_count": result.get("panel_count", 0),
                "page_count": result.get("page_count", 0),
            },
        )
        return result

    def get_comic(self, comic_id: str) -> dict:
        path = self.output_dir / f"{comic_id}.json"
        if not path.exists():
            raise FileNotFoundError(comic_id)
        return json.loads(path.read_text(encoding="utf-8"))

    def get_download_path(self, comic_id: str, image_format: str) -> Path:
        png_path = self.output_dir / f"{comic_id}.png"
        if not png_path.exists():
            raise FileNotFoundError(comic_id)
        fmt = image_format.lower()
        if fmt == "png":
            return png_path
        if fmt == "pdf":
            return self._ensure_pdf(comic_id, png_path)
        if fmt not in {"jpeg", "jpg"}:
            raise ValidationError("Only PDF, PNG, and JPEG downloads are supported.")
        return self._ensure_jpeg(comic_id, png_path)

    def image_path(self, comic_id: str) -> Path:
        path = self.output_dir / f"{comic_id}.png"
        if not path.exists():
            raise FileNotFoundError(comic_id)
        return path

    def page_image_path(self, comic_id: str, page_number: int) -> Path:
        page_paths = self._comic_page_paths(comic_id)
        if page_number < 1 or page_number > len(page_paths):
            raise FileNotFoundError(comic_id)
        return page_paths[page_number - 1]

    def list_theme_datasets(self) -> list[dict[str, Any]]:
        return [
            {
                "slug": theme["slug"],
                "title": theme["title"],
                "subject": theme["subject"],
                "recommended_style": theme["recommended_style"],
                "recommended_style_label": self._style_display_name(theme["recommended_style"]),
                "theme_profile": theme.get("theme_profile", "cartoon_kids"),
                "theme_profile_label": theme.get("theme_profile_label", "Cartoon Kids Theme"),
                "theme_character_hint": theme.get("theme_character_hint", ""),
                "theme_background_hint": theme.get("theme_background_hint", ""),
                "font_label": theme.get("font_label", ""),
                "layout_label": theme.get("layout_label", ""),
                "ui_display_font": theme.get("ui_display_font", "\"Bangers\", cursive"),
                "ui_body_font": theme.get("ui_body_font", "\"Comic Neue\", cursive"),
                "theme_accent": theme.get("theme_accent", "#f4631e"),
                "theme_surface": theme.get("theme_surface", "#ffffff"),
                "theme_render": theme.get("theme_render", {}),
                "preview": theme["preview"],
                "word_count": theme["word_count"],
                "segment_count": len(theme["segments"]),
            }
            for theme in sorted(self._theme_datasets, key=lambda item: item["title"].casefold())
        ]

    # ─── Input handling ──────────────────────────────────────────────────────

    def _collect_input(
        self, *, text: str, uploaded_file, theme_slug: str = ""
    ) -> tuple[str, str, dict[str, Any] | None, list[SourceSection]]:
        title_text = self._normalise_whitespace(text)
        file_sections: list[SourceSection] = []
        source_parts: list[str] = []
        theme_meta = None
        theme_text = ""

        if uploaded_file:
            file_sections = self._read_uploaded_sections(uploaded_file)
            source_parts.append("Upload")

        text_sections = self._text_to_sections(title_text, source="Text", label_prefix="Text Section") if title_text else []
        if title_text:
            source_parts.append("Text")

        if theme_slug:
            theme_meta = self._theme_datasets_by_slug.get(theme_slug)
            if theme_meta is None:
                raise ValidationError("The selected theme could not be found.")
            guidance = self._normalise_whitespace(
                " ".join([title_text, " ".join(section.text for section in file_sections)])
            )
            if not title_text and not file_sections:
                theme_text = self._theme_excerpt(theme_meta, guidance)
            if theme_text:
                source_parts.append("Theme")

        sections = [*text_sections, *file_sections]
        if theme_text:
            sections.append(
                SourceSection(
                    label=f"{theme_meta['title']} Theme",
                    text=theme_text,
                    source="Theme",
                    index=len(sections) + 1,
                )
            )

        combined_parts = [f"{section.label}: {section.text}" for section in sections if section.text]
        combined = self._normalise_whitespace("\n\n".join(combined_parts).strip())
        if not combined:
            raise ValidationError("Please add notes, upload a file, or choose a theme.")
        if len(combined) > MAX_TEXT_CHARS:
            raise ValidationError("Input exceeds the maximum supported size. Please stay under 25 million characters.")
        source_label = " + ".join(source_parts) if source_parts else "Text Input"
        return combined, source_label, theme_meta, sections

    def _aggregate_chunks(
        self, chunks: list[str], subject: str, progress_callback: ProgressCallback | None = None
    ) -> tuple[list[Concept], list[Relation], list[str]]:
        all_concepts: list[Concept] = []
        all_relations: list[Relation] = []
        all_sentences: list[str] = []

        for index, chunk in enumerate(chunks):
            progress = 20 + int(((index + 1) / max(len(chunks), 1)) * 50)
            self._notify(
                progress_callback, "reading", progress,
                f"Processing chunk {index + 1} of {len(chunks)}.",
                {"chunk_index": index + 1, "chunk_count": len(chunks)},
            )
            sentences = self._sentences(chunk)
            all_sentences.extend(sentences)
            try:
                concepts = self._concepts(chunk, sentences, subject)
                all_concepts.extend(concepts)
            except ValidationError:
                pass
            all_relations.extend(self._relations(sentences))

        merged: dict[str, Concept] = {}
        for concept in all_concepts:
            key = concept.label.lower()
            if key in merged:
                merged[key] = Concept(
                    label=concept.label,
                    score=round(merged[key].score + concept.score, 2),
                    kind=concept.kind,
                )
            else:
                merged[key] = concept

        top_concepts = sorted(merged.values(), key=lambda item: item.score, reverse=True)[:6]

        seen_relations: set[tuple[str, str, str]] = set()
        unique_relations: list[Relation] = []
        for relation in all_relations:
            key = (relation.subject.lower(), relation.verb.lower(), relation.object.lower())
            if key not in seen_relations:
                seen_relations.add(key)
                unique_relations.append(relation)
            if len(unique_relations) == 5:
                break

        if not top_concepts:
            raise ValidationError("Cominote could not find enough meaningful words to build a comic.")

        all_sentences = all_sentences[:MAX_ANALYSIS_SENTENCES]
        return top_concepts, unique_relations, all_sentences

    def _read_uploaded_file(self, uploaded_file) -> str:
        return "\n\n".join(section.text for section in self._read_uploaded_sections(uploaded_file))

    def _read_uploaded_sections(self, uploaded_file) -> list[SourceSection]:
        if isinstance(uploaded_file, (str, Path)):
            path = Path(uploaded_file)
            filename = path.name
            suffix = path.suffix.lower()
            if suffix not in ALLOWED_EXTENSIONS:
                raise ValidationError("Only .txt, .pdf, and .docx uploads are supported.")
            if not path.exists():
                raise ValidationError("Uploaded file could not be found for processing.")
            if path.stat().st_size > MAX_FILE_SIZE:
                raise ValidationError("Uploaded files must be 50 MB or smaller.")
            if suffix == ".txt":
                return self._text_to_sections(path.read_text(encoding="utf-8", errors="ignore"), source="TXT", label_prefix="Text Section")
            if suffix == ".pdf":
                return self._extract_pdf_sections(path)
            if suffix == ".docx":
                return self._extract_docx_sections(path)
            raise ValidationError(f"Unsupported upload type for {filename}.")

        filename = uploaded_file.filename or ""
        suffix = Path(filename).suffix.lower()
        if suffix not in ALLOWED_EXTENSIONS:
            raise ValidationError("Only .txt, .pdf, and .docx uploads are supported.")

        raw = uploaded_file.read()
        if hasattr(uploaded_file, "stream"):
            uploaded_file.stream.seek(0)
        if len(raw) > MAX_FILE_SIZE:
            raise ValidationError("Uploaded files must be 50 MB or smaller.")

        if suffix == ".txt":
            return self._text_to_sections(raw.decode("utf-8", errors="ignore"), source="TXT", label_prefix="Text Section")
        if suffix == ".pdf":
            return self._extract_pdf_sections(raw)
        if suffix == ".docx":
            return self._extract_docx_sections(raw)
        raise ValidationError("Unsupported upload type.")

    def _text_to_sections(self, text: str, *, source: str, label_prefix: str) -> list[SourceSection]:
        cleaned = self._normalise_whitespace(text)
        if not cleaned:
            return []
        chunks = _split_text_smart(cleaned, chunk_size=2_800)
        return [
            SourceSection(label=f"{label_prefix} {index}", text=chunk, source=source, index=index)
            for index, chunk in enumerate(chunks, start=1)
            if chunk
        ]

    def _extract_pdf_text(self, source: bytes | Path) -> str:
        return "\n\n".join(section.text for section in self._extract_pdf_sections(source))

    def _extract_pdf_sections(self, source: bytes | Path) -> list[SourceSection]:
        if fitz is None:
            raise DependencyError("PDF upload support needs PyMuPDF installed in the Python environment.")
        sections: list[SourceSection] = []
        if isinstance(source, Path):
            with fitz.open(str(source)) as document:
                for index, page in enumerate(document, start=1):
                    text = self._clean_extracted_page_text(page.get_text("text"))
                    if text:
                        sections.append(SourceSection(label=f"Page {index}", text=text, source="PDF", index=index))
        else:
            with fitz.open(stream=source, filetype="pdf") as document:
                for index, page in enumerate(document, start=1):
                    text = self._clean_extracted_page_text(page.get_text("text"))
                    if text:
                        sections.append(SourceSection(label=f"Page {index}", text=text, source="PDF", index=index))
        if not sections:
            raise ValidationError("The uploaded PDF did not contain extractable text.")
        return sections

    @staticmethod
    def _clean_extracted_page_text(text: str) -> str:
        lines: list[str] = []
        for raw_line in (text or "").splitlines():
            line = re.sub(r"\s+", " ", raw_line).strip()
            if not line:
                continue
            if re.fullmatch(r"(page\s*)?\d{1,4}", line, flags=re.IGNORECASE):
                continue
            lines.append(line)
        return "\n".join(lines)

    def _extract_docx_text(self, source: bytes | Path) -> str:
        return "\n\n".join(section.text for section in self._extract_docx_sections(source))

    def _extract_docx_sections(self, source: bytes | Path) -> list[SourceSection]:
        if Document is None:
            raise DependencyError("DOCX upload support needs python-docx installed in the Python environment.")
        if isinstance(source, Path):
            document = Document(str(source))
        else:
            from io import BytesIO
            document = Document(BytesIO(source))
        text = "\n".join(
            self._normalise_whitespace(paragraph.text)
            for paragraph in document.paragraphs
            if self._normalise_whitespace(paragraph.text)
        )
        sections = self._text_to_sections(text, source="DOCX", label_prefix="DOCX Section")
        if not sections:
            raise ValidationError("The uploaded DOCX did not contain extractable text.")
        return sections

    def _resolve_theme_alias(self, theme_slug: str) -> str:
        slug = self._slugify(theme_slug or "")
        if slug in self._theme_datasets_by_slug:
            return slug
        aliases = {
            "anime": "naruto",
            "manga-anime": "naruto",
            "superhero": "marvel",
            "sci-fi": "space",
            "sci-fi-comic": "space",
            "scifi": "space",
            "science-fiction": "space",
            "cartoon-kids": "cartoon",
        }
        return aliases.get(slug, slug)

    def _extract_ocr_sections_from_images(
        self,
        image_files: list[Path],
        progress_callback: ProgressCallback | None = None,
    ) -> list[dict[str, Any]]:
        sections: list[dict[str, Any]] = []
        for index, path in enumerate(image_files):
            self._validate_image_upload_path(path)
            progress = 14 + int((index / max(len(image_files), 1)) * 24)
            self._notify(
                progress_callback,
                "reading",
                progress,
                f"Reading image {index + 1} of {len(image_files)}...",
                {"image_index": index + 1, "image_count": len(image_files)},
            )

            with tempfile.NamedTemporaryFile(prefix="cominote-ocr-", suffix=".png", delete=False) as handle:
                preprocessed_path = Path(handle.name)
            try:
                clarity = self._preprocess_image_for_ocr(path, preprocessed_path)
                self._notify(
                    progress_callback,
                    "reading",
                    progress + 4,
                    f"Extracting text from image {index + 1} of {len(image_files)}...",
                    {"image_index": index + 1, "image_count": len(image_files)},
                )
                ocr_result = self._run_ocr(preprocessed_path)
            finally:
                preprocessed_path.unlink(missing_ok=True)

            raw_text = str(ocr_result.get("text") or "")
            cleaned = self._clean_ocr_text(raw_text)
            line_count = len([line for line in (ocr_result.get("lines") or []) if str(line.get("text") or "").strip()])
            confidence_values = [
                float(line.get("confidence") or 0)
                for line in (ocr_result.get("lines") or [])
                if isinstance(line, dict)
            ]
            avg_confidence = round(sum(confidence_values) / len(confidence_values), 3) if confidence_values else 0.0

            if not cleaned:
                if clarity["contrast"] >= 0.035 and clarity["edge_density"] < 0.018:
                    raise ValidationError("Image unclear. Please upload a clearer image.")
                raise ValidationError(
                    "This image has no readable study content. Please upload notes, textbook screenshots, or educational text images."
                )
            if len(cleaned) < 12:
                raise ValidationError(
                    "This image has no readable study content. Please upload notes, textbook screenshots, or educational text images."
                )
            if line_count == 0:
                raise ValidationError("Image unclear. Please upload a clearer image.")
            if avg_confidence and avg_confidence < 0.28:
                raise ValidationError("Image unclear. Please upload a clearer image.")

            sections.append(
                {
                    "filename": path.name,
                    "text": cleaned,
                    "raw_text": raw_text,
                    "line_count": line_count,
                    "confidence": avg_confidence,
                    "clarity": clarity,
                }
            )

        return sections

    def _validate_image_upload_path(self, path: Path) -> None:
        suffix = path.suffix.lower()
        if suffix not in ALLOWED_IMAGE_EXTENSIONS:
            raise ValidationError("Only JPG, JPEG, and PNG text image uploads are supported.")
        if not path.exists():
            raise ValidationError("Uploaded text image could not be found for processing.")
        if path.stat().st_size > MAX_IMAGE_FILE_SIZE:
            raise ValidationError("Each text image must be 15 MB or smaller.")

    def _preprocess_image_for_ocr(self, source_path: Path, output_path: Path) -> dict[str, float]:
        if Image is None or ImageOps is None or ImageEnhance is None:
            raise DependencyError("Image OCR preprocessing needs Pillow installed in the Python environment.")
        try:
            with Image.open(source_path) as source:
                image = ImageOps.exif_transpose(source)
                image = image.convert("RGB")
        except Exception as exc:
            raise ValidationError("Uploaded text image could not be opened as a valid JPG or PNG.") from exc

        image.thumbnail((2600, 2600), Image.Resampling.LANCZOS)
        if min(image.size) < 900:
            scale = min(3.0, 900 / max(min(image.size), 1))
            image = image.resize((int(image.width * scale), int(image.height * scale)), Image.Resampling.LANCZOS)

        gray = ImageOps.grayscale(image)
        stat = ImageStat.Stat(gray) if ImageStat is not None else None
        mean = stat.mean[0] / 255 if stat else 0.5
        stddev = stat.stddev[0] / 255 if stat else 0.1
        gray = ImageEnhance.Contrast(gray).enhance(1.7 if stddev < 0.18 else 1.35)
        gray = ImageEnhance.Sharpness(gray).enhance(1.45)
        gray = gray.filter(ImageFilter.MedianFilter(size=3)) if ImageFilter is not None else gray

        if ImageStat is not None and ImageFilter is not None:
            edges = gray.filter(ImageFilter.FIND_EDGES)
            edge_density = min(1.0, (ImageStat.Stat(edges).mean[0] / 255) * 2.2)
        else:
            edge_density = stddev

        gray.save(output_path, format="PNG")
        return {
            "brightness": round(mean, 3),
            "contrast": round(stddev, 3),
            "edge_density": round(edge_density, 3),
        }

    def _run_ocr(self, image_path: Path) -> dict[str, Any]:
        self._last_ocr_provider = ""
        errors: list[str] = []
        empty_result: dict[str, Any] | None = None
        for provider in (self._ocr_with_macos_vision, self._ocr_with_tesseract_cli, self._ocr_with_pytesseract):
            try:
                result = provider(image_path)
            except DependencyError as exc:
                errors.append(str(exc))
                continue
            if result.get("error"):
                errors.append(str(result["error"]))
                continue
            text = str(result.get("text") or "").strip()
            if text:
                return result
            if empty_result is None:
                empty_result = result

        if empty_result is not None:
            return empty_result

        if errors:
            raise DependencyError(
                "OCR is unavailable or failed. Install Tesseract OCR, pytesseract, or run ComiNote on macOS with Vision OCR available. "
                f"Details: {' | '.join(errors[:3])}"
            )
        raise DependencyError("OCR is unavailable. Install Tesseract OCR or run ComiNote on macOS with Vision OCR available.")

    def _ocr_with_macos_vision(self, image_path: Path) -> dict[str, Any]:
        if sys.platform != "darwin":
            raise DependencyError("macOS Vision OCR is not available on this platform.")
        swift = shutil.which("swift")
        helper = self.base_dir / "scripts" / "ocr_vision.swift"
        if not swift or not helper.exists():
            raise DependencyError("macOS Vision OCR helper is not available.")

        env = os.environ.copy()
        env.setdefault("CLANG_MODULE_CACHE_PATH", "/tmp/cominote-clang-cache")
        completed = subprocess.run(
            [swift, str(helper), str(image_path)],
            cwd=str(self.base_dir),
            env=env,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=45,
            check=False,
        )
        if completed.returncode != 0 and not completed.stdout.strip():
            raise DependencyError((completed.stderr or "macOS Vision OCR failed.").strip())
        try:
            payload = json.loads(completed.stdout.strip() or "{}")
        except json.JSONDecodeError as exc:
            raise DependencyError("macOS Vision OCR returned an unreadable response.") from exc
        if payload.get("error"):
            return payload
        self._last_ocr_provider = "macos_vision"
        return payload

    def _ocr_with_tesseract_cli(self, image_path: Path) -> dict[str, Any]:
        tesseract = shutil.which("tesseract")
        if not tesseract:
            raise DependencyError("Tesseract CLI is not installed.")
        completed = subprocess.run(
            [tesseract, str(image_path), "stdout", "--psm", "6", "-l", "eng"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=45,
            check=False,
        )
        if completed.returncode != 0:
            raise DependencyError((completed.stderr or "Tesseract OCR failed.").strip())
        self._last_ocr_provider = "tesseract_cli"
        text = completed.stdout.strip()
        return {
            "text": text,
            "lines": [{"text": line, "confidence": 0.75} for line in text.splitlines() if line.strip()],
            "error": None,
        }

    def _ocr_with_pytesseract(self, image_path: Path) -> dict[str, Any]:
        try:
            import pytesseract  # type: ignore
        except ImportError as exc:
            raise DependencyError("pytesseract is not installed.") from exc
        with Image.open(image_path) as image:
            text = pytesseract.image_to_string(image, config="--psm 6")
        self._last_ocr_provider = "pytesseract"
        text = text.strip()
        return {
            "text": text,
            "lines": [{"text": line, "confidence": 0.7} for line in text.splitlines() if line.strip()],
            "error": None,
        }

    def _clean_ocr_text(self, text: str) -> str:
        if not text:
            return ""
        cleaned = text.replace("\ufeff", " ").replace("\u00a0", " ")
        cleaned = re.sub(r"(?<=\w)-\s*\n\s*(?=\w)", "", cleaned)
        cleaned = re.sub(r"[|~`^*_={}<>\\]+", " ", cleaned)
        cleaned = re.sub(r"[^\S\n]+", " ", cleaned)
        raw_lines = [line.strip(" \t-•·") for line in cleaned.splitlines()]
        lines: list[str] = []
        for line in raw_lines:
            line = re.sub(r"\s+", " ", line).strip()
            if not line:
                continue
            if re.fullmatch(r"[\W\d_]{1,6}", line):
                continue
            if len(line) <= 2:
                continue
            lines.append(line)

        paragraphs: list[str] = []
        current = ""
        for line in lines:
            if not current:
                current = line
                continue
            if current.endswith((".", "!", "?", ":")) or self._looks_like_heading(line):
                paragraphs.append(current)
                current = line
            else:
                current = f"{current} {line}"
        if current:
            paragraphs.append(current)

        normalized = "\n\n".join(self._normalise_whitespace(paragraph) for paragraph in paragraphs)
        normalized = re.sub(r"\s+([,.;:!?])", r"\1", normalized)
        normalized = re.sub(r"([.!?])(?=[A-Z])", r"\1 ", normalized)
        for pattern, replacement in OCR_TEXT_CORRECTIONS:
            normalized = re.sub(pattern, replacement, normalized, flags=re.IGNORECASE)
        return normalized.strip()

    def _looks_like_study_content(self, text: str) -> bool:
        words = self._tokens(text)
        if len(words) < 8:
            return False
        lower = text.lower()
        topic_hits = sum(
            1
            for keywords in OCR_TOPIC_KEYWORDS.values()
            for keyword in keywords
            if keyword in lower
        )
        academic_markers = {
            "define", "means", "concept", "example", "study", "notes", "chapter", "section",
            "process", "system", "structure", "function", "cause", "effect", "important",
            "therefore", "because", "explain", "classification", "types",
        }
        marker_hits = sum(1 for marker in academic_markers if marker in lower)
        sentence_count = len(self._sentences(text))
        return topic_hits >= 1 or marker_hits >= 2 or (len(words) >= 24 and sentence_count >= 2)

    def _detect_ocr_topic(self, text: str) -> tuple[str, str]:
        lower = text.lower()
        scores: dict[str, int] = {}
        for topic, keywords in OCR_TOPIC_KEYWORDS.items():
            score = 0
            for keyword in keywords:
                hits = lower.count(keyword)
                if hits:
                    score += hits * (3 if len(keyword) > 5 else 2)
            scores[topic] = score

        topic = max(scores, key=scores.get)
        if scores.get(topic, 0) <= 0:
            topic = "Science"
        return topic, OCR_TOPIC_SUBJECT_MAP.get(topic, "science")

    def _ocr_learning_points(self, text: str, detected_topic: str) -> list[str]:
        sentences = self._section_candidate_sentences(text)
        if not sentences:
            sentences = self._sentences(text)
        scored: list[tuple[int, int, str]] = []
        topic_terms = OCR_TOPIC_KEYWORDS.get(detected_topic, set())
        for index, sentence in enumerate(sentences):
            lower = sentence.lower()
            score = min(len(sentence.split()), 24)
            score += sum(6 for keyword in topic_terms if keyword in lower)
            score += sum(3 for cue in ("means", "is", "are", "because", "therefore", "process", "function", "example") if cue in lower)
            if self._looks_like_heading(sentence):
                score += 4
            scored.append((score, -index, self._finish_sentence(self._humanize_text(sentence))))
        selected: list[str] = []
        seen: set[str] = set()
        for _score, _neg_index, sentence in sorted(scored, reverse=True):
            key = self._sentence_selection_key(sentence)
            if key in seen:
                continue
            seen.add(key)
            selected.append(self._shorten_text(sentence, width=180))
            if len(selected) >= 10:
                break
        return selected or [self._shorten_text(text, width=180)]

    def _build_ocr_comic_source_text(
        self,
        *,
        cleaned_text: str,
        learning_points: list[str],
        detected_topic: str,
        guidance: str,
    ) -> str:
        point_text = "\n".join(f"- {point}" for point in learning_points)
        sections = [
            f"{detected_topic} lesson overview.",
            point_text,
            cleaned_text,
        ]
        if guidance:
            sections.append(f"Extra user guidance:\n{guidance}")
        return "\n\n".join(section.strip() for section in sections if section.strip())

    def _prepare_image_comic_inputs(
        self,
        image_files: list[Path],
        progress_callback: ProgressCallback | None = None,
    ) -> list[dict[str, Any]]:
        prepared: list[dict[str, Any]] = []
        for index, path in enumerate(image_files):
            progress = 18 + int(((index + 1) / max(len(image_files), 1)) * 42)
            image = self._load_image_for_comic(path)
            analysis = self._analyze_image_for_comic(image, path.name)
            prepared.append({"image": image, "analysis": analysis})
            self._notify(
                progress_callback,
                "reading",
                progress,
                f"Analyzed image {index + 1} of {len(image_files)}: {analysis.scene}.",
                {"image_index": index + 1, "image_count": len(image_files)},
            )
        return prepared

    def _load_image_for_comic(self, path: Path) -> Any:
        suffix = path.suffix.lower()
        if suffix not in ALLOWED_IMAGE_EXTENSIONS:
            raise ValidationError("Only JPG, JPEG, and PNG image uploads are supported.")
        if not path.exists():
            raise ValidationError("Uploaded image could not be found for processing.")
        if path.stat().st_size > MAX_IMAGE_FILE_SIZE:
            raise ValidationError("Each image must be 15 MB or smaller.")

        try:
            with Image.open(path) as source:
                image = ImageOps.exif_transpose(source)
                if image.mode in {"RGBA", "LA"}:
                    alpha = image.getchannel("A")
                    background = Image.new("RGB", image.size, "#ffffff")
                    background.paste(image.convert("RGB"), mask=alpha)
                    image = background
                else:
                    image = image.convert("RGB")

                image.thumbnail(
                    (MAX_IMAGE_RENDER_DIMENSION, MAX_IMAGE_RENDER_DIMENSION),
                    Image.Resampling.LANCZOS,
                )
                return image.copy()
        except ValidationError:
            raise
        except Exception as exc:
            raise ValidationError("One of the uploaded images could not be opened as a valid JPG or PNG.") from exc

    def _analyze_image_for_comic(self, image: Any, filename: str) -> ImageComicAnalysis:
        sample = image.copy()
        sample.thumbnail((320, 320), Image.Resampling.BICUBIC)
        pixels = list(sample.getdata())
        total = max(len(pixels), 1)

        brightness_values: list[float] = []
        saturation_total = 0.0
        green = blue = warm = dark = skin = 0
        for r, g, b in pixels:
            luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255
            brightness_values.append(luminance)
            mx = max(r, g, b)
            mn = min(r, g, b)
            saturation_total += 0 if mx == 0 else (mx - mn) / mx
            if g > r * 1.12 and g > b * 1.02 and g > 70:
                green += 1
            if b > r * 1.08 and b > g * 0.92 and b > 85:
                blue += 1
            if r > b * 1.18 and r > 90 and g > 50:
                warm += 1
            if luminance < 0.25:
                dark += 1
            if r > 95 and g > 45 and b > 20 and r > g and r > b and max(r, g, b) - min(r, g, b) > 15:
                skin += 1

        brightness = sum(brightness_values) / total
        saturation = saturation_total / total
        contrast = (sum((value - brightness) ** 2 for value in brightness_values) / total) ** 0.5
        if ImageStat is not None:
            edge_image = sample.convert("L").filter(ImageFilter.FIND_EDGES)
            edge_density = min(1.0, (ImageStat.Stat(edge_image).mean[0] / 255) * 2.2)
        else:
            edge_density = contrast

        green_ratio = green / total
        blue_ratio = blue / total
        warm_ratio = warm / total
        dark_ratio = dark / total
        skin_ratio = skin / total

        width, height = image.size
        if width > height * 1.18:
            orientation = "landscape"
        elif height > width * 1.18:
            orientation = "portrait"
        else:
            orientation = "square"

        if dark_ratio > 0.42:
            scene = "night or low-light scene"
        elif green_ratio > 0.24:
            scene = "outdoor nature scene"
        elif blue_ratio > 0.26 and brightness > 0.42:
            scene = "sky or water scene"
        elif edge_density > 0.28 and contrast > 0.24:
            scene = "busy action scene"
        elif saturation < 0.16:
            scene = "quiet indoor scene"
        else:
            scene = "colorful everyday scene"

        if skin_ratio > 0.055:
            people = "one or more people likely visible"
        elif skin_ratio > 0.018:
            people = "possible person or character detail"
        else:
            people = "no clear person detected"

        if "people" in people or "person" in people:
            if dark_ratio > 0.34:
                expression = "serious or dramatic expression"
            elif saturation > 0.32 and brightness > 0.48:
                expression = "bright cheerful expression"
            elif edge_density > 0.3:
                expression = "alert action-focused expression"
            else:
                expression = "focused expression"
        else:
            expression = "visual mood led by the main subject"

        if edge_density > 0.32:
            action = "fast movement or busy detail"
        elif contrast > 0.25:
            action = "strong pose with dramatic contrast"
        else:
            action = "calm staged moment"

        if green_ratio > 0.24:
            background = "green outdoor background"
        elif blue_ratio > 0.24:
            background = "open blue sky or water background"
        elif dark_ratio > 0.36:
            background = "shadow-heavy background"
        elif warm_ratio > 0.22:
            background = "warm indoor or sunset background"
        else:
            background = "mixed everyday background"

        if dark_ratio > 0.36:
            mood = "dramatic"
        elif saturation > 0.34 and edge_density > 0.24:
            mood = "energetic"
        elif brightness > 0.62:
            mood = "bright"
        else:
            mood = "balanced"

        objects: list[str] = ["main subject", "background shapes"]
        if "people" in people or "person" in people:
            objects.insert(0, "people or character")
        if green_ratio > 0.18:
            objects.append("plants or outdoor elements")
        if blue_ratio > 0.2:
            objects.append("sky or water")
        if warm_ratio > 0.2:
            objects.append("warm colored objects")
        if edge_density > 0.26:
            objects.append("bold edges and action detail")
        objects = list(dict.fromkeys(objects))[:5]

        return ImageComicAnalysis(
            filename=filename,
            width=width,
            height=height,
            orientation=orientation,
            scene=scene,
            objects=objects,
            people=people,
            expression=expression,
            action=action,
            background=background,
            mood=mood,
            palette=self._dominant_palette(sample),
            brightness=round(brightness, 3),
            saturation=round(saturation, 3),
            contrast=round(contrast, 3),
            edge_density=round(edge_density, 3),
        )

    def _dominant_palette(self, image: Any, color_count: int = 5) -> list[str]:
        palette_image = image.copy()
        palette_image.thumbnail((96, 96), Image.Resampling.BICUBIC)
        try:
            quantized = palette_image.quantize(colors=color_count, method=Image.Quantize.MEDIANCUT)
            colors = quantized.convert("RGB").getcolors(maxcolors=96 * 96) or []
        except Exception:
            colors = palette_image.convert("RGB").getcolors(maxcolors=96 * 96) or []
        colors = sorted(colors, key=lambda item: item[0], reverse=True)
        dominant: list[str] = []
        for _count, color in colors:
            if isinstance(color, int):
                continue
            r, g, b = color[:3]
            hex_color = _rgb_to_hex((r, g, b))
            if hex_color not in dominant:
                dominant.append(hex_color)
            if len(dominant) >= color_count:
                break
        return dominant or ["#ffd93d", "#f4631e", "#1a1a1a"]

    def _infer_image_subject(self, analyses: list[ImageComicAnalysis], theme_meta: dict[str, Any]) -> str:
        theme_profile = str(theme_meta.get("theme_profile") or "").lower()
        scene_blob = " ".join(
            f"{analysis.scene} {analysis.background} {' '.join(analysis.objects)}"
            for analysis in analyses
        ).lower()
        if theme_profile in {"superhero", "anime", "fantasy"}:
            return "literature"
        if "plant" in scene_blob or "sky" in scene_blob or "water" in scene_blob:
            return "science"
        if "edge" in scene_blob or "shape" in scene_blob:
            return "mathematics"
        return "literature"

    def _image_concepts(self, analyses: list[ImageComicAnalysis], theme_meta: dict[str, Any]) -> list[Concept]:
        counts: Counter[str] = Counter()
        for analysis in analyses:
            counts.update([analysis.scene.title(), analysis.mood.title(), analysis.action.title()])
            counts.update(item.title() for item in analysis.objects)
        counts.update([str(theme_meta.get("title") or "Image Comic")])
        concepts = [
            Concept(label=label, score=round(1.0 + count * 0.4, 2), kind="visual")
            for label, count in counts.most_common(6)
        ]
        return concepts or [Concept(label="Image Comic", score=1.0, kind="visual")]

    def _image_scenes(
        self,
        *,
        title: str,
        subject: str,
        analyses: list[ImageComicAnalysis],
        theme_title: str,
        guidance: str = "",
    ) -> list[Scene]:
        backgrounds = BACKGROUND_LIBRARY.get(subject, BACKGROUND_LIBRARY["literature"])
        cast = self._story_cast(subject, self._image_concepts(analyses, {"title": theme_title}), [], theme_title)
        scenes: list[Scene] = []
        single_image_modes = ("Opening Shot", "Character Beat", "Action Beat", "Final Frame")
        multi_image_modes = ("Scene Reveal", "Action Beat")
        if len(analyses) == 1:
            plan = [(0, mode) for mode in single_image_modes]
        elif len(analyses) == 2:
            plan = [(image_index, mode) for image_index in range(2) for mode in multi_image_modes]
        else:
            plan = [(image_index, "Scene Reveal") for image_index in range(len(analyses))]

        while len(plan) < 3:
            plan.append((len(plan) % len(analyses), single_image_modes[len(plan) % len(single_image_modes)]))

        for index, (image_index, mode) in enumerate(plan[:MAX_IMAGE_UPLOADS]):
            analysis = analyses[image_index]
            speaker, role, _kind = cast[index % len(cast)]
            object_phrase = ", ".join(analysis.objects[:3])
            if mode == "Opening Shot":
                dialogue = f"This {analysis.scene} sets the world for {title}."
                caption = f"Detected {analysis.background}, {object_phrase}, and a {analysis.mood} mood."
            elif mode == "Character Beat":
                dialogue = f"The focus feels like {analysis.expression}."
                caption = f"Cominote keeps the full {analysis.orientation} image visible while adding comic polish."
            elif mode == "Action Beat":
                dialogue = f"The action reads as {analysis.action}."
                caption = f"Bold outlines, theme colors, and motion accents turn the image into a comic panel."
            else:
                dialogue = f"{theme_title} style gives this moment a finished comic-page look."
                caption = guidance[:180] if guidance else f"Scene cue: {analysis.scene}; background cue: {analysis.background}."

            scenes.append(
                Scene(
                    title=mode,
                    speaker=speaker,
                    role=role,
                    dialogue=self._shorten_text(dialogue, width=118),
                    caption=self._shorten_text(caption, width=148),
                    background=backgrounds[index % len(backgrounds)],
                    focus=analysis.scene.title(),
                )
            )
        return scenes

    @staticmethod
    def _normalise_whitespace(text: str) -> str:
        return re.sub(r"\s+", " ", (text or "")).strip()

    # ─── NLP layer ──────────────────────────────────────────────────────────

    def _sentences(self, text: str) -> list[str]:
        if sent_tokenize is not None and nltk is not None:
            try:
                return [s.strip() for s in sent_tokenize(text) if s.strip()]
            except LookupError:
                pass
        return [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]

    def _tokens(self, text: str) -> list[str]:
        if word_tokenize is not None:
            try:
                return [t for t in word_tokenize(text) if re.search(r"[A-Za-z]", t)]
            except LookupError:
                pass
        return re.findall(r"[A-Za-z][A-Za-z'-]*", text)

    def _stopwords(self) -> set[str]:
        if stopwords is not None:
            try:
                return set(stopwords.words("english"))
            except LookupError:
                return FALLBACK_STOPWORDS
        return FALLBACK_STOPWORDS

    def _lemmatize(self, token: str) -> str:
        word = token.lower()
        if WordNetLemmatizer is not None:
            try:
                return WordNetLemmatizer().lemmatize(word)
            except LookupError:
                pass
        if word.endswith("ies") and len(word) > 4:
            return f"{word[:-3]}y"
        if word.endswith("ing") and len(word) > 5:
            return word[:-3]
        if word.endswith("ed") and len(word) > 4:
            return word[:-2]
        if word.endswith(("ches", "shes", "sses", "xes", "zes")) and len(word) > 5:
            return word[:-2]
        if word.endswith("s") and not word.endswith(("ics", "ss", "us")) and len(word) > 3:
            return word[:-1]
        return word

    def _concepts(self, text: str, sentences: list[str], subject: str) -> list[Concept]:
        stop = self._stopwords() | {
            "study",
            "note",
            "lesson",
            "overview",
            "detail",
            "learning",
            "point",
            "key",
        }
        sentence_terms = []
        totals: Counter[str] = Counter()

        for sentence in sentences:
            cleaned = []
            for token in self._tokens(sentence):
                lemma = self._lemmatize(token)
                if lemma in stop or len(lemma) < 3:
                    continue
                cleaned.append(lemma)
            if cleaned:
                sentence_terms.append(cleaned)
                totals.update(cleaned)

        if not totals:
            raise ValidationError("Cominote could not find enough meaningful words to build a comic.")

        doc_count = len(sentence_terms) or 1
        doc_frequency: Counter[str] = Counter()
        for sentence in sentence_terms:
            doc_frequency.update(set(sentence))

        entities = self._entities(text)
        entity_labels = {e.lower() for e in entities}
        subject_terms = SUBJECT_KEYWORDS.get(subject, set())
        scored = []

        total_terms = sum(totals.values())
        for term, count in totals.items():
            tf = count / total_terms
            idf = math.log((1 + doc_count) / (1 + doc_frequency[term])) + 1
            score = tf * idf
            if term in subject_terms:
                score *= 1.3
            if term in entity_labels:
                score *= 1.2
            kind = "entity" if term in entity_labels else "concept"
            scored.append(Concept(label=term.title(), score=round(score * 12, 2), kind=kind))

        scored.sort(key=lambda item: item.score, reverse=True)
        unique = []
        seen = set()
        for concept in scored:
            key = concept.label.lower()
            if key in seen:
                continue
            seen.add(key)
            unique.append(concept)
            if len(unique) == 6:
                break
        return unique

    def _entities(self, text: str) -> list[str]:
        model = self._load_spacy_model()
        if model is not None:
            doc = model(text)
            entities = [ent.text.strip() for ent in doc.ents if ent.text.strip()]
            if entities:
                return entities[:8]

        matches = re.findall(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b", text)
        unique = []
        for match in matches:
            if match.lower() in {"the", "this"}:
                continue
            if match not in unique:
                unique.append(match)
            if len(unique) == 8:
                break
        return unique

    def _load_spacy_model(self):
        if spacy is None:
            return None
        if self._spacy_model is not None:
            return self._spacy_model
        try:
            self._spacy_model = spacy.load("en_core_web_sm")
        except Exception:
            try:
                self._spacy_model = spacy.blank("en")
            except Exception:
                self._spacy_model = None
        return self._spacy_model

    def _relations(self, sentences: list[str]) -> list[Relation]:
        relations = []
        common_verbs = {
            "is", "are", "was", "were", "means", "shows", "uses", "forms",
            "creates", "produces", "explains", "changes", "moves", "affects", "reveals",
        }

        for sentence in sentences:
            tokens = self._tokens(sentence)
            if len(tokens) < 3:
                continue
            subject = tokens[0].title()
            verb = ""
            object_tokens = []
            for index, token in enumerate(tokens[1:], start=1):
                lowered = token.lower()
                if lowered in common_verbs or lowered.endswith(("s", "ed", "ing")):
                    verb = lowered
                    object_tokens = tokens[index + 1:]
                    break
            if not verb:
                continue
            object_text = " ".join(object_tokens[:5]).strip().title()
            if not object_text:
                continue
            relations.append(Relation(subject=subject, verb=verb, object=object_text))
            if len(relations) == 5:
                break

        return relations

    def _scenes(
        self, *, title: str, subject: str, concepts: list[Concept],
        relations: list[Relation], sentences: list[str], theme_title: str = "",
    ) -> list[Scene]:
        cast = self._story_cast(subject, concepts, sentences, theme_title)
        backgrounds = BACKGROUND_LIBRARY[subject]
        if theme_title:
            ordered_concepts = [concept for concept in concepts if concept.kind == "entity"] + [
                concept for concept in concepts if concept.kind != "entity"
            ]
            focus_concepts = ordered_concepts[:4] or [Concept(label=subject.title(), score=1.0, kind="concept")]
        else:
            focus_concepts = concepts[:4] or [Concept(label=subject.title(), score=1.0, kind="concept")]

        intro = Scene(
            title="Hook",
            speaker=cast[0][0], role=cast[0][1],
            dialogue=f"I need to understand {focus_concepts[0].label.lower()} for {title}.",
            caption=self._sentence_snippet(sentences[0]),
            background=backgrounds[0], focus=focus_concepts[0].label,
        )
        develop_one = Scene(
            title="Key Idea",
            speaker=cast[1][0], role=cast[1][1],
            dialogue=f"Start with {focus_concepts[0].label.lower()}: it anchors the whole lesson.",
            caption=self._sentence_snippet(sentences[min(1, len(sentences) - 1)]),
            background=backgrounds[1], focus=focus_concepts[0].label,
        )
        relation_focus = relations[0] if relations else None
        middle_dialogue = (
            f"{relation_focus.subject} {relation_focus.verb} {relation_focus.object.lower()}, which gives the story its next beat."
            if relation_focus and not theme_title
            else f"{focus_concepts[1].label if len(focus_concepts) > 1 else focus_concepts[0].label} connects to the main idea."
        )
        develop_two = Scene(
            title="Connection",
            speaker=cast[2][0], role=cast[2][1],
            dialogue=middle_dialogue,
            caption=self._sentence_snippet(sentences[min(2, len(sentences) - 1)]),
            background=backgrounds[2], focus=focus_concepts[min(1, len(focus_concepts) - 1)].label,
        )
        examples = []
        for index, concept in enumerate(focus_concepts[2:4], start=3):
            sentence_index = min(index, len(sentences) - 1)
            examples.append(Scene(
                title=f"Panel {index + 1}",
                speaker=cast[index % len(cast)][0], role=cast[index % len(cast)][1],
                dialogue=f"{concept.label} makes the lesson easier to remember with a concrete example.",
                caption=self._sentence_snippet(sentences[sentence_index]),
                background=backgrounds[index % len(backgrounds)], focus=concept.label,
            ))

        outro = Scene(
            title="Wrap Up",
            speaker=cast[3][0], role=cast[3][1],
            dialogue=f"Now the core ideas of {title} fit together like a comic sequence!",
            caption="Review the full strip, then download it for revision.",
            background=backgrounds[-1], focus=focus_concepts[0].label,
        )

        scenes = [intro, develop_one, develop_two, *examples, outro]
        return scenes[: min(max(len(scenes), 3), 6)]

    def _scenes_from_sections(
        self,
        *,
        title: str,
        subject: str,
        concepts: list[Concept],
        relations: list[Relation],
        sentences: list[str],
        sections: list[SourceSection],
        theme_title: str = "",
    ) -> list[Scene]:
        usable_sections = [section for section in sections if section.text.strip()]
        if not usable_sections:
            return self._scenes(
                title=title,
                subject=subject,
                concepts=concepts,
                relations=relations,
                sentences=sentences,
                theme_title=theme_title,
            )

        if len(usable_sections) > MAX_GENERATED_SCENES:
            usable_sections = self._compact_sections_for_scene_budget(usable_sections, MAX_GENERATED_SCENES)

        cast = self._story_cast(subject, concepts, sentences, theme_title)
        backgrounds = BACKGROUND_LIBRARY[subject]
        scenes: list[Scene] = []
        concept_fallbacks = concepts or [Concept(label=subject.title(), score=1.0, kind="concept")]

        for section in usable_sections:
            section_sentences = self._section_candidate_sentences(section.text)
            if not section_sentences:
                continue

            target = self._section_scene_target(section, section_sentences)
            remaining_budget = MAX_GENERATED_SCENES - len(scenes)
            if remaining_budget <= 0:
                break
            target = max(1, min(target, remaining_budget))
            selected_sentences = self._select_section_sentences(section_sentences, target)

            for local_index, sentence in enumerate(selected_sentences):
                global_index = len(scenes)
                focus = self._focus_for_sentence(sentence, concept_fallbacks, global_index)
                speaker, role, _kind = cast[global_index % len(cast)]
                mode = self._panel_mode(sentence, global_index, local_index, len(selected_sentences))
                scenes.append(
                    Scene(
                        title=self._section_scene_title(mode, global_index),
                        speaker=speaker,
                        role=role,
                        dialogue=self._scene_dialogue(
                            section=section,
                            sentence=sentence,
                            focus=focus,
                            title=title,
                            local_index=local_index,
                            section_total=len(selected_sentences),
                            global_index=global_index,
                            mode=mode,
                        ),
                        caption=self._scene_caption(sentence, focus, mode),
                        background=backgrounds[global_index % len(backgrounds)],
                        focus=focus,
                    )
                )

        if not scenes:
            return self._scenes(
                title=title,
                subject=subject,
                concepts=concepts,
                relations=relations,
                sentences=sentences,
                theme_title=theme_title,
            )

        while len(scenes) < 3 and sentences:
            sentence = sentences[min(len(scenes), len(sentences) - 1)]
            focus = self._focus_for_sentence(sentence, concept_fallbacks, len(scenes))
            speaker, role, _kind = cast[len(scenes) % len(cast)]
            mode = self._panel_mode(sentence, len(scenes), len(scenes), 3)
            scenes.append(
                Scene(
                    title=self._section_scene_title(mode, len(scenes)),
                    speaker=speaker,
                    role=role,
                    dialogue=self._scene_dialogue(
                        section=SourceSection(label="Review", text=sentence, source="Text", index=len(scenes) + 1),
                        sentence=sentence,
                        focus=focus,
                        title=title,
                        local_index=len(scenes),
                        section_total=3,
                        global_index=len(scenes),
                        mode=mode,
                    ),
                    caption=self._scene_caption(sentence, focus, mode),
                    background=backgrounds[len(scenes) % len(backgrounds)],
                    focus=focus,
                )
            )

        return scenes

    def _compact_sections_for_scene_budget(self, sections: list[SourceSection], budget: int) -> list[SourceSection]:
        if len(sections) <= budget:
            return sections
        group_size = max(1, math.ceil(len(sections) / budget))
        compacted: list[SourceSection] = []
        for group_start in range(0, len(sections), group_size):
            group = sections[group_start: group_start + group_size]
            if not group:
                continue
            first, last = group[0], group[-1]
            label = first.label if first.label == last.label else f"{first.label}-{last.label}"
            text = " ".join(self._sentence_snippet(section.text) for section in group if section.text)
            compacted.append(SourceSection(label=label, text=text, source=first.source, index=len(compacted) + 1))
        return compacted[:budget]

    def _section_candidate_sentences(self, text: str) -> list[str]:
        candidates: list[str] = []
        pending_heading = ""
        raw_lines = self._logical_content_lines(text)
        for raw_line in raw_lines:
            clean = self._clean_source_sentence(raw_line)
            if not clean:
                continue
            word_count = len(clean.split())
            is_bullet = bool(re.match(r"^\s*(?:[-*•·]+|\d+[\).]|[A-Za-z][\).])\s+", raw_line))
            if self._looks_like_heading(clean):
                pending_heading = clean.rstrip(":")
                continue

            if is_bullet or word_count <= 32:
                units = [clean]
            else:
                units = [self._clean_source_sentence(sentence) for sentence in self._sentences(clean)]

            for unit in units:
                if len(unit.split()) < 4:
                    continue
                if pending_heading and pending_heading.lower() not in unit.lower():
                    unit = f"{pending_heading}: {unit}"
                candidates.append(unit)

            if word_count > 16 or re.search(r"[.!?]$", clean):
                pending_heading = ""

        if candidates:
            return self._dedupe_sentence_candidates(candidates)[:60]

        candidates = [
            self._clean_source_sentence(sentence)
            for sentence in self._sentences(self._normalise_whitespace(text))
            if len(self._clean_source_sentence(sentence).split()) >= 4
        ]
        if candidates:
            return self._dedupe_sentence_candidates(candidates)[:60]

        shortened = self._clean_source_sentence(text)
        return [shortened] if shortened else []

    def _logical_content_lines(self, text: str) -> list[str]:
        lines: list[str] = []
        buffer = ""
        for raw_line in (text or "").splitlines():
            stripped = raw_line.strip()
            clean = self._clean_source_sentence(stripped)
            if not clean:
                continue
            is_bullet = bool(re.match(r"^\s*(?:[-*•·]+|\d+[\).]|[A-Za-z][\).])\s+", stripped))
            is_heading = self._looks_like_heading(clean)
            if is_heading or is_bullet:
                if buffer:
                    lines.append(buffer)
                    buffer = ""
                lines.append(stripped)
                continue

            buffer = f"{buffer} {clean}".strip() if buffer else clean
            if re.search(r"[.!?]$", clean):
                lines.append(buffer)
                buffer = ""

        if buffer:
            lines.append(buffer)
        return lines

    def _section_scene_target(self, section: SourceSection, sentences: list[str]) -> int:
        word_count = len(section.text.split())
        if section.source == "PDF":
            if word_count >= 700 or len(sentences) >= 5:
                return 4
            if word_count >= 320 or len(sentences) >= 3:
                return 3
            if word_count >= 220 or len(sentences) >= 2:
                return 2
            return 1
        if word_count >= 1100 or len(sentences) >= 14:
            return 4
        if word_count >= 620 or len(sentences) >= 9:
            return 3
        if word_count >= 260 or len(sentences) >= 5:
            return 2
        return 1

    def _select_section_sentences(self, sentences: list[str], target: int) -> list[str]:
        if len(sentences) <= target:
            return sentences

        selected_indexes: list[int] = []
        selected_keys: set[str] = set()

        def add(index: int) -> None:
            key = self._sentence_selection_key(sentences[index])
            if key and key in selected_keys:
                return
            selected_indexes.append(index)
            if key:
                selected_keys.add(key)

        if target >= 2:
            add(0)

        scored: list[tuple[int, int]] = []
        for index, sentence in enumerate(sentences):
            scored.append((self._sentence_importance_score(sentence, index, len(sentences)), index))

        for _score, index in sorted(scored, reverse=True):
            add(index)
            if len(selected_indexes) >= target:
                break

        if not selected_indexes and scored:
            selected_indexes.append(max(scored)[1])

        return [sentences[index] for index in sorted(selected_indexes)[:target]]

    def _dedupe_sentence_candidates(self, candidates: list[str]) -> list[str]:
        seen: set[str] = set()
        deduped: list[str] = []
        for candidate in candidates:
            key = self._sentence_selection_key(candidate)
            if key and key in seen:
                continue
            if key:
                seen.add(key)
            deduped.append(candidate)
        return deduped

    def _sentence_importance_score(self, sentence: str, index: int, total: int) -> int:
        lower = sentence.lower()
        score_terms = (
            "important", "main", "means", "because", "therefore", "example",
            "for example", "such as", "including", "first", "second", "third",
            "result", "cause", "effect", "however", "finally", "conclusion",
            "definition", "explains", "advantage", "disadvantage", "step",
            "process", "feature", "quality", "used to", "helps", "must",
        )
        score = min(len(sentence.split()), 44)
        score += sum(16 for term in score_terms if term in lower)
        if self._looks_like_heading(sentence.split(":", 1)[0]):
            score += 24
        if re.search(r"\b(?:is|are|means|refers to|represents|defined as)\b", lower):
            score += 22
        if re.search(r"^\s*(?:[-*•·]+|\d+[\).]|[A-Za-z][\).])\s+", sentence):
            score += 18
        score += len(re.findall(r"\b[A-Z]{2,}\b", sentence[:240])) * 8
        score += len(re.findall(r"\b[A-Z][a-z]+\b", sentence[:240])) * 2
        if re.search(r"\d", sentence):
            score += 8
        if index == 0:
            score += 10
        if index == total - 1:
            score += 6
        return score

    def _sentence_selection_key(self, sentence: str) -> str:
        tokens = [
            token.lower()
            for token in self._tokens(self._clean_source_sentence(sentence))
            if len(token) > 3 and token.lower() not in self._stopwords()
        ]
        return " ".join(tokens[:8])

    def _focus_for_sentence(self, sentence: str, concepts: list[Concept], scene_index: int) -> str:
        content = sentence.split(":", 1)[-1] if ":" in sentence else sentence
        lower = content.lower()
        priority_terms = (
            "furps", "software design", "software testing", "user interface design",
            "interface design", "maintainability", "usability", "modularity",
            "coupling", "cohesion", "abstraction", "encapsulation", "requirements",
        )
        ordered_terms = [
            *[term for term in priority_terms if term in TECH_TERM_TEACHING],
            *sorted(
                (term for term in TECH_TERM_TEACHING if term not in priority_terms),
                key=len,
                reverse=True,
            ),
        ]
        for term in ordered_terms:
            if re.search(rf"\b{re.escape(term)}\b", lower):
                return term.upper() if term in {"api", "furps"} else term.capitalize()
        for concept in concepts:
            if concept.label.lower() in lower:
                return concept.label
        words = [
            token
            for token in self._tokens(content)
            if len(token) > 3
            and token.lower() not in self._stopwords()
            and token.lower() not in {"should", "would", "could", "becomes", "because", "important", "using"}
        ]
        if words:
            counts = Counter(token.title() for token in words)
            return counts.most_common(1)[0][0]
        return concepts[scene_index % len(concepts)].label if concepts else "Main Idea"

    def _panel_mode(self, sentence: str, global_index: int, local_index: int, section_total: int) -> str:
        lower = sentence.lower()
        if any(term in lower for term in ("for example", "such as", "including", "instance")):
            return "example"
        if re.search(r"\b(?:is|are|means|refers to|represents|defined as)\b", lower):
            return "concept"
        if any(term in lower for term in ("because", "therefore", "so that", "important", "purpose", "helps")):
            return "why"
        if section_total > 1 and local_index == section_total - 1:
            return "recap"
        return PANEL_MODE_SEQUENCE[global_index % len(PANEL_MODE_SEQUENCE)]

    def _section_scene_title(self, mode: str, global_index: int) -> str:
        title_map = {
            "question": "Concept Intro",
            "concept": "Concept Intro",
            "analogy": "Think Like This",
            "why": "Why It Matters",
            "example": "Real Example",
            "tip": "Quick Tip",
            "connect": "Connect Ideas",
            "recap": "Final Idea",
        }
        return title_map.get(mode, EDUCATIONAL_FLOW_LABELS[global_index % len(EDUCATIONAL_FLOW_LABELS)])

    def _scene_dialogue(
        self,
        *,
        section: SourceSection,
        sentence: str,
        focus: str,
        title: str,
        local_index: int,
        section_total: int,
        global_index: int,
        mode: str,
    ) -> str:
        focus_name = self._friendly_focus(focus)
        core = self._core_teaching_sentence(sentence, focus_name)
        analogy = self._analogy_for_focus(focus_name, sentence)
        benefit = self._benefit_phrase(sentence, focus_name)
        memory = self._memory_cue(focus_name, sentence)

        if mode == "question":
            line = f"Ever wondered where {focus_name.lower()} fits? {core}"
        elif mode == "concept":
            line = core
        elif mode == "analogy":
            line = f"Think of {focus_name.lower()} like {analogy}. {core}"
        elif mode == "why":
            line = f"{focus_name} matters because {benefit}. {core}"
        elif mode == "example":
            line = f"Example time: {self._example_line(sentence, focus_name)}"
        elif mode == "tip":
            line = f"Quick memory trick: {memory}. {core}"
        elif mode == "connect":
            line = f"This connects to the next idea: {core}"
        else:
            line = f"Final idea: {core}"

        return self._finish_sentence(self._shorten_text(line, width=88))

    def _scene_caption(self, sentence: str, focus: str, mode: str) -> str:
        focus_name = self._friendly_focus(focus)
        core = self._core_teaching_sentence(sentence, focus_name)
        if mode == "why":
            caption = f"Why it sticks: {self._benefit_phrase(sentence, focus_name)}."
        elif mode == "example":
            caption = f"Try it with a small example: {self._example_line(sentence, focus_name)}"
        elif mode == "tip":
            caption = f"Memory cue: {self._memory_cue(focus_name, sentence)}."
        elif mode == "connect":
            caption = f"Connect it: ask what changes before and after {focus_name.lower()}."
        elif mode == "recap":
            caption = f"Takeaway: {core}"
        else:
            caption = f"Read it simply: {core}"
        return self._finish_sentence(self._shorten_text(caption, width=84))

    def _core_teaching_sentence(self, sentence: str, focus: str) -> str:
        clean = self._clean_source_sentence(sentence)
        term_line = self._term_teaching_line(clean, focus)
        if term_line:
            return term_line

        definition = self._definition_teaching_line(clean)
        if definition:
            return definition

        human = self._humanize_text(clean)
        if not human:
            human = f"{focus} is the idea to remember here."
        return self._finish_sentence(self._shorten_text(human, width=142))

    def _term_teaching_line(self, sentence: str, focus: str) -> str:
        lower = f"{sentence} {focus}".lower()
        focus_lower = focus.lower()
        content_lower = sentence.lower().split(":", 1)[-1]
        priority_terms = (
            "furps", "software design", "software testing", "user interface design",
            "interface design", "maintainability", "usability", "reliability",
            "performance", "supportability", "functionality",
        )
        ordered_terms = [
            *[term for term in priority_terms if term in TECH_TERM_TEACHING],
            *sorted(
                (term for term in TECH_TERM_TEACHING if term not in priority_terms),
                key=len,
                reverse=True,
            ),
        ]

        for search_area in (focus_lower, content_lower, lower):
            for term in ordered_terms:
                if re.search(rf"\b{re.escape(term)}\b", search_area):
                    display = term.upper() if term in {"api", "furps"} else term.capitalize()
                    verb = "are" if term.endswith("s") and term not in {"furps"} else "is"
                    return f"{display} {verb} {TECH_TERM_TEACHING[term]}."
        return ""

    def _definition_teaching_line(self, sentence: str) -> str:
        candidate = re.sub(r"^[A-Z][A-Za-z0-9 /&()_-]{1,70}:\s*", "", sentence.strip())
        match = re.match(
            r"^([A-Za-z][A-Za-z0-9 /&()_-]{1,70}?)\s+(?:is|are|means|refers to|represents|is defined as|can be defined as)\s+(.{8,})",
            candidate,
            flags=re.IGNORECASE,
        )
        if not match:
            return ""
        subject = self._clean_source_sentence(match.group(1)).strip(" :-")
        detail = self._humanize_text(match.group(2)).strip(" .")
        if not subject or len(subject.split()) > 8 or subject.lower().endswith((" that", " which")) or not detail:
            return ""
        return self._finish_sentence(f"{subject} simply means {self._shorten_text(detail, width=104)}")

    def _humanize_text(self, sentence: str) -> str:
        clean = self._clean_source_sentence(sentence)
        clean = re.sub(r"\bcreates?\s+a\s+representation\s+of\b", "turns ideas into a clear model of", clean, flags=re.IGNORECASE)
        clean = re.sub(r"\bcan\s+be\s+used\s+to\b", "helps you", clean, flags=re.IGNORECASE)
        clean = re.sub(r"\bis\s+used\s+to\b", "helps you", clean, flags=re.IGNORECASE)
        for pattern, replacement in JARGON_REPLACEMENTS:
            clean = re.sub(pattern, replacement, clean, flags=re.IGNORECASE)
        clean = re.sub(r"\s+([,.;:!?])", r"\1", clean)
        clean = re.sub(r"\s+", " ", clean).strip()
        return self._finish_sentence(self._shorten_text(clean, width=142)) if clean else ""

    def _clean_source_sentence(self, sentence: str) -> str:
        clean = self._normalise_whitespace(sentence)
        clean = re.sub(r"^\s*(?:[-*•·]+|\d+[\).]|[A-Za-z][\).])\s+", "", clean)
        clean = re.sub(r"\bPage\s+\d+(?:\s+of\s+\d+)?\b", "", clean, flags=re.IGNORECASE)
        clean = re.sub(r"\b(?:Figure|Fig\.|Table)\s+\d+(?:\.\d+)?\b", "", clean, flags=re.IGNORECASE)
        clean = re.sub(r"\[[^\]]{1,20}\]", "", clean)
        clean = re.sub(r"https?://\S+", "", clean)
        clean = re.sub(r"\s+([,.;:!?])", r"\1", clean)
        return self._normalise_whitespace(clean).strip(" -")

    def _looks_like_heading(self, text: str) -> bool:
        clean = self._clean_source_sentence(text).strip(":")
        words = clean.split()
        if len(words) < 2 or len(words) > 12:
            return False
        if re.search(r"[.!?]$", clean):
            return False
        if sum(1 for word in words if len(word) > 2) < 2:
            return False
        titleish = sum(1 for word in words if word[:1].isupper() or word.isupper())
        return titleish >= max(1, len(words) // 2) or clean.isupper()

    def _friendly_focus(self, focus: str) -> str:
        cleaned = self._clean_source_sentence(focus).strip(" .:;")
        words = cleaned.split()
        if not words:
            return "this idea"
        if len(words) > 4:
            return "this idea"
        joined = " ".join(words)
        return joined.upper() if joined.lower() in {"api", "furps"} else joined[:1].upper() + joined[1:]

    def _analogy_for_focus(self, focus: str, sentence: str) -> str:
        lower = f"{focus} {sentence}".lower()
        if any(term in lower for term in ("software design", "architecture", "design")):
            return "a blueprint before building a house"
        if "algorithm" in lower:
            return "a recipe that tells you each step in order"
        if "database" in lower:
            return "a labelled cupboard where information is easy to find"
        if "requirement" in lower:
            return "a project checklist before anyone starts building"
        if "testing" in lower:
            return "practice questions before the final exam"
        if "model" in lower:
            return "a map that makes a big topic easier to explore"
        if any(term in lower for term in ("quality", "furps", "performance", "usability", "reliability")):
            return "a report card for how well software works"
        return "a sticky note that keeps the main idea visible"

    def _benefit_phrase(self, sentence: str, focus: str) -> str:
        clean = self._without_terminal_period(self._humanize_text(sentence))
        lower = clean.lower()
        if "because" in lower:
            return self._shorten_text(clean[lower.index("because") + len("because"):].strip(" ,."), width=86)
        if "so that" in lower:
            return self._shorten_text(clean[lower.index("so that") + len("so that"):].strip(" ,."), width=86)
        if "helps" in lower:
            phrase = self._shorten_text(clean[lower.index("helps"):].strip(" ,."), width=86)
            return phrase if phrase.lower().startswith("it ") else f"it {phrase}"
        term_line = self._term_teaching_line(sentence, focus)
        if term_line:
            return "it gives you a quick way to judge the topic"
        return "it tells you what to notice, remember, and use later"

    def _example_line(self, sentence: str, focus: str) -> str:
        raw = self._clean_source_sentence(sentence)
        raw_lower = raw.lower()
        for marker in ("for example", "such as", "including", "instance"):
            if marker in raw_lower:
                example = raw[raw_lower.index(marker):].strip(" ,:")
                return self._finish_sentence(self._shorten_text(self._humanize_text(example), width=126))
        clean = self._without_terminal_period(self._humanize_text(sentence))
        lower = clean.lower()
        if "software" in lower or "app" in lower:
            return f"imagine a class app where {focus.lower()} helps the team make cleaner choices."
        return self._finish_sentence(self._shorten_text(clean, width=126))

    def _memory_cue(self, focus: str, sentence: str) -> str:
        if re.search(r"\bFURPS\b", sentence, flags=re.IGNORECASE):
            return "FURPS equals five checks for software quality"
        acronym = re.search(r"\b[A-Z]{3,}\b", sentence)
        if acronym:
            return f"{acronym.group(0)} is your shortcut word for this part"
        return f"link {focus.lower()} with what it does, why it matters, and one example"

    def _shorten_text(self, text: str, *, width: int) -> str:
        clean = self._normalise_whitespace(text)
        if len(clean) <= width:
            return clean
        return textwrap.shorten(clean, width=width, placeholder="...")

    @staticmethod
    def _without_terminal_period(text: str) -> str:
        clean = text.strip()
        if clean.endswith(".") and not clean.endswith("..."):
            return clean[:-1]
        return clean

    @staticmethod
    def _finish_sentence(text: str) -> str:
        clean = text.strip()
        if not clean:
            return clean
        if clean.endswith(("...", ".", "!", "?")):
            return clean
        return f"{clean}."

    @staticmethod
    def _sentence_snippet(sentence: str) -> str:
        return textwrap.shorten(CominoteEngine._normalise_whitespace(sentence), width=86, placeholder="...")

    def _story_cast(
        self,
        subject: str,
        concepts: list[Concept],
        sentences: list[str],
        theme_title: str = "",
    ) -> list[tuple[str, str, str]]:
        base_cast = CAST_LIBRARY[subject]
        if not theme_title:
            return base_cast

        candidates: list[str] = []
        title_blob = theme_title.lower()
        for concept in concepts:
            if concept.kind != "entity":
                continue
            label = concept.label.strip()
            lowered = label.lower()
            if not label or lowered in GENERIC_ENTITY_STOPWORDS or lowered == title_blob:
                continue
            if lowered in {item.lower() for item in candidates}:
                continue
            if len(label) > 22 or re.fullmatch(r"[ivxlcdm]+", lowered):
                continue
            candidates.append(label)
            if len(candidates) == 3:
                break

        if len(candidates) < 3:
            for entity in self._entities(" ".join(sentences[:10])):
                label = self._normalise_whitespace(entity)
                lowered = label.lower()
                if not label or lowered in GENERIC_ENTITY_STOPWORDS or lowered == title_blob:
                    continue
                if any(lowered == existing.lower() for existing in candidates):
                    continue
                if len(label) > 22 or re.fullmatch(r"[ivxlcdm]+", lowered):
                    continue
                candidates.append(label)
                if len(candidates) == 3:
                    break

        themed_cast = list(base_cast)
        for index, name in enumerate(candidates[:3]):
            base_name, base_role, base_kind = base_cast[index]
            themed_cast[index] = (textwrap.shorten(name, width=18, placeholder=""), base_role, base_kind)
        return themed_cast

    def _load_theme_datasets(self) -> None:
        self.theme_dir = self.base_dir / "JSON"
        self._theme_datasets: list[dict[str, Any]] = []
        self._theme_datasets_by_slug: dict[str, dict[str, Any]] = {}
        if not self.theme_dir.exists():
            self._load_virtual_theme_datasets()
            return

        for path in sorted(self.theme_dir.glob("*.json")):
            theme = self._load_theme_dataset(path)
            if theme is None:
                continue
            self._theme_datasets.append(theme)
            self._theme_datasets_by_slug[theme["slug"]] = theme
        self._load_virtual_theme_datasets()

    def _load_theme_dataset(self, path: Path) -> dict[str, Any] | None:
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            return None

        segments: list[str] = []
        if isinstance(payload, dict):
            ordered_items = sorted(payload.items(), key=lambda item: self._theme_sort_key(item[0]))
            for _, value in ordered_items:
                text = self._coerce_theme_text(value)
                text = self._clean_theme_text(text)
                if text:
                    segments.append(text)
        else:
            text = self._clean_theme_text(self._coerce_theme_text(payload))
            if text:
                segments.append(text)

        segments = [segment for segment in segments if len(segment.split()) >= 8]
        if not segments:
            return None

        title = self._format_theme_title(path.stem)
        source_text = "\n\n".join(segments)
        subject = self._infer_theme_subject(title, source_text)
        recommended_style = self._infer_theme_style(title, subject, source_text)
        theme_profile = self._resolve_theme_profile(theme_meta=None, requested_style=recommended_style, title=title, source_text=source_text)
        theme_render = self._resolve_theme_render(
            title=title,
            source_text=source_text,
            theme_profile_slug=str(theme_profile.get("slug") or "cartoon_kids"),
        )
        render_profile_slug = str(theme_render.get("theme_profile") or "")
        if render_profile_slug in self._theme_profiles_by_slug:
            theme_profile = self._default_theme_profile(render_profile_slug)
        featured_names = [
            str(item.get("name") or "").strip()
            for item in (theme_render.get("featured_cast") or [])
            if isinstance(item, dict) and str(item.get("name") or "").strip()
        ]
        featured_background = theme_render.get("featured_background") if isinstance(theme_render.get("featured_background"), dict) else {}
        preview = textwrap.shorten(segments[0], width=240, placeholder="...")
        word_count = sum(len(segment.split()) for segment in segments)
        return {
            "slug": self._slugify(title),
            "title": title,
            "subject": subject,
            "recommended_style": recommended_style,
            "theme_profile": theme_profile["slug"],
            "theme_profile_label": theme_profile["label"],
            "theme_character_hint": ", ".join(featured_names) or theme_profile.get("character_summary", ""),
            "theme_background_hint": str(featured_background.get("name") or theme_profile.get("background_summary", "")),
            "font_label": theme_profile.get("font_label", ""),
            "layout_label": theme_profile.get("layout_label", ""),
            "ui_display_font": theme_profile.get("ui_display_font", "\"Bangers\", cursive"),
            "ui_body_font": theme_profile.get("ui_body_font", "\"Comic Neue\", cursive"),
            "theme_accent": (theme_render.get("ui") or {}).get("accent", theme_profile.get("accent", "#f4631e")),
            "theme_surface": (theme_render.get("ui") or {}).get("surface", theme_profile.get("surface", "#ffffff")),
            "theme_render": theme_render,
            "preview": preview,
            "word_count": word_count,
            "segments": segments,
        }

    def _load_virtual_theme_datasets(self) -> None:
        for entry in self._theme_render_entries:
            if not isinstance(entry, dict):
                continue
            label = str(entry.get("label") or "").strip()
            if not label:
                continue
            slug = self._slugify(label)
            if not slug or slug in self._theme_datasets_by_slug:
                continue
            theme = self._build_virtual_theme_dataset(entry)
            if theme is None:
                continue
            self._theme_datasets.append(theme)
            self._theme_datasets_by_slug[slug] = theme

    def _build_virtual_theme_dataset(self, entry: dict[str, Any]) -> dict[str, Any] | None:
        label = str(entry.get("label") or "").strip()
        if not label:
            return None

        preferred_profile = str(entry.get("theme_profile") or "cartoon_kids")
        theme_render = self._resolve_theme_render(
            title=label,
            source_text=str(entry.get("headline") or ""),
            theme_profile_slug=preferred_profile,
        )
        featured_names = [
            str(item.get("name") or "").strip()
            for item in (theme_render.get("featured_cast") or [])
            if isinstance(item, dict) and str(item.get("name") or "").strip()
        ]
        featured_background = (
            theme_render.get("featured_background")
            if isinstance(theme_render.get("featured_background"), dict)
            else {}
        )
        segments = self._build_virtual_theme_segments(theme_render)
        if not segments:
            return None

        source_text = "\n\n".join(segments)
        subject = self._infer_theme_subject(label, source_text)
        recommended_style = self._infer_theme_style(label, subject, source_text)
        theme_profile = self._resolve_theme_profile(
            theme_meta=None,
            requested_style=recommended_style,
            title=label,
            source_text=source_text,
        )
        render_profile_slug = str(theme_render.get("theme_profile") or "")
        if render_profile_slug in self._theme_profiles_by_slug:
            theme_profile = self._default_theme_profile(render_profile_slug)

        preview_source = str(theme_render.get("headline") or segments[0])
        preview = textwrap.shorten(preview_source, width=240, placeholder="...")
        word_count = sum(len(segment.split()) for segment in segments)
        return {
            "slug": self._slugify(label),
            "title": label,
            "subject": subject,
            "recommended_style": recommended_style,
            "theme_profile": theme_profile["slug"],
            "theme_profile_label": theme_profile["label"],
            "theme_character_hint": ", ".join(featured_names) or theme_profile.get("character_summary", ""),
            "theme_background_hint": str(
                featured_background.get("name") or theme_profile.get("background_summary", "")
            ),
            "font_label": theme_profile.get("font_label", ""),
            "layout_label": theme_profile.get("layout_label", ""),
            "ui_display_font": theme_profile.get("ui_display_font", "\"Bangers\", cursive"),
            "ui_body_font": theme_profile.get("ui_body_font", "\"Comic Neue\", cursive"),
            "theme_accent": (theme_render.get("ui") or {}).get("accent", theme_profile.get("accent", "#f4631e")),
            "theme_surface": (theme_render.get("ui") or {}).get("surface", theme_profile.get("surface", "#ffffff")),
            "theme_render": theme_render,
            "preview": preview,
            "word_count": word_count,
            "segments": segments,
            "is_virtual_theme": True,
        }

    def _build_virtual_theme_segments(self, theme_render: dict[str, Any]) -> list[str]:
        label = str(theme_render.get("label") or "Theme").strip() or "Theme"
        cast = [
            item
            for item in (theme_render.get("featured_cast") or [])
            if isinstance(item, dict) and str(item.get("name") or "").strip()
        ]
        cast_names = ", ".join(str(item.get("name") or "").strip() for item in cast[:3])
        comic_elements = [
            str(item).strip()
            for item in (theme_render.get("comic_elements") or [])
            if str(item).strip()
        ]
        background = theme_render.get("featured_background") if isinstance(theme_render.get("featured_background"), dict) else {}
        background_name = str(background.get("name") or "Comic Stage").strip() or "Comic Stage"
        background_description = str(background.get("description") or "A signature comic environment designed for this theme.").strip()
        headline = str(
            theme_render.get("headline")
            or f"{label} swaps in its own featured cast, background, and comic energy."
        ).strip()
        badge_text = str(theme_render.get("badge_text") or "Theme Pulse").strip() or "Theme Pulse"

        segments = [
            self._clean_theme_text(
                f"{label} theme overview. {headline} The dashboard badge reads {badge_text}. "
                f"The featured cast includes {cast_names or 'signature comic heroes'}."
            ),
            self._clean_theme_text(
                f"The signature setting is {background_name}. {background_description} "
                f"Comic panels emphasize {', '.join(comic_elements[:4]) or 'dynamic caption boxes and energetic panel framing'}."
            ),
        ]

        for member in cast[:3]:
            name = str(member.get("name") or "").strip()
            description = str(member.get("description") or "").strip()
            costume = member.get("costume") if isinstance(member.get("costume"), dict) else {}
            accessory = str(costume.get("accessory") or "comic signature gear").strip()
            if not name:
                continue
            segments.append(
                self._clean_theme_text(
                    f"{name} appears as a featured {label} character. "
                    f"{description or 'The character carries the signature mood of the selected theme.'} "
                    f"Signature detail: {accessory}."
                )
            )

        return [segment for segment in segments if len(segment.split()) >= 8]

    @staticmethod
    def _theme_sort_key(key: Any) -> tuple[int, Any]:
        text = str(key)
        if re.fullmatch(r"-?\d+", text):
            return (0, int(text))
        return (1, text.lower())

    @staticmethod
    def _format_theme_title(raw_name: str) -> str:
        cleaned = re.sub(r"[_-]+", " ", raw_name).strip()
        cleaned = re.sub(r"(?<=[a-z0-9])(?=[A-Z])", " ", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        if not cleaned:
            return "Theme"

        small_words = {"a", "an", "and", "as", "at", "by", "for", "from", "in", "of", "on", "or", "the", "to", "with"}

        def title_token(word: str, index: int) -> str:
            if word.isupper() and len(word) > 1:
                return word

            lowered = word.lower()
            if index > 0 and lowered in small_words:
                return lowered

            def capitalise(match: re.Match[str]) -> str:
                token = match.group(0)
                return token[:1].upper() + token[1:].lower()

            return re.sub(r"[A-Za-zÀ-ÿ]+(?:['’][A-Za-zÀ-ÿ]+)?", capitalise, lowered)

        return " ".join(title_token(word, index) for index, word in enumerate(cleaned.split()))

    def _coerce_theme_text(self, value: Any) -> str:
        if isinstance(value, str):
            return value
        if isinstance(value, list):
            flattened = [self._coerce_theme_text(item) for item in value]
            flattened = [item for item in flattened if item]
            if not flattened:
                return ""
            average_len = sum(len(item) for item in flattened) / max(len(flattened), 1)
            separator = " " if average_len <= 24 else "\n"
            return separator.join(flattened)
        if isinstance(value, dict):
            ordered_items = sorted(value.items(), key=lambda item: self._theme_sort_key(item[0]))
            return "\n".join(
                text for _, item_value in ordered_items
                if (text := self._coerce_theme_text(item_value))
            )
        return str(value).strip()

    def _clean_theme_text(self, text: str) -> str:
        if not text:
            return ""
        cleaned = text.replace("\ufeff", " ").replace("\xa0", " ").replace("\u000c", " ")
        cleaned = re.sub(r"\s+([,.;:!?])", r"\1", cleaned)
        cleaned = re.sub(r"([(\[{])\s+", r"\1", cleaned)
        cleaned = re.sub(r"\s+([)\]}])", r"\1", cleaned)
        cleaned = cleaned.replace(" n't", "n't").replace(" 's", "'s").replace(" ’ ", "’")
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        cleaned = re.sub(
            r"\b(?:copyright|contents|title page|image credits|acknowledgements|about the author)\b",
            " ",
            cleaned,
            flags=re.IGNORECASE,
        )
        return self._normalise_whitespace(cleaned)

    def _infer_theme_subject(self, title: str, source_text: str) -> str:
        title_blob = title.lower()
        for subject, hints in TITLE_SUBJECT_HINTS.items():
            if any(hint in title_blob for hint in hints):
                return subject

        sample = f"{title} {source_text[:12000]}".lower()
        best_subject = "literature"
        best_score = -1
        for subject, keywords in SUBJECT_KEYWORDS.items():
            score = sum(1 for keyword in keywords if keyword in sample)
            if score > best_score:
                best_subject = subject
                best_score = score
        return best_subject if best_score > 0 else "literature"

    def _infer_theme_style(self, title: str, subject: str, source_text: str) -> str:
        title_blob = title.lower()
        for style_slug, hints in TITLE_STYLE_HINTS.items():
            if any(hint in title_blob for hint in hints):
                return style_slug
        sample = f"{title_blob} {source_text[:4000].lower()}"
        if any(token in sample for token in ("manga", "black white", "black-and-white", "ink", "speed line")):
            return "manga_story"
        if any(token in sample for token in ("fantasy", "wizard", "magic", "magical", "castle", "dragon", "ring")):
            return "fantasy_adventure"
        if any(token in sample for token in ("pixar", "3d", "rounded", "soft")):
            return "pixar_colorful"
        if any(token in sample for token in ("anime", "academy", "shonen", "shojo", "tokyo")):
            return "anime_school"
        if any(token in sample for token in ("hero", "marvel", "mission", "spy")):
            return "superhero_comic"
        if any(token in sample for token in ("magic", "wonderland", "animal", "forest", "storybook")):
            return "fantasy_adventure" if "wonderland" in sample else "cute_cartoon"
        return SUBJECT_STYLE_DEFAULTS.get(subject, "anime_school")

    def _theme_excerpt(self, theme: dict[str, Any], guidance_text: str = "") -> str:
        segments = theme.get("segments") or []
        if not segments:
            return ""

        trimmed_segments = [textwrap.shorten(segment, width=2800, placeholder="...") for segment in segments]
        if not guidance_text.strip():
            ranked = sorted(
                ((self._theme_segment_score(segment), -index, segment) for index, segment in enumerate(trimmed_segments)),
                reverse=True,
            )
            best_segments = [segment for score, _, segment in ranked if score > 0][:3]
            return "\n\n".join(best_segments or trimmed_segments[:3])

        guide_tokens = {token.lower() for token in self._tokens(guidance_text) if len(token) > 2}
        if not guide_tokens:
            ranked = sorted(
                ((self._theme_segment_score(segment), -index, segment) for index, segment in enumerate(trimmed_segments)),
                reverse=True,
            )
            best_segments = [segment for score, _, segment in ranked if score > 0][:3]
            return "\n\n".join(best_segments or trimmed_segments[:3])

        scored: list[tuple[int, int, str]] = []
        for index, segment in enumerate(trimmed_segments):
            segment_tokens = {token.lower() for token in self._tokens(segment[:1800]) if len(token) > 2}
            overlap = len(guide_tokens & segment_tokens) * 8
            overlap += self._theme_segment_score(segment)
            scored.append((overlap, -index, segment))

        best_segments = [segment for score, _, segment in sorted(scored, reverse=True) if score > 0][:3]
        if not best_segments:
            best_segments = trimmed_segments[:3]
        return "\n\n".join(best_segments)

    def _theme_segment_score(self, segment: str) -> int:
        lower = segment.lower()
        heading_hits = sum(lower.count(token) for token in (
            "contents", "chapter", "copyright", "bibliography", "image credits",
            "acknowledgements", "title page", "introduction", "about the author",
        ))
        sentence_count = len(self._sentences(segment[:1600]))
        dialogue_hits = segment.count('"') + segment.count("'")
        named_hits = len(re.findall(r"\b[A-Z][a-z]+\b", segment[:1200]))
        word_count = len(segment.split())
        return (sentence_count * 5) + min(word_count // 18, 24) + min(dialogue_hits, 6) + min(named_hits, 12) - (heading_hits * 8)

    @staticmethod
    def _slugify(value: str) -> str:
        return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")

    def _load_dataset_assets(self) -> None:
        self.dataset_dir = self.base_dir / "dataset"
        self.archive_dir = self.base_dir / "archive"
        profile_data = self._load_dataset_json("theme_asset_profiles.json")
        profiles = profile_data.get("profiles") if isinstance(profile_data, dict) else []
        self._theme_profiles = profiles if isinstance(profiles, list) else []
        self._theme_profiles_by_slug = {
            profile.get("slug"): profile
            for profile in self._theme_profiles
            if isinstance(profile, dict) and profile.get("slug")
        }
        keyword_rules = profile_data.get("keyword_rules") if isinstance(profile_data, dict) else []
        self._theme_keyword_rules = keyword_rules if isinstance(keyword_rules, list) else []
        style_rules = profile_data.get("style_rules") if isinstance(profile_data, dict) else {}
        self._theme_style_rules = style_rules if isinstance(style_rules, dict) else {}
        render_data = self._load_dataset_json("theme_character_manifest.json")
        self._theme_render_default = render_data.get("default") if isinstance(render_data, dict) else {}
        render_entries = render_data.get("themes") if isinstance(render_data, dict) else []
        self._theme_render_entries = render_entries if isinstance(render_entries, list) else []
        self._dataset_characters = self._load_dataset_records("characters.json", "characters")
        self._dataset_backgrounds = self._load_dataset_records("backgrounds.json", "backgrounds")
        self._dataset_emotions = self._load_dataset_records("emotions.json", "emotions")
        self._dataset_poses = self._load_dataset_records("poses.json", "poses")
        panel_data = self._load_dataset_json("panels_text_styles.json")
        text_elements = panel_data.get("text_elements", {})
        self._dataset_layouts = panel_data.get("panel_layouts", [])
        self._dataset_layouts_by_id = {item.get("id"): item for item in self._dataset_layouts if item.get("id")}
        self._dataset_bubbles = text_elements.get("speech_bubbles", [])
        self._dataset_bubbles_by_type = {item.get("type"): item for item in self._dataset_bubbles if item.get("type")}
        self._dataset_sfx = text_elements.get("sound_effects", [])
        self._dataset_sfx_by_word = {item.get("word"): item for item in self._dataset_sfx if item.get("word")}
        self._dataset_styles = panel_data.get("styles", [])
        self._dataset_styles_by_name = {item.get("name"): item for item in self._dataset_styles if item.get("name")}
        self._archive_characters = self._load_archive_characters()

    def _load_dataset_json(self, filename: str) -> dict[str, Any]:
        path = self.dataset_dir / filename
        if not path.exists():
            return {}
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}

    def _load_dataset_records(self, filename: str, key: str) -> list[dict[str, Any]]:
        data = self._load_dataset_json(filename)
        records = data.get(key)
        return records if isinstance(records, list) else []

    @staticmethod
    def _merge_dicts(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
        merged = dict(base)
        for key, value in override.items():
            if isinstance(value, dict) and isinstance(merged.get(key), dict):
                merged[key] = CominoteEngine._merge_dicts(merged[key], value)
            else:
                merged[key] = value
        return merged

    def _theme_render_ui_defaults(self, theme_profile_slug: str) -> dict[str, str]:
        defaults = {
            "anime": {
                "accent": "#ff5fa2",
                "accent_2": "#6c5ce7",
                "accent_3": "#00b4d8",
                "surface": "#fff4fb",
                "bg_start": "#fff5fb",
                "bg_end": "#e7eeff",
                "glow": "rgba(255, 95, 162, 0.2)",
            },
            "pixar": {
                "accent": "#ff8744",
                "accent_2": "#20b2aa",
                "accent_3": "#5e60ce",
                "surface": "#fff7ef",
                "bg_start": "#fff6ed",
                "bg_end": "#eefbff",
                "glow": "rgba(255, 135, 68, 0.18)",
            },
            "superhero": {
                "accent": "#ea233f",
                "accent_2": "#0f52ba",
                "accent_3": "#ffd447",
                "surface": "#fff1ef",
                "bg_start": "#fff1ef",
                "bg_end": "#eef3ff",
                "glow": "rgba(234, 35, 63, 0.18)",
            },
            "cartoon_kids": {
                "accent": "#fb5607",
                "accent_2": "#ffb703",
                "accent_3": "#2ec4b6",
                "surface": "#fff8dd",
                "bg_start": "#fff8dd",
                "bg_end": "#edfff2",
                "glow": "rgba(255, 183, 3, 0.2)",
            },
            "manga": {
                "accent": "#232323",
                "accent_2": "#575757",
                "accent_3": "#a0a0a0",
                "surface": "#f2f2f2",
                "bg_start": "#fafafa",
                "bg_end": "#dcdcdc",
                "glow": "rgba(35, 35, 35, 0.1)",
            },
            "fantasy": {
                "accent": "#7d4dff",
                "accent_2": "#2dd4bf",
                "accent_3": "#facc15",
                "surface": "#f8f2ff",
                "bg_start": "#f8f2ff",
                "bg_end": "#ecfbf7",
                "glow": "rgba(125, 77, 255, 0.16)",
            },
            "horror": {
                "accent": "#4c1d95",
                "accent_2": "#dc2626",
                "accent_3": "#f8fafc",
                "surface": "#f3f0ff",
                "bg_start": "#ddd6fe",
                "bg_end": "#111827",
                "glow": "rgba(76, 29, 149, 0.22)",
            },
        }
        return dict(defaults.get(theme_profile_slug, defaults["cartoon_kids"]))

    def _theme_render_comic_elements(self, theme_profile_slug: str) -> list[str]:
        defaults = {
            "anime": ["speed lines", "spark bursts", "hero captions"],
            "pixar": ["soft glows", "rounded bubbles", "story cards"],
            "superhero": ["impact bursts", "city motion", "bold captions"],
            "cartoon_kids": ["halftone dots", "playful bubbles", "bright stickers"],
            "manga": ["ink speed lines", "dramatic boxes", "tone textures"],
            "fantasy": ["glow runes", "quest frames", "magic trails"],
            "horror": ["cinematic fog", "spotlight shadows", "mystery glows"],
        }
        return list(defaults.get(theme_profile_slug, defaults["cartoon_kids"]))

    def _default_theme_render(self, theme_profile_slug: str = "cartoon_kids") -> dict[str, Any]:
        profile = self._default_theme_profile(theme_profile_slug)
        ui = self._theme_render_ui_defaults(str(profile.get("slug") or theme_profile_slug))
        ui["accent"] = self._safe_hex(profile.get("accent"), ui["accent"])
        ui["surface"] = self._safe_hex(profile.get("surface"), ui["surface"])
        base = self._theme_render_default if isinstance(self._theme_render_default, dict) else {}
        return self._merge_dicts(
            base,
            {
                "id": str(base.get("id") or "theme-explorer"),
                "label": str(base.get("label") or profile.get("label") or "Theme Explorer"),
                "badge_text": str(base.get("badge_text") or "Comic Preview"),
                "headline": str(base.get("headline") or profile.get("description") or "Select a theme to preview it."),
                "sound_effect": str(base.get("sound_effect") or "POP!"),
                "theme_profile": str(profile.get("slug") or theme_profile_slug),
                "ui": ui,
                "comic_elements": list(base.get("comic_elements") or self._theme_render_comic_elements(str(profile.get("slug") or theme_profile_slug))),
            },
        )

    def _theme_render_match_score(self, entry: dict[str, Any], blob: str) -> int:
        score = 0
        keywords = [str(keyword).strip().lower() for keyword in (entry.get("keywords") or []) if str(keyword).strip()]
        for keyword in keywords:
            if keyword and keyword in blob:
                score += max(3, len(keyword.split()) * 6)
        label = str(entry.get("label") or "").strip().lower()
        if label and label in blob:
            score += max(10, len(label.split()) * 8)
        return score

    def _exact_theme_render_entry(self, *, theme_meta: dict[str, Any] | None = None, title: str = "") -> dict[str, Any] | None:
        keys = {
            self._slugify(str(value))
            for value in (
                title,
                (theme_meta or {}).get("slug", ""),
                (theme_meta or {}).get("title", ""),
            )
            if str(value).strip()
        }
        if not keys:
            return None
        for entry in self._theme_render_entries:
            if not isinstance(entry, dict):
                continue
            entry_keys = {
                self._slugify(str(entry.get("id") or "")),
                self._slugify(str(entry.get("label") or "")),
                *[self._slugify(str(keyword)) for keyword in (entry.get("keywords") or [])],
            }
            if keys & {key for key in entry_keys if key}:
                return entry
        return None

    def _pick_theme_anchor(self, title: str) -> str:
        tokens = [
            token
            for token in re.findall(r"[A-Za-z0-9]+", title)
            if len(token) > 2 and token.lower() not in FALLBACK_STOPWORDS
        ]
        generic = {"theme", "story", "comic", "book", "journey", "great", "lord", "west", "things"}
        refined = [token for token in tokens if token.lower() not in generic] or tokens
        if not refined:
            return "Theme"
        return max(refined, key=len).title()

    def _synthesise_theme_render(self, title: str, theme_profile_slug: str) -> dict[str, Any]:
        profile_slug = theme_profile_slug or "cartoon_kids"
        ui = self._theme_render_ui_defaults(profile_slug)
        anchor = self._pick_theme_anchor(title)
        digest = hashlib.md5(f"{title}|{profile_slug}".encode("utf-8")).hexdigest()
        tint = int(digest[:2], 16)
        accent_mix = (
            min(255, int(ui["accent"][1:3], 16) + tint // 8),
            min(255, int(ui["accent"][3:5], 16) + tint // 10),
            min(255, int(ui["accent"][5:7], 16) + tint // 12),
        )
        accent = _rgb_to_hex(accent_mix)
        profile_map = {
            "anime": ("Blaze", "anime-hero", "student", "academy jacket", "hero trousers", "energy band", "comic runners", "Anime"),
            "pixar": ("Buddy", "cartoon-pal", "student", "soft hoodie", "rounded trousers", "story satchel", "comic sneakers", "Pixar"),
            "superhero": ("Sentinel", "hero-guardian", "teacher", "hero jacket", "powered leggings", "comic emblem", "hero boots", "Hero"),
            "manga": ("Echo", "manga-hero", "student", "ink-lined coat", "dramatic trousers", "speed scarf", "comic boots", "Manga"),
            "fantasy": ("Quest", "fantasy-ranger", "student", "travel cloak", "quest trousers", "rune compass", "trail boots", "Fantasy"),
            "horror": ("Nocturne", "mystic-hero", "student", "shadow jacket", "night trousers", "glow charm", "silent boots", "Horror"),
            "cartoon_kids": ("Spark", "explorer", "student", "playful tee", "adventure pants", "story badge", "comic sneakers", "Story"),
        }
        suffix, variant, category, top, bottom, accessory, shoes, badge = profile_map.get(
            profile_slug,
            profile_map["cartoon_kids"],
        )
        return {
            "id": f"generated-{self._slugify(title) or 'theme'}",
            "label": title,
            "badge_text": f"{badge} Mode",
            "headline": f"{title} now has its own live character preview, palette, and comic setup.",
            "sound_effect": "WOW!",
            "theme_profile": profile_slug,
            "ui": {
                **ui,
                "accent": accent,
            },
            "comic_elements": self._theme_render_comic_elements(profile_slug),
            "featured_cast": [
                {
                    "name": f"{anchor} {suffix}",
                    "description": f"The signature {profile_slug.replace('_', ' ')} guide for {title}.",
                    "variant": variant,
                    "category": category,
                    "palette": {
                        "hair": "#47311f",
                        "skin": "#efc39b",
                        "costume_primary": accent,
                        "costume_accent": ui["accent_2"],
                        "eye": ui["accent_3"],
                    },
                    "costume": {
                        "top": top,
                        "bottom": bottom,
                        "accessory": accessory,
                        "shoes": shoes,
                    },
                    "tags": [self._slugify(title), profile_slug, "theme-generated", "preview"],
                }
            ],
            "featured_background": {
                "name": f"{anchor} Realm",
                "description": f"A {profile_slug.replace('_', ' ')} environment tuned for {title}.",
                "category": "theme_world",
                "prompt_fragment": f"{profile_slug.replace('_', ' ')} themed story world for {title}, comic depth, bright panels",
                "palette": {
                    "primary": ui["bg_start"],
                    "secondary": ui["bg_end"],
                    "accent": accent,
                    "floor": "#8b5e3c",
                },
                "tags": [self._slugify(title), profile_slug, "theme-world"],
            },
        }

    def _resolve_theme_render(
        self,
        *,
        theme_meta: dict[str, Any] | None = None,
        title: str = "",
        source_text: str = "",
        theme_profile_slug: str = "",
    ) -> dict[str, Any]:
        explicit = (theme_meta or {}).get("theme_render")
        if isinstance(explicit, dict) and explicit.get("id"):
            profile_slug = str(
                explicit.get("theme_profile")
                or theme_profile_slug
                or (theme_meta or {}).get("theme_profile")
                or "cartoon_kids"
            )
            base = self._default_theme_render(profile_slug)
            merged = self._merge_dicts(base, explicit)
            merged["theme_profile"] = profile_slug
            return merged

        exact_entry = self._exact_theme_render_entry(theme_meta=theme_meta, title=title)
        if exact_entry:
            profile_slug = str(
                exact_entry.get("theme_profile")
                or theme_profile_slug
                or (theme_meta or {}).get("theme_profile")
                or "cartoon_kids"
            )
            base = self._default_theme_render(profile_slug)
            merged = self._merge_dicts(base, exact_entry)
            merged["theme_profile"] = profile_slug
            return merged

        title_blob = " ".join(
            part
            for part in (
                title,
                (theme_meta or {}).get("title", ""),
                (theme_meta or {}).get("slug", ""),
            )
            if part
        ).lower()
        content_blob = source_text[:1600].lower() if not title_blob.strip() else ""
        best_entry: dict[str, Any] | None = None
        best_score = 0
        for entry in self._theme_render_entries:
            if not isinstance(entry, dict):
                continue
            score = self._theme_render_match_score(entry, title_blob) * 4
            if score == 0 and content_blob:
                score = self._theme_render_match_score(entry, content_blob)
            if score > best_score:
                best_score = score
                best_entry = entry

        resolved_profile = str(
            (best_entry or {}).get("theme_profile")
            or theme_profile_slug
            or (theme_meta or {}).get("theme_profile")
            or "cartoon_kids"
        )
        base = self._default_theme_render(resolved_profile)
        if best_entry:
            merged = self._merge_dicts(base, best_entry)
            merged["theme_profile"] = resolved_profile
            return merged
        if title or (theme_meta or {}).get("title"):
            synthesized = self._synthesise_theme_render(title or str((theme_meta or {}).get("title") or ""), resolved_profile)
            merged = self._merge_dicts(base, synthesized)
            merged["theme_profile"] = resolved_profile
            return merged
        return base

    def _build_theme_cast_records(self, theme_render: dict[str, Any], theme_profile_slug: str) -> list[dict[str, Any]]:
        featured = theme_render.get("featured_cast") or []
        if not isinstance(featured, list):
            return []
        ui = theme_render.get("ui") or {}
        records: list[dict[str, Any]] = []
        for index, entry in enumerate(featured):
            if not isinstance(entry, dict):
                continue
            name = str(entry.get("name") or "").strip()
            if not name:
                continue
            palette = entry.get("palette") or {}
            record = {
                "id": f"THEME_{self._slugify(str(theme_render.get('id') or name))}_{index}",
                "name": name,
                "category": str(entry.get("category") or "student"),
                "sub_type": str(entry.get("variant") or theme_profile_slug or "theme_character"),
                "source": "theme_manifest",
                "theme_fit": [theme_profile_slug] if theme_profile_slug else [],
                "subject_fit": ["all"],
                "image_prompt": str(entry.get("description") or f"{name} themed comic character"),
                "tags": list(dict.fromkeys([
                    *(entry.get("tags") or []),
                    str(theme_render.get("id") or ""),
                    str(theme_render.get("label") or ""),
                    theme_profile_slug,
                    "theme-manifest",
                ])),
                "palette": {
                    "hair": self._safe_hex(palette.get("hair"), "#47311f"),
                    "skin": self._safe_hex(palette.get("skin"), "#efc39b"),
                    "costume_primary": self._safe_hex(palette.get("costume_primary"), self._safe_hex(ui.get("accent"), "#f4631e")),
                    "costume_accent": self._safe_hex(palette.get("costume_accent"), self._safe_hex(ui.get("accent_2"), "#ffd93d")),
                    "eye": self._safe_hex(palette.get("eye"), self._safe_hex(ui.get("accent_3"), "#1e6fd9")),
                },
                "costume": entry.get("costume") or {},
                "render_variant": str(entry.get("variant") or ""),
            }
            records.append(record)
        return records

    def _build_theme_background_records(self, theme_render: dict[str, Any], theme_profile_slug: str) -> list[dict[str, Any]]:
        background = theme_render.get("featured_background")
        if not isinstance(background, dict):
            return []
        ui = theme_render.get("ui") or {}
        palette = background.get("palette") or {}
        return [
            {
                "id": f"THEME_BG_{self._slugify(str(theme_render.get('id') or background.get('name') or 'theme'))}",
                "name": str(background.get("name") or "Theme World"),
                "category": str(background.get("category") or "theme_world"),
                "source": "theme_manifest",
                "theme_fit": [theme_profile_slug] if theme_profile_slug else [],
                "subject_fit": ["all"],
                "prompt_fragment": str(background.get("prompt_fragment") or background.get("description") or "themed comic background"),
                "palette": {
                    "primary": self._safe_hex(palette.get("primary"), self._safe_hex(ui.get("bg_start"), "#fff8dd")),
                    "secondary": self._safe_hex(palette.get("secondary"), self._safe_hex(ui.get("bg_end"), "#edfff2")),
                    "accent": self._safe_hex(palette.get("accent"), self._safe_hex(ui.get("accent"), "#f4631e")),
                    "floor": self._safe_hex(palette.get("floor"), "#8b5e3c"),
                },
                "tags": list(dict.fromkeys([
                    *(background.get("tags") or []),
                    str(theme_render.get("id") or ""),
                    theme_profile_slug,
                    "theme-manifest",
                ])),
            }
        ]

    @staticmethod
    def _prioritise_theme_records(
        featured: list[dict[str, Any]],
        existing: list[dict[str, Any]],
        limit: int,
        *,
        unique_key: str = "id",
    ) -> list[dict[str, Any]]:
        merged: list[dict[str, Any]] = []
        seen: set[str] = set()
        for source in (featured, existing):
            for record in source:
                record_key = str(record.get(unique_key) or record.get("name") or "")
                if not record_key or record_key in seen:
                    continue
                merged.append(record)
                seen.add(record_key)
                if len(merged) >= limit:
                    return merged
        return merged

    def _default_theme_profile(self, slug: str = "cartoon_kids") -> dict[str, Any]:
        fallback = self._theme_profiles_by_slug.get(slug) or self._theme_profiles_by_slug.get("cartoon_kids")
        if fallback:
            return dict(fallback)
        return {
            "slug": slug,
            "label": slug.replace("_", " ").title(),
            "dataset_style": "cute_cartoon",
            "palette_style": "cartoon",
            "layout_style": "balanced_grid",
            "preferred_bubble": "normal",
            "title_font_family": "playful",
            "body_font_family": "rounded",
            "effect_font_family": "playful",
            "ui_display_font": "\"Bangers\", cursive",
            "ui_body_font": "\"Comic Neue\", cursive",
            "font_label": "Comic Neue",
            "character_summary": "friendly comic learners",
            "background_summary": "study-friendly bright spaces",
            "layout_label": "Balanced story grid",
            "accent": "#f4631e",
            "surface": "#ffffff",
        }

    def _resolve_theme_profile(
        self,
        *,
        theme_meta: dict[str, Any] | None,
        requested_style: str = "",
        title: str = "",
        source_text: str = "",
    ) -> dict[str, Any]:
        explicit_slug = (theme_meta or {}).get("theme_profile")
        if explicit_slug and explicit_slug in self._theme_profiles_by_slug:
            return dict(self._theme_profiles_by_slug[explicit_slug])

        style_slug = (requested_style or "").strip().lower()
        style_profile = self._theme_style_rules.get(style_slug) or THEME_PROFILE_FALLBACKS.get(style_slug)

        title_blob = " ".join(
            part for part in (
                title,
                (theme_meta or {}).get("title", ""),
                (theme_meta or {}).get("slug", ""),
            )
            if part
        ).lower()
        sample_blob = f"{title_blob} {(source_text or '')[:6000].lower()}"

        best_profile = ""
        best_score = -1
        for rule in self._theme_keyword_rules:
            if not isinstance(rule, dict):
                continue
            profile_slug = str(rule.get("profile") or "").strip().lower()
            if not profile_slug:
                continue
            keywords = [str(keyword).lower() for keyword in (rule.get("keywords") or []) if str(keyword).strip()]
            score = 0
            for keyword in keywords:
                if keyword in title_blob:
                    score += 8
                elif keyword in sample_blob:
                    score += 3
            if score > best_score:
                best_score = score
                best_profile = profile_slug

        resolved_slug = best_profile if best_score >= 8 else (style_profile or best_profile or "")
        if not resolved_slug and style_slug in self._theme_profiles_by_slug:
            resolved_slug = style_slug
        if not resolved_slug:
            palette_slug, dataset_style = self._resolve_style_pair(style_slug or "cute_cartoon")
            resolved_slug = self._theme_style_rules.get(dataset_style) or THEME_PROFILE_FALLBACKS.get(dataset_style) or palette_slug

        return self._default_theme_profile(resolved_slug or "cartoon_kids")

    def _load_archive_characters(self) -> list[dict[str, Any]]:
        archive_files = [
            ("dc", self.archive_dir / "dc-wikia-data.csv"),
            ("marvel", self.archive_dir / "marvel-wikia-data.csv"),
        ]
        records: list[dict[str, Any]] = []
        for publisher, path in archive_files:
            if not path.exists():
                continue
            with path.open(encoding="utf-8", errors="ignore", newline="") as handle:
                reader = csv.DictReader(handle)
                for row in reader:
                    record = self._normalise_archive_character(row, publisher)
                    if record is not None:
                        records.append(record)
        records.sort(key=lambda item: item.get("appearances", 0), reverse=True)
        return records

    def _normalise_archive_character(self, row: dict[str, str], publisher: str) -> dict[str, Any] | None:
        name = (row.get("name") or "").strip()
        if not name:
            return None
        align = (row.get("ALIGN") or "").strip()
        if "Bad" in align:
            return None
        sex = (row.get("SEX") or "").strip()
        if sex and "Male" not in sex and "Female" not in sex:
            return None

        try:
            appearances = int(float((row.get("APPEARANCES") or "0").strip() or "0"))
        except ValueError:
            appearances = 0
        if appearances < 25:
            return None

        display_name = re.sub(r"\s*\([^)]*\)\s*$", "", name).strip() or name
        hair_name = (row.get("HAIR") or "").strip().lower()
        eye_name = (row.get("EYE") or "").strip().lower()
        primary, accent = ARCHIVE_PUBLISHER_COLORS.get(publisher, ("#3b82f6", "#60a5fa"))
        palette = self._archive_palette(display_name, publisher, hair_name, eye_name, sex)
        secret_identity = (row.get("ID") or "").strip()
        prompt = (
            f"kid-friendly anime superhero, {display_name}, {publisher} comic inspired hero, "
            f"{hair_name or 'stylized hair'}, {eye_name or 'bright eyes'}, vibrant heroic costume, "
            "clean line art, friendly expression, safe colorful comic-book design"
        )
        return {
            "id": f"{publisher.upper()}_{row.get('page_id') or hashlib.md5(display_name.encode('utf-8')).hexdigest()[:8]}",
            "name": display_name,
            "category": "archive_hero",
            "sub_type": "superhero",
            "source": "archive",
            "publisher": publisher,
            "align": align,
            "identity": secret_identity,
            "sex": sex,
            "hair_name": hair_name,
            "eye_name": eye_name,
            "appearances": appearances,
            "palette": palette,
            "image_prompt": prompt,
            "tags": [publisher, align.lower(), sex.lower(), secret_identity.lower(), "hero", "archive"],
            "costume": {
                "top": "bright superhero suit",
                "bottom": "comic hero suit",
                "accessory": "hero badge and cape",
                "shoes": "hero boots",
            },
            "publisher_colors": {"primary": primary, "accent": accent},
        }

    def _archive_palette(self, display_name: str, publisher: str, hair_name: str, eye_name: str, sex: str) -> dict[str, str]:
        primary, accent = ARCHIVE_PUBLISHER_COLORS.get(publisher, ("#3b82f6", "#60a5fa"))
        name_hash = hashlib.md5(display_name.encode("utf-8")).hexdigest()
        alt_seed = int(name_hash[:2], 16)
        alt_mix = (
            min(255, int(primary[1:3], 16) + alt_seed // 4),
            min(255, int(primary[3:5], 16) + alt_seed // 6),
            min(255, int(primary[5:7], 16) + alt_seed // 8),
        )
        return {
            "hair": HAIR_COLOR_MAP.get(hair_name, "#2f2a26"),
            "skin": "#f2c7a5" if "Female" in sex else "#edc39b",
            "costume_primary": primary,
            "costume_accent": _rgb_to_hex(alt_mix),
            "eye": EYE_COLOR_MAP.get(eye_name, accent),
        }

    def _resolve_style_pair(self, requested_style: str) -> tuple[str, str]:
        style_slug = (requested_style or "").strip().lower()
        if style_slug in STYLE_RESOLUTION_MAP:
            return STYLE_RESOLUTION_MAP[style_slug]
        if style_slug in STYLE_LIBRARY:
            return style_slug, STYLE_RESOLUTION_MAP.get(style_slug, (style_slug, "cute_cartoon"))[1]
        if style_slug in self._dataset_styles_by_name:
            profile_slug = self._theme_style_rules.get(style_slug) or THEME_PROFILE_FALLBACKS.get(style_slug, "cartoon_kids")
            profile = self._theme_profiles_by_slug.get(profile_slug, {})
            palette_style = profile.get("palette_style") or STYLE_RESOLUTION_MAP.get(profile_slug, ("cartoon", style_slug))[0]
            return str(palette_style), style_slug
        return STYLE_RESOLUTION_MAP["cute_cartoon"]

    def _style_display_name(self, requested_style: str) -> str:
        palette_style, dataset_style = self._resolve_style_pair(requested_style)
        style_record = self._dataset_styles_by_name.get(dataset_style, {})
        return style_record.get("label") or DATASET_STYLE_LABELS.get(dataset_style) or STYLE_LIBRARY[palette_style]["name"]

    @staticmethod
    def _safe_hex(value: Any, fallback: str) -> str:
        if isinstance(value, str):
            long_match = re.search(r"#[0-9A-Fa-f]{6}", value)
            if long_match:
                return long_match.group(0)
            short_match = re.search(r"#[0-9A-Fa-f]{3}", value)
            if short_match:
                short = short_match.group(0).lstrip("#")
                return f"#{short[0] * 2}{short[1] * 2}{short[2] * 2}"
        return fallback

    @staticmethod
    def _blend_hex(color: str, target: str = "#ffffff", amount: float = 0.5) -> str:
        amount = max(0.0, min(1.0, amount))
        return _rgb_to_hex(_lerp_color(_hex_to_rgb(color), _hex_to_rgb(target), amount))

    def _record_blob(self, record: dict[str, Any]) -> str:
        costume = record.get("costume") or {}
        parts = [
            record.get("name", ""),
            record.get("category", ""),
            record.get("sub_type", ""),
            record.get("source", ""),
            record.get("publisher", ""),
            record.get("align", ""),
            record.get("identity", ""),
            record.get("hair_name", ""),
            record.get("eye_name", ""),
            record.get("image_prompt", ""),
            " ".join(record.get("tags") or []),
            " ".join(str(value) for value in costume.values() if value),
        ]
        return " ".join(parts).lower()

    @staticmethod
    def _seeded_rng(*parts: Any) -> random.Random:
        seed = "|".join(str(part) for part in parts if part is not None)
        digest = hashlib.md5(seed.encode("utf-8")).hexdigest()
        return random.Random(int(digest[:12], 16))

    def _theme_fit_score(self, record: dict[str, Any], theme_profile_slug: str) -> int:
        fits = {str(item).lower() for item in (record.get("theme_fit") or [])}
        if not theme_profile_slug:
            return 0
        if theme_profile_slug in fits:
            return 18
        if "all" in fits:
            return 6
        return -8 if fits else 0

    def _subject_fit_score(self, record: dict[str, Any], subject: str) -> int:
        fits = {str(item).lower() for item in (record.get("subject_fit") or [])}
        if subject in fits:
            return 8
        if "all" in fits:
            return 4
        return 0

    def _pick_ranked_candidates(
        self,
        ranked: list[dict[str, Any]],
        target: int,
        rng: random.Random,
        *,
        window: int = 4,
        unique_key: str = "id",
    ) -> list[dict[str, Any]]:
        selected: list[dict[str, Any]] = []
        pool = list(ranked)
        used: set[str] = set()
        while pool and len(selected) < target:
            current_window = [item for item in pool[: max(window, 1)] if (item.get(unique_key) or item.get("name")) not in used]
            if not current_window:
                break
            choice = rng.choice(current_window)
            choice_key = choice.get(unique_key) or choice.get("name")
            if choice_key:
                used.add(str(choice_key))
            selected.append(choice)
            pool = [item for item in pool if (item.get(unique_key) or item.get("name")) != choice_key]
        return selected

    def _theme_filtered_records(self, records: list[dict[str, Any]], theme_profile_slug: str) -> list[dict[str, Any]]:
        filtered = []
        for record in records:
            fits = {str(item).lower() for item in (record.get("theme_fit") or [])}
            if not fits or theme_profile_slug in fits or "all" in fits:
                filtered.append(record)
        return filtered or records

    def _character_score(self, record: dict[str, Any], subject: str, desired_category: str, theme_profile_slug: str) -> int:
        blob = self._record_blob(record)
        score = 0
        category = (record.get("category") or "").lower()
        score += self._theme_fit_score(record, theme_profile_slug)
        score += self._subject_fit_score(record, subject)
        if category == desired_category:
            score += 12
        elif category != "mascot":
            score += 1
        if category == "mascot" and theme_profile_slug not in {"cartoon_kids", "pixar"}:
            score -= 6
        for hint in SUBJECT_CHARACTER_HINTS.get(subject, set()):
            if hint in blob:
                score += 3
        if subject in blob:
            score += 5
        if theme_profile_slug == "superhero" and any(token in blob for token in ("hero", "cape", "captain", "shield", "action")):
            score += 6
        if theme_profile_slug == "manga" and any(token in blob for token in ("ink", "dramatic", "monochrome", "manga")):
            score += 6
        if theme_profile_slug == "fantasy" and any(token in blob for token in ("wizard", "knight", "magic", "dragon", "enchanted")):
            score += 6
        if "cute" in blob or "colorful" in blob or "friendly" in blob:
            score += 1
        return score

    def _archive_character_score(self, record: dict[str, Any], subject: str, theme_profile_slug: str) -> int:
        if theme_profile_slug != "superhero":
            return -100
        blob = self._record_blob(record)
        score = min(int(record.get("appearances", 0) // 150), 24)
        if "good characters" in blob:
            score += 10
        elif "neutral characters" in blob:
            score += 4
        if "female characters" in blob:
            score += 2
        for hint in SUBJECT_CHARACTER_HINTS.get(subject, set()):
            if hint in blob:
                score += 2
        if any(token in blob for token in ("captain", "doctor", "professor", "atom", "iron", "flash", "lantern", "super")):
            score += 2
        return score

    def _background_score(self, record: dict[str, Any], subject: str, theme_profile_slug: str) -> int:
        blob = f"{record.get('name', '')} {record.get('prompt_fragment', '')} {' '.join(record.get('tags') or [])}".lower()
        score = self._theme_fit_score(record, theme_profile_slug) + self._subject_fit_score(record, subject)
        priority = SUBJECT_BACKGROUND_PRIORITY.get(subject, [])
        if record.get("id") in priority:
            score += max(1, len(priority) - priority.index(record["id"]))
        if "battle" in blob:
            score -= 3
        if theme_profile_slug == "manga" and any(token in blob for token in ("ink", "monochrome", "speed")):
            score += 6
        if theme_profile_slug == "fantasy" and any(token in blob for token in ("magic", "castle", "forest", "enchanted")):
            score += 6
        if theme_profile_slug == "superhero" and any(token in blob for token in ("rooftop", "command", "arena", "skyline")):
            score += 6
        if any(token in blob for token in ("bright", "colorful", "sunny", "magical", "cheerful", "warm")):
            score += 2
        return score

    def _assemble_dataset_cast(self, subject: str, panel_count: int, theme_profile_slug: str, rng: random.Random) -> list[dict[str, Any]]:
        if not self._dataset_characters:
            return []
        theme_characters = self._theme_filtered_records(self._dataset_characters, theme_profile_slug)

        selected: list[dict[str, Any]] = []
        used_ids: set[str] = set()
        for desired_category in ("teacher", "student", "subject_expert"):
            ranked = [
                record
                for record in sorted(
                    (record for record in theme_characters if record.get("id") not in used_ids),
                    key=lambda item: self._character_score(item, subject, desired_category, theme_profile_slug),
                    reverse=True,
                )
                if self._character_score(record, subject, desired_category, theme_profile_slug) > 0
            ]
            if ranked:
                chosen = self._pick_ranked_candidates(ranked, 1, rng, window=min(4, len(ranked)))
                if chosen:
                    selected.append(chosen[0])
                    if chosen[0].get("id"):
                        used_ids.add(chosen[0]["id"])

        fill_target = min(4, max(2, panel_count))
        fillers = [
            record
            for record in sorted(
                (record for record in theme_characters if record.get("id") not in used_ids),
                key=lambda item: self._character_score(item, subject, (item.get("category") or "").lower(), theme_profile_slug),
                reverse=True,
            )
            if self._character_score(record, subject, (record.get("category") or "").lower(), theme_profile_slug) > 0
            and (record.get("category") or "").lower() != "mascot"
        ]
        for record in self._pick_ranked_candidates(fillers, max(0, fill_target - len(selected)), rng, window=min(6, len(fillers) or 1)):
            selected.append(record)
            if record.get("id"):
                used_ids.add(record["id"])

        if len(selected) < fill_target and theme_profile_slug in {"cartoon_kids", "pixar", "anime", "fantasy"}:
            mascot_ranked = [
                record
                for record in sorted(
                    (record for record in theme_characters if record.get("id") not in used_ids and (record.get("category") or "").lower() == "mascot"),
                    key=lambda item: self._character_score(item, subject, "mascot", theme_profile_slug),
                    reverse=True,
                )
                if self._character_score(record, subject, "mascot", theme_profile_slug) > 0
            ]
            for record in self._pick_ranked_candidates(mascot_ranked, max(0, fill_target - len(selected)), rng, window=min(3, len(mascot_ranked) or 1)):
                selected.append(record)
                if record.get("id"):
                    used_ids.add(record["id"])
        return selected or theme_characters[:3]

    def _assemble_archive_cast(self, subject: str, panel_count: int, theme_profile_slug: str, rng: random.Random) -> list[dict[str, Any]]:
        if not self._archive_characters or theme_profile_slug != "superhero":
            return []
        ranked = [
            item
            for item in sorted(self._archive_characters, key=lambda record: self._archive_character_score(record, subject, theme_profile_slug), reverse=True)
            if self._archive_character_score(item, subject, theme_profile_slug) > 0
        ]
        selected: list[dict[str, Any]] = []
        used_names: set[str] = set()
        used_publishers: set[str] = set()
        for record in ranked:
            name = (record.get("name") or "").lower()
            publisher = (record.get("publisher") or "").lower()
            if not name or name in used_names:
                continue
            if publisher in used_publishers and len(selected) < 2:
                continue
            selected.append(record)
            used_names.add(name)
            if publisher:
                used_publishers.add(publisher)
            if len(selected) >= min(4, max(2, panel_count)):
                break
        if selected:
            return selected
        return self._pick_ranked_candidates(ranked, min(3, len(ranked)), rng, window=min(5, len(ranked) or 1), unique_key="name")

    def _merge_casts(
        self,
        dataset_cast: list[dict[str, Any]],
        archive_cast: list[dict[str, Any]],
        panel_count: int,
        cast_mode: str,
        dataset_style: str,
        theme_profile_slug: str,
    ) -> list[dict[str, Any]]:
        target = min(4, max(2, panel_count))
        if cast_mode == "cominote":
            return dataset_cast[:target] or archive_cast[:target]
        if cast_mode == "archive":
            return archive_cast[:target] or dataset_cast[:target]

        merged: list[dict[str, Any]] = []
        used_ids: set[str] = set()
        if cast_mode == "mixed":
            source_order = (archive_cast, dataset_cast)
        elif theme_profile_slug == "superhero" and archive_cast:
            source_order = (archive_cast, dataset_cast)
        else:
            source_order = (dataset_cast, archive_cast if theme_profile_slug == "superhero" else [])

        pointers = [0 for _ in source_order]
        while len(merged) < target:
            made_progress = False
            for idx, source in enumerate(source_order):
                while pointers[idx] < len(source):
                    record = source[pointers[idx]]
                    pointers[idx] += 1
                    record_id = record.get("id") or record.get("name")
                    if record_id in used_ids:
                        continue
                    merged.append(record)
                    used_ids.add(record_id)
                    made_progress = True
                    break
                if len(merged) >= target:
                    return merged
            if not made_progress:
                break
        return merged[:target]

    def _assemble_dataset_backgrounds(self, subject: str, panel_count: int, theme_profile_slug: str, rng: random.Random) -> list[dict[str, Any]]:
        if not self._dataset_backgrounds:
            return []
        theme_backgrounds = self._theme_filtered_records(self._dataset_backgrounds, theme_profile_slug)
        ranked = [
            item
            for item in sorted(theme_backgrounds, key=lambda record: self._background_score(record, subject, theme_profile_slug), reverse=True)
            if self._background_score(item, subject, theme_profile_slug) > 0
        ]
        selected: list[dict[str, Any]] = []
        used_ids: set[str] = set()
        used_categories: set[str] = set()
        for record in ranked:
            score = self._background_score(record, subject, theme_profile_slug)
            if score <= 0:
                continue
            record_id = record.get("id", "")
            category = (record.get("category") or "").lower()
            if record_id in used_ids:
                continue
            if category in used_categories and len(selected) >= 3:
                continue
            selected.append(record)
            if record_id:
                used_ids.add(record_id)
            if category:
                used_categories.add(category)
            if len(selected) >= min(max(panel_count, 3), 6):
                break
        if selected:
            return selected
        fallback = self._pick_ranked_candidates(ranked, max(1, min(panel_count, 3)), rng, window=min(5, len(ranked) or 1))
        return fallback or theme_backgrounds[: max(1, min(panel_count, 3))]

    def _emotion_sequence(self, panel_count: int) -> list[str]:
        sequences = {
            1: ["happy"],
            2: ["happy", "proud"],
            3: ["happy", "excited", "proud"],
            4: ["happy", "thinking", "excited", "proud"],
            5: ["happy", "thinking", "excited", "surprised", "proud"],
            6: ["happy", "thinking", "excited", "surprised", "proud", "happy"],
        }
        if panel_count in sequences:
            return sequences[panel_count]
        pattern = ["happy", "thinking", "excited", "surprised", "proud", "happy"]
        return [pattern[index % len(pattern)] for index in range(panel_count)]

    def _pick_dataset_emotion(self, emotion_name: str) -> dict[str, Any]:
        for record in self._dataset_emotions:
            if (record.get("name") or "").lower() == emotion_name:
                return record
        return {"name": emotion_name, "sound_effect": "WOW!", "prompt_fragment": emotion_name}

    def _pick_dataset_pose(self, scene: Scene, emotion: dict[str, Any], theme_profile_slug: str) -> dict[str, Any]:
        preferred: list[str]
        scene_title = scene.title.lower()
        emotion_name = (emotion.get("name") or "happy").lower()
        if theme_profile_slug == "superhero" and emotion_name in {"excited", "proud"}:
            preferred = ["heroic_landing", "power_stance", "pointing"]
        elif theme_profile_slug == "fantasy" and emotion_name in {"excited", "thinking", "proud"}:
            preferred = ["spell_cast", "reading", "power_stance"]
        elif theme_profile_slug == "manga" and emotion_name in {"thinking", "surprised", "shocked"}:
            preferred = ["thinking", "writing", "shocked_reaction"]
        elif emotion_name == "thinking" or "connection" in scene_title:
            preferred = ["thinking", "writing", "reading"]
        elif emotion_name in {"surprised", "shocked"}:
            preferred = ["shocked_reaction", "jumping", "pointing"]
        elif emotion_name == "proud" or "wrap" in scene_title:
            preferred = ["celebrating", "standing"]
        elif emotion_name == "excited":
            preferred = ["explaining", "pointing", "celebrating"]
        else:
            preferred = ["standing", "explaining", "pointing"]

        for pose_name in preferred:
            for record in self._dataset_poses:
                theme_fit = {str(item).lower() for item in (record.get("theme_fit") or [])}
                if (record.get("name") or "").lower() == pose_name and (not theme_fit or theme_profile_slug in theme_fit or "all" in theme_fit):
                    return record
        return {"name": "standing", "prompt_fragment": "standing upright", "good_for_emotions": [emotion_name]}

    def _pick_dataset_bubble(self, scene: Scene, emotion: dict[str, Any], theme_profile: dict[str, Any]) -> dict[str, Any]:
        theme_profile_slug = str(theme_profile.get("slug") or "")
        emotion_name = (emotion.get("name") or "happy").lower()
        preferred_type = str(theme_profile.get("preferred_bubble") or "").lower()
        if theme_profile_slug == "manga":
            return self._dataset_bubbles_by_type.get("manga_box", {"type": "manga_box", "fill": "#ffffff"})
        if theme_profile_slug == "pixar":
            return self._dataset_bubbles_by_type.get("pixar_round", {"type": "pixar_round", "fill": "#fff8f1"})
        if theme_profile_slug == "superhero" and emotion_name in {"excited", "surprised", "shocked"} and len(scene.dialogue) <= 72:
            return self._dataset_bubbles_by_type.get("action_burst", {"type": "action_burst", "fill": "#fff9c4"})
        if emotion_name == "thinking":
            return self._dataset_bubbles_by_type.get("thought_bubble", {"type": "thought_bubble", "fill": "#e8eaf6"})
        if emotion_name in {"excited", "surprised", "shocked"} and len(scene.dialogue) <= 42:
            return self._dataset_bubbles_by_type.get("shouting", {"type": "shouting", "fill": "#fff9c4"})
        if preferred_type and not (preferred_type == "action_burst" and len(scene.dialogue) > 72):
            preferred = self._dataset_bubbles_by_type.get(preferred_type)
            if preferred:
                return preferred
        return self._dataset_bubbles_by_type.get("normal", {"type": "normal", "fill": "#ffffff"})

    def _pick_dataset_sfx(self, subject: str, emotion: dict[str, Any], index: int, theme_profile_slug: str) -> dict[str, Any]:
        emotion_name = (emotion.get("name") or "happy").lower()
        preferred = list(EMOTION_TO_SFX.get(emotion_name, []))
        if theme_profile_slug == "manga":
            preferred = ["BOOM!", "AHA!", *preferred]
        elif theme_profile_slug == "superhero":
            preferred = ["POW!", "BOOM!", "ZAP!", *preferred]
        elif theme_profile_slug == "pixar":
            preferred = ["WOW!", "YES!", *preferred]
        elif subject == "science":
            preferred = ["ZAP!", "AHA!", *preferred]
        elif subject == "mathematics":
            preferred = ["AHA!", "YES!", *preferred]
        elif subject == "literature":
            preferred = ["WOW!", "POW!", *preferred]

        for word in preferred:
            if word in self._dataset_sfx_by_word:
                return self._dataset_sfx_by_word[word]
        if self._dataset_sfx:
            return self._dataset_sfx[index % len(self._dataset_sfx)]
        return {"word": EFFECT_WORDS[index % len(EFFECT_WORDS)], "color_primary": "#e53935", "color_secondary": "#ffd54f"}

    def _build_visual_plan(
        self,
        title: str,
        subject: str,
        style: str,
        scenes: list[Scene],
        cast_mode: str,
        *,
        theme_meta: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        theme_profile = self._resolve_theme_profile(theme_meta=theme_meta, requested_style=style, title=title)
        fallback_palette_style, fallback_dataset_style = self._resolve_style_pair(style)
        palette_style = str(theme_profile.get("palette_style") or fallback_palette_style)
        dataset_style = str(theme_profile.get("dataset_style") or fallback_dataset_style)
        base_style_record = self._dataset_styles_by_name.get(dataset_style, {"name": dataset_style, "label": dataset_style.replace("_", " ").title()})
        style_record = {**theme_profile, **base_style_record}
        theme_profile_slug = str(theme_profile.get("slug") or "")
        theme_render = self._resolve_theme_render(
            theme_meta=theme_meta,
            title=title,
            theme_profile_slug=theme_profile_slug,
        )
        theme_cast = self._build_theme_cast_records(theme_render, theme_profile_slug)
        theme_backgrounds = self._build_theme_background_records(theme_render, theme_profile_slug)
        rng = self._seeded_rng(title, subject, dataset_style, theme_profile_slug, len(scenes))
        dataset_cast = self._assemble_dataset_cast(subject, len(scenes), theme_profile_slug, rng)
        archive_cast = self._assemble_archive_cast(subject, len(scenes), theme_profile_slug, rng)
        cast = self._merge_casts(dataset_cast, archive_cast, len(scenes), cast_mode, dataset_style, theme_profile_slug)
        cast = self._prioritise_theme_records(theme_cast, cast, min(4, max(2, len(scenes))))
        backgrounds = self._assemble_dataset_backgrounds(subject, len(scenes), theme_profile_slug, rng)
        backgrounds = self._prioritise_theme_records(
            theme_backgrounds,
            backgrounds,
            max(1, min(max(len(scenes), 3), 6)),
        )
        emotions = self._emotion_sequence(len(scenes))

        panels = []
        for index, scene in enumerate(scenes):
            character = cast[index % len(cast)] if cast else {"name": scene.speaker, "category": "student", "palette": {}}
            background = (
                backgrounds[index % len(backgrounds)]
                if backgrounds
                else {"name": scene.background, "category": "school", "palette": {}}
            )
            emotion = self._pick_dataset_emotion(emotions[index])
            pose = self._pick_dataset_pose(scene, emotion, theme_profile_slug)
            bubble = self._pick_dataset_bubble(scene, emotion, theme_profile)
            sfx = self._pick_dataset_sfx(subject, emotion, index, theme_profile_slug)
            theme_elements = theme_render.get("comic_elements") or self._theme_render_comic_elements(theme_profile_slug)
            panels.append(
                {
                    "character": character,
                    "background": background,
                    "emotion": emotion,
                    "pose": pose,
                    "bubble": bubble,
                    "sfx": sfx,
                    "theme_profile": theme_profile,
                    "theme_render_id": theme_render.get("id", ""),
                    "theme_comic_elements": theme_elements,
                }
            )

        return {
            "palette_style": palette_style,
            "dataset_style": dataset_style,
            "style": style_record,
            "panels": panels,
            "cast_mode": cast_mode,
            "theme_profile": theme_profile,
            "theme_render": theme_render,
            "cast_sources": sorted({(panel.get("character") or {}).get("source", "dataset") for panel in panels}),
        }

    def _paginate_story_indices(self, scene_count: int) -> list[list[int]]:
        if scene_count <= 0:
            return [[]]
        if scene_count <= 2:
            return [list(range(scene_count))]

        pages: list[list[int]] = []
        index = 0
        page_index = 0
        while index < scene_count:
            remaining = scene_count - index
            if remaining <= 4:
                target = remaining
            else:
                target = PANEL_PAGE_PATTERN[page_index % len(PANEL_PAGE_PATTERN)]
                target = min(target, remaining)
                if remaining - target == 1:
                    target = target + 1 if target < 4 else target - 1
            target = max(2, min(4, target, remaining))
            pages.append(list(range(index, index + target)))
            index += target
            page_index += 1
        return pages

    def _layout_variant(self, panel_count: int, layout_style: str, page_number: int) -> str:
        if panel_count <= 1:
            return "solo_splash"
        if panel_count == 2:
            return "wide_pair" if page_number % 2 else "stacked_pair"
        if panel_count == 3:
            variants = (
                "hero_top_two_bottom",
                "three_columns",
                "tall_left_stack_right",
                "two_top_wide_bottom",
                "manga_mixed_three",
            )
            if layout_style == "manga_cascade":
                variants = ("tall_left_stack_right", "manga_mixed_three", "three_columns", "hero_top_two_bottom")
            return variants[(page_number - 1) % len(variants)]
        variants = (
            "four_grid",
            "wide_top_three_bottom",
            "tall_left_two_right_bottom",
            "zig_zag_four",
            "manga_mixed_four",
        )
        if layout_style == "action_splash":
            variants = ("wide_top_three_bottom", "four_grid", "zig_zag_four", "tall_left_two_right_bottom")
        elif layout_style == "cinematic_grid":
            variants = ("tall_left_two_right_bottom", "wide_top_three_bottom", "four_grid", "manga_mixed_four")
        elif layout_style == "manga_cascade":
            variants = ("manga_mixed_four", "tall_left_two_right_bottom", "zig_zag_four", "four_grid")
        return variants[(page_number - 1) % len(variants)]

    def _layout_boxes(
        self,
        panel_count: int,
        margin: int,
        gutter: int,
        layout_style: str = "balanced_grid",
        page_number: int = 1,
    ) -> list[tuple[int, int, int, int, int]]:
        content_w = 1420 if layout_style in {"action_splash", "cinematic_grid"} else 1400
        if layout_style == "manga_cascade":
            content_w = 1380
        half_w = (content_w - gutter) // 2
        third_w = (content_w - gutter * 2) // 3
        variant = self._layout_variant(panel_count, layout_style, page_number)

        if panel_count <= 1:
            return [(margin, 0, content_w, 560, 0)]
        if panel_count == 2:
            if variant == "stacked_pair":
                return [
                    (margin, 0, content_w, 330, 0),
                    (margin, 330 + gutter, content_w, 330, 1),
                ]
            return [
                (margin, 0, half_w, 420, 0),
                (margin + half_w + gutter, 0, half_w, 420, 1),
            ]
        if panel_count == 3:
            if variant == "hero_top_two_bottom":
                top_h = 360
                bottom_h = 330
                return [
                    (margin, 0, content_w, top_h, 0),
                    (margin, top_h + gutter, half_w, bottom_h, 1),
                    (margin + half_w + gutter, top_h + gutter, half_w, bottom_h, 2),
                ]
            if variant == "three_columns":
                return [
                    (margin + index * (third_w + gutter), 0, third_w, 680, index)
                    for index in range(3)
                ]
            if variant == "tall_left_stack_right":
                left_w = int(content_w * 0.48)
                right_w = content_w - left_w - gutter
                stacked_h = 330
                return [
                    (margin, 0, left_w, stacked_h * 2 + gutter, 0),
                    (margin + left_w + gutter, 0, right_w, stacked_h, 1),
                    (margin + left_w + gutter, stacked_h + gutter, right_w, stacked_h, 2),
                ]
            if variant == "manga_mixed_three":
                left_w = int(content_w * 0.56)
                right_w = content_w - left_w - gutter
                return [
                    (margin, 0, left_w, 700, 0),
                    (margin + left_w + gutter, 0, right_w, 335, 1),
                    (margin + left_w + gutter, 335 + gutter, right_w, 341, 2),
                ]
            return [
                (margin, 0, half_w, 340, 0),
                (margin + half_w + gutter, 0, half_w, 340, 1),
                (margin, 340 + gutter, content_w, 370, 2),
            ]
        if panel_count == 4:
            if variant == "wide_top_three_bottom":
                top_h = 350
                bottom_h = max(320, MIN_RENDER_PANEL_HEIGHT)
                return [
                    (margin, 0, content_w, top_h, 0),
                    *[
                        (margin + index * (third_w + gutter), top_h + gutter, third_w, bottom_h, index + 1)
                        for index in range(3)
                    ],
                ]
            if variant == "tall_left_two_right_bottom":
                left_w = int(content_w * 0.46)
                right_w = content_w - left_w - gutter
                top_right_h = 300
                bottom_h = 330
                return [
                    (margin, 0, left_w, top_right_h * 2 + gutter, 0),
                    (margin + left_w + gutter, 0, right_w, top_right_h, 1),
                    (margin + left_w + gutter, top_right_h + gutter, right_w, top_right_h, 2),
                    (margin, top_right_h * 2 + gutter * 2, content_w, bottom_h, 3),
                ]
            if variant == "zig_zag_four":
                large_w = int(content_w * 0.56)
                small_w = content_w - large_w - gutter
                top_h = 330
                bottom_h = 330
                return [
                    (margin, 0, large_w, top_h, 0),
                    (margin + large_w + gutter, 0, small_w, top_h, 1),
                    (margin, top_h + gutter, small_w, bottom_h, 2),
                    (margin + small_w + gutter, top_h + gutter, large_w, bottom_h, 3),
                ]
            if variant == "manga_mixed_four":
                top_h = 340
                bottom_h = max(330, MIN_RENDER_PANEL_HEIGHT)
                return [
                    (margin, 0, content_w, top_h, 0),
                    *[
                        (margin + index * (third_w + gutter), top_h + gutter, third_w, bottom_h, index + 1)
                        for index in range(3)
                    ],
                ]
            return [
                (margin, 0, half_w, 320, 0),
                (margin + half_w + gutter, 0, half_w, 320, 1),
                (margin, 320 + gutter, half_w, 320, 2),
                (margin + half_w + gutter, 320 + gutter, half_w, 320, 3),
            ]

        panel_h = 320
        boxes: list[tuple[int, int, int, int, int]] = []
        row = 0
        for index in range(panel_count):
            col = index % 2
            lx = margin + col * (half_w + gutter)
            ly = row * (panel_h + gutter)
            boxes.append((lx, ly, half_w, panel_h, index))
            if col == 1:
                row += 1
        return boxes

    def _panel_theme(self, page_palette: dict[str, str], visual: dict[str, Any]) -> dict[str, str]:
        char_palette = (visual.get("character") or {}).get("palette") or {}
        background_palette = (visual.get("background") or {}).get("palette") or {}
        profile = visual.get("theme_profile") or {}
        return {
            "accent": self._safe_hex(char_palette.get("costume_accent"), self._safe_hex(profile.get("accent"), page_palette["accent"])),
            "accent_alt": self._safe_hex(background_palette.get("accent"), page_palette["accent_alt"]),
            "surface": self._safe_hex(background_palette.get("primary"), self._safe_hex(profile.get("surface"), page_palette["panel_bg"])),
            "surface_alt": self._safe_hex(background_palette.get("secondary"), page_palette["panel_bot"]),
            "floor": self._safe_hex(background_palette.get("floor"), page_palette["ground"]),
            "bubble": self._safe_hex((visual.get("bubble") or {}).get("fill"), page_palette["bubble"]),
            "caption": self._safe_hex(self._dataset_bubbles_by_type.get("narration_box", {}).get("fill"), page_palette["panel_top"]),
            "caption_text": page_palette["caption_bg"],
        }

    def _emotion_face(self, emotion_name: str) -> tuple[str, str]:
        mapping = {
            "happy": ("happy", "smile"),
            "excited": ("excited", "wide"),
            "thinking": ("think", "hmm"),
            "proud": ("happy", "grin"),
            "surprised": ("shock", "o"),
            "shocked": ("shock", "o"),
            "confused": ("think", "hmm"),
        }
        return mapping.get(emotion_name.lower(), ("happy", "smile"))

    @staticmethod
    def _clamp(value: int, minimum: int, maximum: int) -> int:
        return max(minimum, min(maximum, value))

    def _text_size(self, draw, text: str, font, *, spacing: int = 4) -> tuple[int, int]:
        if not text:
            return 0, 0
        if "\n" in text:
            box = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing)
        else:
            box = draw.textbbox((0, 0), text, font=font)
        return box[2] - box[0], box[3] - box[1]

    def _wrap_text_pixels(self, draw, text: str, font, max_width: int) -> list[str]:
        words = text.split()
        if not words:
            return [""]
        lines: list[str] = []
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}".strip()
            width, _ = self._text_size(draw, candidate, font)
            if width <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
        return lines

    def _truncate_line(self, draw, text: str, font, max_width: int) -> str:
        line = text.strip()
        if not line:
            return ""
        if self._text_size(draw, line, font)[0] <= max_width:
            return line
        shortened = line
        while shortened:
            candidate = shortened.rstrip(" ,.;:") + "..."
            if self._text_size(draw, candidate, font)[0] <= max_width:
                return candidate
            shortened = " ".join(shortened.split()[:-1])
        return "..."

    def _fit_text_block(
        self,
        draw,
        text: str,
        max_width: int,
        max_height: int,
        *,
        start_size: int,
        min_size: int = 10,
        bold: bool = False,
        family: str = "comic",
        spacing: int = 4,
        max_lines: int | None = None,
    ) -> tuple[str, Any]:
        cleaned = self._normalise_whitespace(text)
        if not cleaned:
            font = self._font(min_size, bold=bold, family=family)
            return "", font
        for font_size in range(start_size, min_size - 1, -1):
            font = self._font(font_size, bold=bold, family=family)
            lines = self._wrap_text_pixels(draw, cleaned, font, max_width)
            if max_lines is not None and len(lines) > max_lines:
                lines = lines[:max_lines]
                lines[-1] = self._truncate_line(draw, " ".join(cleaned.split()[sum(len(line.split()) for line in lines[:-1]):]), font, max_width)
            wrapped = "\n".join(lines)
            width, height = self._text_size(draw, wrapped, font, spacing=spacing)
            if width <= max_width and height <= max_height and (max_lines is None or len(lines) <= max_lines):
                return wrapped, font

        font = self._font(min_size, bold=bold, family=family)
        lines = self._wrap_text_pixels(draw, cleaned, font, max_width)
        if max_lines is not None and len(lines) > max_lines:
            visible = lines[:max_lines]
            remaining_words = cleaned.split()[sum(len(line.split()) for line in visible[:-1]):]
            visible[-1] = self._truncate_line(draw, " ".join(remaining_words), font, max_width)
            lines = visible
        while len(lines) > 1 and self._text_size(draw, "\n".join(lines), font, spacing=spacing)[1] > max_height:
            visible = lines[:-1]
            remaining_words = cleaned.split()[sum(len(line.split()) for line in visible[:-1]):]
            visible[-1] = self._truncate_line(draw, " ".join(remaining_words), font, max_width)
            lines = visible
        return "\n".join(lines), font

    def _estimate_bubble_fit(
        self,
        draw,
        text: str,
        bubble_w: int,
        max_bubble_h: int,
        bubble_type: str,
        body_family: str,
        *,
        start_size: int,
        min_size: int,
        max_lines: int,
    ) -> tuple[int, int]:
        bubble_type = "normal" if bubble_type == "action_burst" else bubble_type
        if bubble_type == "shouting":
            pad_x = 70 if bubble_w > 360 else 54 if bubble_w > 220 else 36
            pad_y = 30 if max_bubble_h > 130 else 22
        else:
            pad_x = 50 if bubble_w > 360 else 40 if bubble_w > 230 else 28
            pad_y = 24 if max_bubble_h > 118 else 18
        wrapped, font = self._fit_text_block(
            draw,
            text,
            max(52, bubble_w - pad_x * 2),
            max(28, max_bubble_h - pad_y * 2),
            start_size=start_size,
            min_size=min_size,
            bold=bubble_type == "shouting",
            family=body_family,
            spacing=7,
            max_lines=max_lines,
        )
        _text_w, text_h = self._text_size(draw, wrapped, font, spacing=7)
        min_h = min(112 if bubble_w >= 260 else 96, max_bubble_h)
        desired_h = self._clamp(text_h + pad_y * 2 + 4, min_h, max_bubble_h)
        return desired_h, int(getattr(font, "size", min_size))

    @staticmethod
    def _rect_intersection_area(a: tuple[int, int, int, int], b: tuple[int, int, int, int]) -> int:
        left = max(a[0], b[0])
        top = max(a[1], b[1])
        right = min(a[2], b[2])
        bottom = min(a[3], b[3])
        return max(0, right - left) * max(0, bottom - top)

    @staticmethod
    def _inflate_rect(rect: tuple[int, int, int, int], amount: int) -> tuple[int, int, int, int]:
        return rect[0] - amount, rect[1] - amount, rect[2] + amount, rect[3] + amount

    def _smart_panel_composition(
        self,
        draw,
        dialogue: str,
        bubble_type: str,
        body_family: str,
        inner_x1: int,
        inner_x2: int,
        content_top: int,
        content_bottom: int,
        panel_w: int,
        panel_h: int,
        layout_style: str,
        index: int,
    ) -> dict[str, Any]:
        content_w = max(80, inner_x2 - inner_x1)
        content_h = max(80, content_bottom - content_top)
        safe_gap = 22
        candidates: list[dict[str, Any]] = []

        char_options: list[tuple[str, tuple[int, int, int, int], int]] = []
        if panel_w >= 560 or (layout_style == "cinematic_grid" and panel_w >= 500) or (layout_style == "action_splash" and panel_w >= 480):
            char_w = self._clamp(int(panel_w * 0.24), 130, 270 if panel_w >= 900 else 210)
            char_y = content_top + 18
            char_h = max(72, content_bottom - char_y - 18)
            char_options.append(("left", (inner_x1 + 22, char_y, inner_x1 + 22 + char_w, char_y + char_h), 90))
            if panel_w >= 820:
                char_x = inner_x2 - 22 - char_w
                char_options.append(("right", (char_x, char_y, char_x + char_w, char_y + char_h), 58))
        elif panel_w >= 340 and content_h >= 125:
            char_w = self._clamp(int(panel_w * 0.28), 104, 132)
            char_y = content_top + 10
            char_h = max(72, content_bottom - char_y - 10)
            char_options.append(("left", (inner_x1 + 8, char_y, inner_x1 + 8 + char_w, char_y + char_h), 86))
            if content_w >= 390:
                char_x = inner_x2 - 8 - char_w
                char_options.append(("right", (char_x, char_y, char_x + char_w, char_y + char_h), 54))
        else:
            char_w = max(84, min(content_w - 30, 170))
            bubble_budget_h = self._clamp(int(content_h * 0.48), 88, 140)
            char_h = max(64, content_h - bubble_budget_h - 20)
            char_x = inner_x1 + (content_w - char_w) // 2
            char_y = content_bottom - char_h - 8
            char_options.append(("bottom", (char_x, char_y, char_x + char_w, char_y + char_h), 72))

        for char_side, char_rect, char_score in char_options:
            char_safe = self._inflate_rect(char_rect, safe_gap)
            zones = [
                ("right", (char_rect[2] + safe_gap, content_top + 8, inner_x2 - 4, content_bottom - 8), 72 if char_side == "left" else 32),
                ("left", (inner_x1 + 4, content_top + 8, char_rect[0] - safe_gap, content_bottom - 8), 72 if char_side == "right" else 32),
                ("top", (inner_x1 + 4, content_top + 4, inner_x2 - 4, char_rect[1] - safe_gap), 44 if char_side == "bottom" else 18),
            ]
            for zone_name, zone, zone_score in zones:
                zx1, zy1, zx2, zy2 = zone
                zone_w = zx2 - zx1
                zone_h = zy2 - zy1
                if zone_w < 138 or zone_h < 86:
                    continue
                start_size = 42 if zone_w > 760 else 36 if zone_w > 430 else 30 if zone_w > 240 else 24
                if zone_name == "top" and zone_w > 430:
                    start_size = max(start_size, 34)
                max_lines = 4 if zone_w > 520 and zone_h > 150 else 3
                bubble_h, fitted_size = self._estimate_bubble_fit(
                    draw,
                    dialogue,
                    zone_w,
                    zone_h,
                    bubble_type,
                    body_family,
                    start_size=start_size,
                    min_size=16 if zone_w <= 260 else 18,
                    max_lines=max_lines,
                )
                if zone_name in {"left", "right"}:
                    speaker_y = char_rect[1] + int((char_rect[3] - char_rect[1]) * 0.36)
                    by1 = self._clamp(speaker_y - bubble_h // 2, zy1, max(zy1, zy2 - bubble_h))
                else:
                    by1 = zy1
                bubble_rect = (zx1, by1, zx2, by1 + bubble_h)
                overlap = self._rect_intersection_area(char_safe, bubble_rect)
                if overlap:
                    continue
                area_score = (zone_w * bubble_h) // 120
                balance_penalty = abs(((bubble_rect[0] + bubble_rect[2]) // 2) - ((inner_x1 + inner_x2) // 2)) // 12
                score = char_score + zone_score + area_score + fitted_size * 7 - balance_penalty
                candidates.append(
                    {
                        "score": score,
                        "character": char_rect,
                        "bubble": bubble_rect,
                        "zone": zone_name,
                    }
                )

        if not candidates:
            bubble_x1 = inner_x1 + 4
            bubble_x2 = inner_x2 - 4
            available_for_bubble = max(64, content_h - 92)
            bubble_h, _font_size = self._estimate_bubble_fit(
                draw,
                dialogue,
                bubble_x2 - bubble_x1,
                self._clamp(int(content_h * 0.42), 64, min(130, available_for_bubble)),
                bubble_type,
                body_family,
                start_size=28,
                min_size=15,
                max_lines=3,
            )
            bubble_rect = (bubble_x1, content_top + 4, bubble_x2, content_top + 4 + bubble_h)
            char_w = max(72, min(content_w - 30, 150))
            char_x = inner_x1 + (content_w - char_w) // 2
            char_y = bubble_rect[3] + 12
            char_h = max(60, content_bottom - char_y - 12)
            char_rect = (char_x, char_y, char_x + char_w, char_y + char_h)
        else:
            best = max(candidates, key=lambda item: item["score"])
            char_rect = best["character"]
            bubble_rect = best["bubble"]

        char_rect = (
            max(inner_x1 + 4, char_rect[0]),
            max(content_top + 4, char_rect[1]),
            min(inner_x2 - 4, char_rect[2]),
            min(content_bottom - 4, char_rect[3]),
        )
        bubble_rect = (
            max(inner_x1 + 4, bubble_rect[0]),
            max(content_top + 4, bubble_rect[1]),
            min(inner_x2 - 4, bubble_rect[2]),
            min(content_bottom - 4, bubble_rect[3]),
        )
        if self._rect_intersection_area(self._inflate_rect(char_rect, safe_gap), bubble_rect):
            if char_rect[2] + safe_gap + 138 <= inner_x2 - 4:
                bubble_rect = (char_rect[2] + safe_gap, bubble_rect[1], inner_x2 - 4, bubble_rect[3])
            elif inner_x1 + 4 <= char_rect[0] - safe_gap - 138:
                bubble_rect = (inner_x1 + 4, bubble_rect[1], char_rect[0] - safe_gap, bubble_rect[3])
            else:
                bubble_h = min(max(64, (content_bottom - content_top) // 3), 112)
                bubble_rect = (inner_x1 + 4, content_top + 4, inner_x2 - 4, content_top + 4 + bubble_h)
                char_w = max(72, min(content_w - 30, 150))
                char_x = inner_x1 + (content_w - char_w) // 2
                char_y = min(bubble_rect[3] + 12, content_bottom - 64)
                char_h = max(60, content_bottom - char_y - 12)
                char_rect = (char_x, char_y, char_x + char_w, min(content_bottom - 4, char_y + char_h))

        head_x = (char_rect[0] + char_rect[2]) // 2
        head_y = char_rect[1] + max(28, int((char_rect[3] - char_rect[1]) * 0.34))
        bubble_cx = (bubble_rect[0] + bubble_rect[2]) // 2
        bubble_cy = (bubble_rect[1] + bubble_rect[3]) // 2
        if bubble_cx >= char_rect[2]:
            tail_anchor = (char_rect[2] + 4, head_y)
        elif bubble_cx <= char_rect[0]:
            tail_anchor = (char_rect[0] - 4, head_y)
        elif bubble_cy < char_rect[1]:
            tail_anchor = (head_x, char_rect[1] - 4)
        else:
            tail_anchor = (head_x, char_rect[3] + 4)
        return {
            "character": (char_rect[0], char_rect[1], char_rect[2] - char_rect[0], char_rect[3] - char_rect[1]),
            "bubble": bubble_rect,
            "tail_anchor": tail_anchor,
            "name_x": char_rect[0],
            "name_y": min(content_bottom - 24, char_rect[3] - 24),
        }

    def _draw_dataset_title_banner(
        self,
        img,
        draw,
        x: int,
        y: int,
        w: int,
        h: int,
        title: str,
        subject: str,
        palette: dict[str, str],
        style_record: dict[str, Any],
        cast_sources: list[str],
    ) -> None:
        title_family = str(style_record.get("title_font_family") or "display")
        body_family = str(style_record.get("body_font_family") or "sans")
        theme_label = str(style_record.get("label") or style_record.get("theme_profile") or "Comic Style")
        paper_fill = self._blend_hex(palette["page"], "#ffffff", 0.34)
        soft_accent = self._blend_hex(palette["title_stripe"], "#ffffff", 0.18)
        draw.rounded_rectangle((x + 5, y + 6, x + w + 5, y + h + 6), radius=16, fill=self._blend_hex("#111111", paper_fill, 0.88))
        self._gradient_rect(img, x, y, x + w, y + h, paper_fill, "#ffffff", vertical=True)
        draw.rounded_rectangle((x, y, x + w, y + h), radius=16, outline="#111111", width=4, fill=None)
        draw.rectangle((x + 6, y + 6, x + w - 6, y + 24), fill=soft_accent)
        draw.line((x + 6, y + 25, x + w - 6, y + 25), fill="#111111", width=2)
        draw.rectangle((x + 18, y + 38, x + 150, y + h - 20), fill=self._blend_hex(palette["title_stripe"], "#ffffff", 0.42), outline="#111111", width=2)
        draw.text((x + 84, y + 58), "COMIC", font=self._font(19, bold=True, family=body_family), fill="#111111", anchor="mm")
        draw.text((x + 84, y + 86), "STUDY", font=self._font(19, bold=True, family=body_family), fill="#111111", anchor="mm")

        badge_font = self._font(18, bold=True, family=body_family)
        info_font = self._font(12, bold=True, family=body_family)
        title_text, title_font = self._fit_text_block(
            draw,
            title.upper(),
            w - 220,
            54,
            start_size=50,
            min_size=24,
            bold=True,
            family=title_family,
            max_lines=1,
        )
        title_x = x + 176
        title_y = y + 42
        draw.text((title_x + 3, title_y + 3), title_text, font=title_font, fill=self._blend_hex("#111111", paper_fill, 0.72))
        draw.text((title_x, title_y), title_text, font=title_font, fill="#111111")

        badge_text = f" {subject.upper()}  |  {theme_label.upper()} "
        badge_w = self._text_size(draw, badge_text, badge_font)[0] + 22
        badge_y = y + 94
        draw.rounded_rectangle((title_x, badge_y, title_x + badge_w, badge_y + 31), radius=8, fill=soft_accent, outline="#111111", width=2)
        draw.text((title_x + 12, badge_y + 6), badge_text, font=badge_font, fill="#111111")
        source_label = " + ".join("ARCHIVE HEROES" if source == "archive" else "COMINOTE DATASET" for source in cast_sources)
        info_text = source_label or "COMINOTE DATASET"
        info_text, info_font = self._fit_text_block(
            draw,
            info_text,
            300,
            16,
            start_size=12,
            min_size=10,
            bold=True,
            family=body_family,
            max_lines=1,
        )
        info_width, _ = self._text_size(draw, info_text, info_font)
        draw.text((x + w - info_width - 22, y + h - 23), info_text, font=info_font, fill="#11111188")

    def _stylize_image_for_theme(
        self,
        source_image: Any,
        visual: dict[str, Any],
        panel_theme: dict[str, str],
    ) -> Any:
        theme_profile = visual.get("theme_profile") or {}
        theme_slug = str(theme_profile.get("slug") or "").lower()
        image = source_image.convert("RGB").copy()

        factors = {
            "anime": (1.32, 1.18, 1.28, 4),
            "superhero": (1.42, 1.32, 1.35, 4),
            "manga": (0.15, 1.42, 1.5, 3),
            "fantasy": (1.24, 1.16, 1.2, 4),
            "pixar": (1.2, 1.12, 1.18, 5),
            "cartoon_kids": (1.36, 1.18, 1.22, 4),
            "horror": (0.62, 1.28, 1.32, 3),
        }
        color_factor, contrast_factor, sharpness_factor, poster_bits = factors.get(
            theme_slug,
            factors["cartoon_kids"],
        )

        if theme_slug in {"manga", "horror"}:
            gray = ImageOps.grayscale(image)
            image = ImageOps.colorize(
                gray,
                black="#151515",
                white=self._blend_hex(panel_theme["surface"], "#ffffff", 0.24),
            ).convert("RGB")
        else:
            image = ImageEnhance.Color(image).enhance(color_factor)
        image = ImageEnhance.Contrast(image).enhance(contrast_factor)
        image = ImageEnhance.Sharpness(image).enhance(sharpness_factor)
        image = ImageOps.posterize(image, poster_bits)

        if ImageFilter is not None:
            edges = source_image.convert("L").filter(ImageFilter.FIND_EDGES)
            threshold = 42 if theme_slug != "manga" else 28
            edge_mask = edges.point(lambda value: 190 if value > threshold else 0).convert("L")
            line_color = "#111111" if theme_slug in {"manga", "superhero", "horror"} else panel_theme["accent"]
            line_layer = Image.new("RGB", image.size, line_color)
            image = Image.composite(line_layer, image, edge_mask)
        return image

    def _paste_contained_image(
        self,
        base_img: Any,
        draw,
        source_image: Any,
        rect: tuple[int, int, int, int],
        *,
        fill: str,
        accent: str,
    ) -> None:
        x1, y1, x2, y2 = rect
        if x2 <= x1 + 10 or y2 <= y1 + 10:
            return
        draw.rounded_rectangle((x1 + 5, y1 + 6, x2 + 5, y2 + 6), radius=12, fill="#11111144")
        draw.rounded_rectangle((x1, y1, x2, y2), radius=12, fill=fill, outline="#111111", width=3)
        draw.rounded_rectangle((x1 + 7, y1 + 7, x2 - 7, y2 - 7), radius=8, outline=accent, width=2)

        pad = 14
        canvas_w = max(1, x2 - x1 - pad * 2)
        canvas_h = max(1, y2 - y1 - pad * 2)
        art = source_image.copy()
        art.thumbnail((canvas_w, canvas_h), Image.Resampling.LANCZOS)
        paste_x = x1 + pad + (canvas_w - art.width) // 2
        paste_y = y1 + pad + (canvas_h - art.height) // 2
        base_img.paste(art, (paste_x, paste_y))
        draw.rounded_rectangle((paste_x, paste_y, paste_x + art.width, paste_y + art.height), radius=6, outline="#111111", width=2)

    def _image_panel_regions(
        self,
        px: int,
        py: int,
        pw: int,
        ph: int,
        header_h: int,
        caption_h: int,
    ) -> tuple[tuple[int, int, int, int], tuple[int, int, int, int], int]:
        content_top = py + header_h + 14
        caption_y1 = py + ph - caption_h - 12
        content_bottom = caption_y1 - 14
        inner_x1 = px + 16
        inner_x2 = px + pw - 16

        if pw >= 620:
            image_x2 = px + int(pw * 0.64)
            image_rect = (inner_x1, content_top, image_x2, content_bottom)
            bubble_rect = (image_x2 + 16, content_top + 8, inner_x2, content_bottom - 8)
        else:
            content_h = max(120, content_bottom - content_top)
            image_h = max(92, int(content_h * 0.58))
            image_rect = (inner_x1, content_top, inner_x2, min(content_bottom - 74, content_top + image_h))
            bubble_rect = (inner_x1 + 8, image_rect[3] + 12, inner_x2 - 8, content_bottom)
        return image_rect, bubble_rect, caption_y1

    def _draw_palette_swatch_strip(
        self,
        draw,
        palette: list[str],
        x: int,
        y: int,
        *,
        max_width: int,
    ) -> None:
        swatch_size = 18
        gap = 5
        for index, color in enumerate(palette[:5]):
            sx = x + index * (swatch_size + gap)
            if sx + swatch_size > x + max_width:
                break
            draw.rounded_rectangle((sx, y, sx + swatch_size, y + swatch_size), radius=5, fill=self._safe_hex(color, "#ffd93d"), outline="#111111", width=2)

    def _draw_image_comic_panel(
        self,
        img,
        draw,
        px: int,
        py: int,
        pw: int,
        ph: int,
        scene: Scene,
        visual: dict[str, Any],
        source_image: Any,
        analysis: ImageComicAnalysis,
        page_palette: dict[str, str],
        style_record: dict[str, Any],
        index: int,
    ) -> None:
        panel_theme = self._panel_theme(page_palette, visual)
        theme_profile = visual.get("theme_profile") or {}
        layout_style = str(style_record.get("layout_style") or theme_profile.get("layout_style") or "balanced_grid")
        title_family = str(style_record.get("title_font_family") or "display")
        body_family = str(style_record.get("body_font_family") or "sans")
        panel_radius = 8 if layout_style == "manga_cascade" else 10
        panel_surface = self._blend_hex(panel_theme["surface"], "#ffffff", 0.2)
        panel_surface_alt = self._blend_hex(panel_theme["surface_alt"], "#ffffff", 0.3)

        draw.rounded_rectangle((px + 6, py + 6, px + pw + 6, py + ph + 6), radius=panel_radius, fill="#11111144")
        self._gradient_rect(img, px, py, px + pw, py + ph, panel_surface, panel_surface_alt)
        self._draw_page_halftone(draw, px + 8, py + 8, px + pw - 8, py + ph - 8, panel_theme["accent"], density=20, dot_r=1.0, opacity=0.032)
        draw.rounded_rectangle((px, py, px + pw, py + ph), radius=panel_radius, outline="#111111", width=4)
        draw.rounded_rectangle((px + 5, py + 5, px + pw - 5, py + ph - 5), radius=max(4, panel_radius - 4), outline="#ffffffaa", width=1)

        header_h = 42 if ph >= 300 else 36
        self._gradient_rect(img, px + 4, py + 4, px + pw - 4, py + header_h, panel_theme["accent"], panel_theme["accent_alt"], vertical=False)
        draw.rectangle((px + 4, py + 4, px + pw - 4, py + header_h), outline="#111111", width=1)
        draw.line([(px + 4, py + header_h), (px + pw - 4, py + header_h)], fill="#111111", width=2)

        badge_size = 26
        badge_x = px + 14
        badge_y = py + 8
        badge_font = self._font(14, bold=True, family=body_family)
        draw.ellipse((badge_x, badge_y, badge_x + badge_size, badge_y + badge_size), fill="#ffffff", outline="#111111", width=2)
        draw.text((badge_x + badge_size // 2, badge_y + badge_size // 2), str(index + 1), font=badge_font, fill="#111111", anchor="mm")

        title_text, title_font = self._fit_text_block(
            draw,
            scene.title.upper(),
            max(80, pw - 92),
            21,
            start_size=20,
            min_size=11,
            bold=True,
            family=title_family,
            max_lines=1,
        )
        draw.text((px + 50, py + 11), title_text, font=title_font, fill="#ffffff")

        caption_preview, caption_font = self._fit_text_block(
            draw,
            scene.caption,
            max(90, pw - 64),
            78,
            start_size=18,
            min_size=14,
            bold=True,
            family=body_family,
            spacing=4,
            max_lines=2,
        )
        _caption_w, caption_text_h = self._text_size(draw, caption_preview, caption_font, spacing=4)
        caption_h = self._clamp(caption_text_h + 40, 62, 98 if ph >= 320 else 84)
        image_rect, bubble_rect, caption_y1 = self._image_panel_regions(px, py, pw, ph, header_h, caption_h)

        styled = self._stylize_image_for_theme(source_image, visual, panel_theme)
        self._paste_contained_image(
            img,
            draw,
            styled,
            image_rect,
            fill=self._blend_hex(panel_theme["surface"], "#ffffff", 0.35),
            accent=panel_theme["accent_alt"],
        )
        styled.close()

        if image_rect[2] - image_rect[0] >= 170:
            self._draw_palette_swatch_strip(
                draw,
                analysis.palette,
                image_rect[0] + 16,
                image_rect[1] + 15,
                max_width=image_rect[2] - image_rect[0] - 32,
            )

        self._draw_dataset_speech_bubble(
            draw,
            visual.get("bubble") or {},
            bubble_rect[0],
            bubble_rect[1],
            bubble_rect[2],
            bubble_rect[3],
            scene.dialogue,
            panel_theme,
            body_family,
            speaker_anchor=(image_rect[2] - 18, image_rect[1] + max(30, (image_rect[3] - image_rect[1]) // 2)),
        )

        name_text = (visual.get("character") or {}).get("name") or scene.speaker
        badge_label, name_font = self._fit_text_block(
            draw,
            name_text.upper(),
            max(90, image_rect[2] - image_rect[0] - 30),
            13,
            start_size=11,
            min_size=8,
            bold=True,
            family=body_family,
            max_lines=1,
        )
        name_y1 = max(image_rect[1] + 12, image_rect[3] - 30)
        badge_w = self._text_size(draw, badge_label, name_font)[0] + 18
        badge_x2 = min(image_rect[2] - 12, image_rect[0] + 12 + badge_w)
        draw.rounded_rectangle((image_rect[0] + 12, name_y1, badge_x2, name_y1 + 20), radius=8, fill=panel_theme["accent"], outline="#111111", width=2)
        draw.text((image_rect[0] + 21, name_y1 + 2), badge_label, font=name_font, fill="#ffffff")

        self._draw_dataset_caption(
            draw,
            px + 12,
            caption_y1,
            px + pw - 12,
            py + ph - 12,
            scene.caption,
            panel_theme,
            body_family,
        )

    def _draw_dataset_panel(
        self,
        img,
        draw,
        px: int,
        py: int,
        pw: int,
        ph: int,
        scene: Scene,
        visual: dict[str, Any],
        page_palette: dict[str, str],
        style_record: dict[str, Any],
        index: int,
    ) -> None:
        panel_theme = self._panel_theme(page_palette, visual)
        theme_profile = visual.get("theme_profile") or {}
        layout_style = str(style_record.get("layout_style") or theme_profile.get("layout_style") or "balanced_grid")
        title_family = str(style_record.get("title_font_family") or "display")
        body_family = str(style_record.get("body_font_family") or "sans")
        bubble_type = str((visual.get("bubble") or {}).get("type") or "normal").lower()
        panel_radius = 8 if layout_style == "manga_cascade" else 10
        shadow_offset = 6
        panel_surface = self._blend_hex(panel_theme["surface"], "#ffffff", 0.28)
        panel_surface_alt = self._blend_hex(panel_theme["surface_alt"], "#ffffff", 0.36)
        panel_shadow = self._blend_hex("#111111", self._safe_hex(page_palette.get("page"), "#fff8ef"), 0.82)
        draw.rounded_rectangle((px + shadow_offset, py + shadow_offset, px + pw + shadow_offset, py + ph + shadow_offset), radius=panel_radius, fill=panel_shadow)
        self._gradient_rect(img, px, py, px + pw, py + ph, panel_surface, panel_surface_alt)
        self._draw_page_halftone(draw, px + 8, py + 8, px + pw - 8, py + ph - 8, panel_theme["accent_alt"], density=22, dot_r=1.2, opacity=0.025)
        self._draw_dataset_background(
            draw,
            img,
            px,
            py,
            pw,
            ph,
            visual.get("background") or {},
            panel_theme,
            theme_profile,
        )

        # Clean newspaper-strip panel borders.
        draw.rounded_rectangle((px, py, px + pw, py + ph), radius=panel_radius, outline="#111111", width=4)
        draw.rounded_rectangle((px + 5, py + 5, px + pw - 5, py + ph - 5), radius=max(4, panel_radius - 4), outline="#ffffffaa", width=1)

        header_h = 42 if ph >= 300 else 36
        header_fill = self._blend_hex(panel_theme["accent"], "#ffffff", 0.18)
        self._gradient_rect(img, px + 4, py + 4, px + pw - 4, py + header_h, header_fill, panel_theme["accent_alt"], vertical=False)
        draw.rectangle((px + 4, py + 4, px + pw - 4, py + header_h), outline="#111111", width=1)
        draw.line([(px + 4, py + header_h), (px + pw - 4, py + header_h)], fill="#111111", width=2)

        badge_font = self._font(14, bold=True, family=body_family)
        badge_size = 26
        badge_x = px + 14
        badge_y = py + 8
        draw.ellipse((badge_x, badge_y, badge_x + badge_size, badge_y + badge_size), fill="#ffffff", outline="#111111", width=2)
        draw.text((badge_x + badge_size // 2, badge_y + badge_size // 2), str(index + 1), font=badge_font, fill="#111111", anchor="mm")
        inner_x1 = px + 18
        inner_x2 = px + pw - 18
        content_top = py + header_h + 12
        caption_max_height = 108 if ph >= 320 else 92 if ph >= 260 else 72
        caption_max_lines = 2 if ph >= 230 else 1
        caption_width = max(90, pw - 64)
        caption_preview, caption_font = self._fit_text_block(
            draw,
            scene.caption,
            caption_width - 18,
            caption_max_height - 28,
            start_size=20 if ph >= 260 else 18,
            min_size=16,
            bold=True,
            family=body_family,
            spacing=4,
            max_lines=caption_max_lines,
        )
        _, caption_text_h = self._text_size(draw, caption_preview, caption_font, spacing=4)
        caption_h = self._clamp(caption_text_h + 40, 60, caption_max_height)
        caption_y1 = py + ph - caption_h - 12
        content_bottom = caption_y1 - 14

        title_text, title_font = self._fit_text_block(
            draw,
            scene.title.upper(),
            max(80, pw - 92),
            21,
            start_size=20,
            min_size=11,
            bold=True,
            family=title_family,
            max_lines=1,
        )
        draw.text((px + 50, py + 11), title_text, font=title_font, fill="#ffffff")

        composition = self._smart_panel_composition(
            draw,
            scene.dialogue,
            bubble_type,
            body_family,
            inner_x1,
            inner_x2,
            content_top,
            content_bottom,
            pw,
            ph,
            layout_style,
            index,
        )
        char_x, char_y, char_w, char_h = composition["character"]
        bubble_x1, bubble_y1, bubble_x2, bubble_y2 = composition["bubble"]
        name_y1 = composition["name_y"]
        name_x1 = composition["name_x"]
        tail_anchor = composition["tail_anchor"]

        self._draw_dataset_character(
            draw,
            char_x,
            char_y,
            char_w,
            char_h,
            visual,
            style_record,
            panel_theme,
        )

        self._draw_dataset_speech_bubble(
            draw,
            visual.get("bubble") or {},
            bubble_x1,
            bubble_y1,
            bubble_x2,
            bubble_y2,
            scene.dialogue,
            panel_theme,
            body_family,
            speaker_anchor=tail_anchor,
        )

        name_text = (visual.get("character") or {}).get("name") or scene.speaker
        badge_label, badge_font = self._fit_text_block(
            draw,
            name_text.upper(),
            max(90, pw - 82),
            13,
            start_size=11,
            min_size=8,
            bold=True,
            family=body_family,
            max_lines=1,
        )
        badge_w = self._text_size(draw, badge_label, badge_font)[0] + 18
        badge_x1 = self._clamp(int(name_x1), inner_x1, max(inner_x1, inner_x2 - badge_w))
        badge_x2 = min(inner_x2, badge_x1 + badge_w)
        draw.rounded_rectangle((badge_x1 + 2, name_y1 + 2, badge_x2 + 2, name_y1 + 20), radius=8, fill=self._blend_hex("#111111", panel_surface, 0.78))
        draw.rounded_rectangle((badge_x1, name_y1, badge_x2, name_y1 + 20), radius=8, fill=panel_theme["accent"], outline="#111111", width=2)
        draw.text((badge_x1 + 9, name_y1 + 2), badge_label, font=badge_font, fill="#ffffff")

        self._draw_dataset_caption(
            draw,
            px + 12,
            caption_y1,
            px + pw - 12,
            py + ph - 12,
            scene.caption,
            panel_theme,
            body_family,
        )

    def _draw_dataset_background(
        self,
        draw,
        img,
        px: int,
        py: int,
        pw: int,
        ph: int,
        background: dict[str, Any],
        panel_theme: dict[str, str],
        theme_profile: dict[str, Any] | None = None,
    ) -> None:
        bg_palette = background.get("palette") or {}
        sky_top = self._blend_hex(self._safe_hex(bg_palette.get("primary"), panel_theme["surface"]), "#ffffff", 0.34)
        sky_bottom = self._blend_hex(self._safe_hex(bg_palette.get("secondary"), panel_theme["surface_alt"]), "#ffffff", 0.42)
        accent = self._safe_hex(bg_palette.get("accent"), panel_theme["accent"])
        floor = self._blend_hex(self._safe_hex(bg_palette.get("floor"), panel_theme["floor"]), "#ffffff", 0.28)
        theme_slug = str((theme_profile or {}).get("slug") or "").lower()
        horizon = py + int(ph * 0.72)

        self._gradient_rect(img, px + 4, py + 44, px + pw - 4, horizon, sky_top, sky_bottom)
        self._gradient_rect(img, px + 4, horizon, px + pw - 4, py + ph - 40, floor, sky_bottom)

        for cloud_x in range(px + 44, px + pw - 50, 190):
            cloud_y = py + 66 + (cloud_x % 28)
            draw.ellipse((cloud_x, cloud_y, cloud_x + 64, cloud_y + 34), fill="#ffffff5c")
            draw.ellipse((cloud_x + 22, cloud_y - 12, cloud_x + 86, cloud_y + 30), fill="#ffffff5c")

        category = (background.get("category") or "school").lower()
        blob = f"{background.get('name', '')} {background.get('prompt_fragment', '')} {' '.join(background.get('tags') or [])}".lower()
        if category == "space" or "space" in blob or "rocket" in blob:
            draw.ellipse((px + pw - 260, py + 72, px + pw - 110, py + 222), fill=self._safe_hex(bg_palette.get("secondary"), "#223f88"), outline="#111111", width=3)
            for star_x in range(px + 30, px + pw - 30, 58):
                draw.ellipse((star_x, py + 62 + (star_x % 38), star_x + 4, py + 66 + (star_x % 38)), fill="#ffffff")
            draw.rectangle((px + pw - 300, horizon - 40, px + pw - 36, horizon + 6), fill="#dbe4ff", outline="#111111", width=2)
            draw.rectangle((px + pw - 278, horizon - 120, px + pw - 58, horizon - 46), fill="#8ab4ff", outline="#111111", width=2)
        elif category == "science" or "lab" in blob or "science" in blob:
            bench_y = horizon - 34
            draw.rectangle((px + pw - 280, bench_y, px + pw - 28, horizon + 4), fill="#f0f5ff", outline="#111111", width=2)
            for offset, color in zip((0, 42, 88), ("#66d9ff", "#ff8ab3", "#72f5ab")):
                tx = px + pw - 250 + offset
                draw.rounded_rectangle((tx, bench_y - 78, tx + 20, bench_y), radius=5, fill=color, outline="#111111", width=2)
                draw.ellipse((tx - 6, bench_y - 90, tx + 26, bench_y - 60), fill=f"{color}88")
            draw.rectangle((px + pw - 162, py + 76, px + pw - 34, py + 186), fill="#ffffff88", outline="#111111", width=2)
            draw.text((px + pw - 150, py + 92), "LAB", font=self._font(20, bold=True), fill=accent)
        elif category == "historical" or "palace" in blob or "market" in blob:
            for col_x in (px + pw - 280, px + pw - 188, px + pw - 96):
                draw.rounded_rectangle((col_x, py + 62, col_x + 28, horizon + 10), radius=8, fill="#e8d7b4", outline="#111111", width=2)
                draw.ellipse((col_x - 10, py + 48, col_x + 38, py + 84), fill="#f0e0bb", outline="#111111", width=2)
            draw.polygon([(px + 22, horizon), (px + 120, py + 128), (px + 220, horizon)], fill="#ffffff22")
        elif category == "nature" or "mountain" in blob or "garden" in blob or "jungle" in blob:
            draw.polygon([(px, horizon), (px + 110, py + 170), (px + 240, horizon)], fill="#84c98d")
            draw.polygon([(px + 150, horizon), (px + 320, py + 140), (px + 470, horizon)], fill="#6ab37a")
            for tree_x in range(px + pw - 260, px + pw - 40, 82):
                draw.rectangle((tree_x + 16, horizon - 42, tree_x + 32, horizon + 8), fill="#7c5a36")
                draw.ellipse((tree_x, horizon - 88, tree_x + 48, horizon - 24), fill="#4caf60")
        elif category == "fantasy" or "underwater" in blob:
            if "underwater" in blob:
                for bubble_x in range(px + pw - 260, px + pw - 30, 34):
                    draw.ellipse((bubble_x, py + 88 + (bubble_x % 28), bubble_x + 16, py + 104 + (bubble_x % 28)), outline="#ffffff", width=2)
                draw.arc((px + pw - 230, horizon - 70, px + pw - 110, horizon + 20), start=0, end=180, fill="#ff8ab3", width=8)
                draw.arc((px + pw - 154, horizon - 64, px + pw - 34, horizon + 16), start=0, end=180, fill="#ffd166", width=8)
            else:
                for sparkle_x in range(px + 60, px + pw - 40, 110):
                    self._draw_starburst(draw, sparkle_x, py + 120 + (sparkle_x % 40), 10, "#ffffff", accent, spikes=6)
        else:
            if "library" in blob or "book" in blob:
                shelf_x = px + pw - 220
                for row_index in range(2):
                    shelf_y = py + 88 + row_index * 80
                    draw.rectangle((shelf_x, shelf_y, shelf_x + 170, shelf_y + 10), fill="#7d5a37")
                    book_x = shelf_x + 6
                    for color in ("#ff9f1c", "#2ec4b6", "#e71d36", "#ffbf69", "#4361ee"):
                        draw.rectangle((book_x, shelf_y - 52, book_x + 24, shelf_y), fill=color, outline="#111111", width=1)
                        book_x += 28
            else:
                board_x1 = px + pw - 230
                board_y1 = py + 72
                board_x2 = px + pw - 34
                board_y2 = horizon - 28
                draw.rectangle((board_x1, board_y1, board_x2, board_y2), fill="#18653d", outline="#111111", width=3)
                draw.text((board_x1 + 16, board_y1 + 18), "IDEA", font=self._font(18, bold=True), fill="#ffffff")
                draw.line([(board_x1 + 18, board_y1 + 52), (board_x2 - 18, board_y1 + 52)], fill="#ffffffaa", width=2)

        if theme_slug == "anime":
            for petal_x in range(px + 44, px + pw - 24, 150):
                petal_y = py + 84 + (petal_x % 36)
                draw.ellipse((petal_x, petal_y, petal_x + 10, petal_y + 6), fill="#ffd1ea88")
                draw.ellipse((petal_x + 6, petal_y + 2, petal_x + 16, petal_y + 8), fill="#ffe3f288")
        elif theme_slug == "pixar":
            for orb_x in range(px + 44, px + pw - 30, 190):
                orb_y = py + 90 + (orb_x % 28)
                draw.ellipse((orb_x, orb_y, orb_x + 30, orb_y + 30), fill="#ffffff2f")
                draw.ellipse((orb_x + 6, orb_y + 6, orb_x + 16, orb_y + 16), fill="#ffffff66")
        elif theme_slug == "superhero":
            skyline_y = horizon - 20
            skyline = [
                (px + 18, skyline_y), (px + 18, skyline_y - 72), (px + 52, skyline_y - 72), (px + 52, skyline_y - 28),
                (px + 88, skyline_y - 28), (px + 88, skyline_y - 96), (px + 122, skyline_y - 96), (px + 122, skyline_y - 44),
                (px + 156, skyline_y - 44), (px + 156, skyline_y - 124), (px + 192, skyline_y - 124), (px + 192, skyline_y),
            ]
            draw.polygon(skyline + [(px + 18, skyline_y)], fill=self._blend_hex("#1a1a1a", sky_bottom, 0.62))
        elif theme_slug == "manga":
            for offset in range(0, pw, 38):
                draw.line([(px + offset, py + 50), (px + offset - 80, py + ph - 40)], fill="#1111110f", width=1)
            draw.rectangle((px + 12, py + 58, px + 78, py + 74), fill="#11111122")
            draw.rectangle((px + 12, py + 78, px + 128, py + 88), fill="#11111111")
        elif theme_slug == "cartoon_kids":
            bunting_y = py + 62
            for bx in range(px + 28, px + pw - 28, 86):
                draw.line([(bx, bunting_y), (bx + 18, bunting_y + 14)], fill=accent, width=2)
                draw.polygon([(bx + 16, bunting_y + 12), (bx + 34, bunting_y + 18), (bx + 22, bunting_y + 32)], fill="#ffffff55", outline=accent)
        elif theme_slug == "fantasy":
            for arch_x in range(px + pw - 280, px + pw - 20, 86):
                draw.arc((arch_x, py + 82, arch_x + 46, py + 158), start=180, end=360, fill="#ffffff88", width=3)
                draw.line([(arch_x + 4, py + 120), (arch_x + 4, horizon)], fill="#ffffff66", width=2)
                draw.line([(arch_x + 42, py + 120), (arch_x + 42, horizon)], fill="#ffffff66", width=2)
        elif theme_slug == "horror":
            draw.ellipse((px + pw - 130, py + 60, px + pw - 58, py + 132), fill="#f8fafcbb", outline="#111827", width=2)
            for fog_y in range(horizon - 76, horizon + 18, 42):
                draw.arc((px + 22, fog_y, px + pw - 30, fog_y + 70), start=182, end=358, fill="#ffffff55", width=5)

    def _visual_identity_blob(self, visual: dict[str, Any]) -> str:
        character = visual.get("character") or {}
        background = visual.get("background") or {}
        profile = visual.get("theme_profile") or {}
        parts = [
            visual.get("theme_render_id", ""),
            profile.get("slug", ""),
            character.get("name", ""),
            character.get("sub_type", ""),
            character.get("render_variant", ""),
            " ".join(character.get("tags") or []),
            background.get("name", ""),
            background.get("category", ""),
            background.get("prompt_fragment", ""),
            " ".join(background.get("tags") or []),
            " ".join(visual.get("theme_comic_elements") or []),
        ]
        return " ".join(str(part) for part in parts if part).lower()

    def _draw_premium_panel_atmosphere(
        self,
        draw,
        img,
        px: int,
        py: int,
        pw: int,
        ph: int,
        visual: dict[str, Any],
        panel_theme: dict[str, str],
        index: int,
    ) -> None:
        blob = self._visual_identity_blob(visual)
        accent = panel_theme["accent"]
        accent_alt = panel_theme["accent_alt"]
        horizon = py + int(ph * 0.72)

        if any(term in blob for term in ("pokemon", "pikachu", "electric", "lightning")):
            for bolt_x in range(px + 72 + (index % 2) * 28, px + pw - 80, 160):
                bolt = [
                    (bolt_x, py + 70), (bolt_x + 28, py + 122), (bolt_x + 10, py + 122),
                    (bolt_x + 42, py + 190), (bolt_x - 6, py + 112), (bolt_x + 14, py + 112),
                ]
                draw.polygon(bolt, fill="#fde047cc", outline="#111111")
            for ring_x in range(px + 58, px + pw - 40, 180):
                draw.ellipse((ring_x, horizon - 82, ring_x + 62, horizon - 20), outline=accent_alt, width=4)
        elif any(term in blob for term in ("marvel", "spider", "web")):
            web_cx, web_cy = px + pw - 68, py + 74
            for ray in range(0, 360, 30):
                ex = web_cx + int(150 * math.cos(math.radians(ray)))
                ey = web_cy + int(110 * math.sin(math.radians(ray)))
                draw.line([(web_cx, web_cy), (ex, ey)], fill="#ffffff80", width=2)
            for radius in (28, 58, 88):
                draw.ellipse((web_cx - radius, web_cy - radius, web_cx + radius, web_cy + radius), outline="#ffffff70", width=2)
        elif any(term in blob for term in ("dc", "batman", "gotham", "night")):
            signal_cx, signal_cy = px + pw - 112, py + 98
            draw.ellipse((signal_cx - 52, signal_cy - 36, signal_cx + 52, signal_cy + 36), fill="#facc1544", outline="#111827", width=3)
            draw.polygon(
                [(signal_cx - 30, signal_cy), (signal_cx - 12, signal_cy - 10), (signal_cx, signal_cy - 2), (signal_cx + 12, signal_cy - 10), (signal_cx + 30, signal_cy), (signal_cx + 8, signal_cy + 8), (signal_cx, signal_cy + 2), (signal_cx - 8, signal_cy + 8)],
                fill="#111827cc",
            )
        elif any(term in blob for term in ("naruto", "ninja", "shinobi", "leaf")):
            for leaf_x in range(px + 40, px + pw - 20, 70):
                leaf_y = py + 74 + ((leaf_x + index * 17) % 150)
                draw.arc((leaf_x, leaf_y, leaf_x + 26, leaf_y + 18), 20, 310, fill=accent, width=3)
            draw.line([(px + 26, py + ph - 82), (px + pw - 26, py + 80)], fill="#ffffff66", width=6)
        elif any(term in blob for term in ("dragon", "goku", "saiyan", "energy", "kame")):
            cx, cy = px + int(pw * 0.24), py + int(ph * 0.5)
            for radius in (62, 96, 132):
                draw.arc((cx - radius, cy - radius, cx + radius, cy + radius), 210, 330, fill="#fde047aa", width=5)
            self._draw_speed_lines(draw, cx, cy, pw, ph, accent, count=18, opacity=0.08)
        elif any(term in blob for term in ("space", "astronaut", "galaxy", "orbit", "station")):
            for star_x in range(px + 42, px + pw - 34, 54):
                star_y = py + 62 + ((star_x * 7 + index * 23) % max(60, ph - 130))
                self._draw_starburst(draw, star_x, star_y, 6, "#ffffff", accent_alt, spikes=5)
            draw.arc((px + pw - 240, py + 92, px + pw - 44, py + 230), 195, 350, fill="#ffffffaa", width=5)
            draw.arc((px + pw - 270, py + 118, px + pw - 24, py + 260), 190, 342, fill=accent_alt, width=3)
        elif any(term in blob for term in ("fantasy", "magic", "dragon", "castle", "wizard", "quest")):
            for rune_x in range(px + 52, px + pw - 34, 116):
                rune_y = py + 72 + ((rune_x + index * 19) % 120)
                draw.text((rune_x, rune_y), "✦", font=self._font(28, bold=True, family="display"), fill="#ffffffcc")
                draw.ellipse((rune_x - 12, rune_y - 10, rune_x + 34, rune_y + 36), outline=accent_alt, width=3)
            draw.polygon([(px + pw - 198, horizon), (px + pw - 124, horizon - 86), (px + pw - 52, horizon)], fill="#1111112f")
        elif any(term in blob for term in ("horror", "mystery", "shadow", "fog", "hawkins", "upside")):
            for fog_x in range(px + 24, px + pw - 20, 120):
                draw.arc((fog_x, horizon - 80, fog_x + 180, horizon - 4), 180, 360, fill="#ffffff70", width=5)
            draw.rectangle((px + 4, py + 48, px + pw - 4, py + ph - 36), outline="#11182799", width=5)
        else:
            for sticker_x in range(px + 50, px + pw - 40, 140):
                sticker_y = py + 78 + ((sticker_x + index * 31) % 110)
                self._draw_starburst(draw, sticker_x, sticker_y, 8, "#fff8cf", accent, spikes=6)

    def _draw_character_motion_layer(
        self,
        draw,
        x: int,
        y: int,
        w: int,
        h: int,
        visual: dict[str, Any],
        panel_theme: dict[str, str],
        index: int,
    ) -> None:
        blob = self._visual_identity_blob(visual)
        cx = x + w // 2
        cy = y + max(34, h // 2)
        accent = panel_theme["accent"]
        accent_alt = panel_theme["accent_alt"]
        if any(term in blob for term in ("pokemon", "pikachu", "electric", "lightning")):
            for angle in (20, 76, 128, 205, 292):
                ex = cx + int(w * 0.48 * math.cos(math.radians(angle)))
                ey = cy + int(h * 0.42 * math.sin(math.radians(angle)))
                draw.line([(cx, cy), (ex, ey)], fill="#fde047aa", width=4)
        elif any(term in blob for term in ("superhero", "marvel", "dc", "hero", "gotham")):
            for radius in (max(22, w // 3), max(34, w // 2)):
                draw.arc((cx - radius, cy - radius, cx + radius, cy + radius), 215, 330, fill=accent_alt, width=5)
        elif any(term in blob for term in ("anime", "naruto", "goku", "saiyan", "ninja")):
            for offset in range(-36, 48, 24):
                draw.line([(cx - w // 2, cy + offset), (cx + w // 2, cy + offset - 46)], fill=f"{accent}88", width=4)
        elif any(term in blob for term in ("space", "astronaut", "alien")):
            draw.arc((cx - w // 2, cy - h // 3, cx + w // 2, cy + h // 3), 15, 330, fill=accent_alt, width=4)
        elif any(term in blob for term in ("fantasy", "magic", "wizard", "quest")):
            for radius in (18, 34):
                self._draw_starburst(draw, cx + radius, cy - radius, 7, "#ffffff", accent_alt, spikes=5)
        elif any(term in blob for term in ("horror", "mystery", "shadow")):
            draw.ellipse((cx - w // 2, cy + h // 4, cx + w // 2, cy + h // 2), fill="#11182733")

    def _bubble_tail_points(
        self,
        x1: int,
        y1: int,
        x2: int,
        y2: int,
        speaker_anchor: tuple[int, int] | None,
    ) -> list[tuple[int, int]]:
        if speaker_anchor is None:
            speaker_anchor = (x1 + 32, y2 + 18)
        sx, sy = speaker_anchor
        if sx < x1:
            base_y = self._clamp(sy, y1 + 24, y2 - 24)
            tip_x = min(x1 - 10, max(sx + 22, x1 - 76))
            return [(x1 + 2, base_y - 13), (x1 + 2, base_y + 13), (tip_x, sy)]
        if sx > x2:
            base_y = self._clamp(sy, y1 + 24, y2 - 24)
            tip_x = max(x2 + 10, min(sx - 22, x2 + 76))
            return [(x2 - 2, base_y - 13), (x2 - 2, base_y + 13), (tip_x, sy)]
        if sy < y1:
            base_x = self._clamp(sx, x1 + 28, x2 - 28)
            tip_y = min(y1 - 10, max(sy + 20, y1 - 58))
            return [(base_x - 14, y1 + 2), (base_x + 14, y1 + 2), (sx, tip_y)]
        base_x = self._clamp(sx, x1 + 28, x2 - 28)
        tip_y = max(y2 + 10, min(sy - 20, y2 + 58))
        return [(base_x - 15, y2 - 2), (base_x + 15, y2 - 2), (sx, tip_y)]

    def _draw_bubble_tail(
        self,
        draw,
        x1: int,
        y1: int,
        x2: int,
        y2: int,
        speaker_anchor: tuple[int, int] | None,
        fill: str,
        shadow_fill: str,
        *,
        thought: bool = False,
    ) -> None:
        points = self._bubble_tail_points(x1, y1, x2, y2, speaker_anchor)
        if thought:
            tip_x, tip_y = points[2]
            edge_x = (points[0][0] + points[1][0]) // 2
            edge_y = (points[0][1] + points[1][1]) // 2
            for step, radius in ((0.38, 7), (0.68, 5)):
                cx = int(edge_x + (tip_x - edge_x) * step)
                cy = int(edge_y + (tip_y - edge_y) * step)
                draw.ellipse((cx - radius + 2, cy - radius + 3, cx + radius + 2, cy + radius + 3), fill=shadow_fill)
                draw.ellipse((cx - radius, cy - radius, cx + radius, cy + radius), fill=fill, outline="#111111", width=2)
            return
        draw.polygon([(x + 3, y + 4) for x, y in points], fill=shadow_fill)
        draw.polygon(points, fill=fill, outline="#111111", width=2)

    def _draw_dataset_speech_bubble(
        self,
        draw,
        bubble: dict[str, Any],
        x1: int,
        y1: int,
        x2: int,
        y2: int,
        text: str,
        panel_theme: dict[str, str],
        body_family: str = "sans",
        speaker_anchor: tuple[int, int] | None = None,
    ) -> None:
        if x2 <= x1 or y2 <= y1:
            return
        bubble_type = (bubble.get("type") or "normal").lower()
        if bubble_type == "action_burst":
            bubble_type = "normal"
        fill = self._safe_hex(bubble.get("fill"), panel_theme["bubble"])
        fill = self._blend_hex(fill, "#ffffff", 0.18)
        shadow_fill = self._blend_hex("#111111", fill, 0.84)
        outline_width = 4 if bubble_type == "shouting" else 3
        bubble_w = x2 - x1
        bubble_h = y2 - y1
        if bubble_type == "shouting":
            radius = 16
            draw.rounded_rectangle((x1 + 4, y1 + 4, x2 + 4, y2 + 4), radius=radius, fill=shadow_fill)
            draw.rounded_rectangle((x1, y1, x2, y2), radius=radius, fill=fill, outline="#111111", width=outline_width)
            self._draw_bubble_tail(draw, x1, y1, x2, y2, speaker_anchor, fill, shadow_fill)
        elif bubble_type == "manga_box":
            draw.rounded_rectangle((x1 + 4, y1 + 4, x2 + 4, y2 + 4), radius=8, fill=shadow_fill)
            draw.rounded_rectangle((x1, y1, x2, y2), radius=10, fill=fill, outline="#111111", width=outline_width)
            self._draw_bubble_tail(draw, x1, y1, x2, y2, speaker_anchor, fill, shadow_fill)
        elif bubble_type == "thought_bubble":
            draw.ellipse((x1 + 4, y1 + 4, x2 + 4, y2 + 4), fill=shadow_fill)
            draw.ellipse((x1, y1, x2, y2), fill=fill, outline="#111111", width=outline_width)
            self._draw_bubble_tail(draw, x1, y1, x2, y2, speaker_anchor, fill, shadow_fill, thought=True)
        else:
            radius = 24 if bubble_type == "pixar_round" else 18
            draw.rounded_rectangle((x1 + 4, y1 + 4, x2 + 4, y2 + 4), radius=radius, fill=shadow_fill)
            draw.rounded_rectangle((x1, y1, x2, y2), radius=radius, fill=fill, outline="#111111", width=outline_width)
            draw.rounded_rectangle((x1 + 8, y1 + 8, x2 - 8, y1 + min(34, (y2 - y1) // 3)), radius=max(10, radius - 8), fill="#ffffff55")
            self._draw_bubble_tail(draw, x1, y1, x2, y2, speaker_anchor, fill, shadow_fill)

        if bubble_type in {"shouting", "action_burst"}:
            text_pad_x = max(72, int(bubble_w * 0.24)) if bubble_type == "action_burst" else 72 if bubble_w > 360 else 54 if bubble_w > 220 else 36
            text_pad_y = 34 if bubble_h > 130 else 24
        else:
            text_pad_x = 50 if bubble_w > 360 else 40 if bubble_w > 230 else 28
            text_pad_y = 24 if bubble_h > 118 else 18
        max_text_h = max(30, bubble_h - text_pad_y * 2)
        if bubble_type == "action_burst":
            start_size = 34 if bubble_w > 390 else 30 if bubble_w > 260 else 24
        else:
            start_size = 42 if bubble_w > 390 else 36 if bubble_w > 260 else 30 if bubble_w > 170 else 22
        min_size = 22 if bubble_w > 390 else 18 if bubble_w > 260 else 15
        wrapped, font = self._fit_text_block(
            draw,
            text,
            max(52, bubble_w - text_pad_x * 2),
            max_text_h,
            start_size=start_size,
            min_size=min_size,
            bold=bubble_type in {"shouting", "action_burst"},
            family=body_family,
            spacing=7,
            max_lines=4 if bubble_h > 150 else 3,
        )
        _text_w, text_h = self._text_size(draw, wrapped, font, spacing=7)
        text_y = y1 + max(text_pad_y, (bubble_h - text_h) // 2)
        for line in wrapped.splitlines() or [""]:
            line_w, line_h = self._text_size(draw, line, font)
            text_x = x1 + max(text_pad_x, (bubble_w - line_w) // 2)
            draw.text((text_x, text_y), line, font=font, fill="#111111")
            text_y += line_h + 7

    def _draw_dataset_caption(
        self,
        draw,
        x1: int,
        y1: int,
        x2: int,
        y2: int,
        caption: str,
        panel_theme: dict[str, str],
        body_family: str = "sans",
    ) -> None:
        caption_fill = self._blend_hex(panel_theme["caption"], "#ffffff", 0.38)
        caption_shadow = self._blend_hex("#111111", caption_fill, 0.84)
        draw.rounded_rectangle((x1 + 3, y1 + 4, x2 + 3, y2 + 4), radius=10, fill=caption_shadow)
        draw.rounded_rectangle((x1, y1, x2, y2), radius=10, fill=caption_fill, outline="#111111", width=2)
        draw.rectangle((x1, y1 + 3, x1 + 10, y2 - 3), fill=panel_theme["accent"])
        label_font = self._font(12, bold=True, family=body_family)
        draw.text((x1 + 18, y1 + 6), "STUDY NOTE", font=label_font, fill=panel_theme["caption_text"])
        wrapped, cap_font = self._fit_text_block(
            draw,
            caption,
            max(66, x2 - x1 - 38),
            max(22, y2 - y1 - 32),
            start_size=17,
            min_size=15,
            bold=True,
            family=body_family,
            spacing=5,
            max_lines=2 if (y2 - y1) >= 58 else 1,
        )
        draw.multiline_text((x1 + 18, y1 + 23), wrapped, font=cap_font, fill="#111111", spacing=5)

    @staticmethod
    def _mono_hex(hex_color: str, contrast: float = 0.0) -> str:
        r, g, b = _hex_to_rgb(hex_color)
        gray = int((r * 0.299) + (g * 0.587) + (b * 0.114))
        gray = max(24, min(232, gray + int(contrast)))
        return _rgb_to_hex((gray, gray, gray))

    def _character_palette_values(self, character: dict[str, Any], panel_theme: dict[str, str]) -> tuple[str, str, str, str, str]:
        character_palette = character.get("palette") or {}
        hair = self._safe_hex(character_palette.get("hair"), "#5b3d24")
        skin = self._safe_hex(character_palette.get("skin"), "#f3c89c")
        outfit = self._safe_hex(character_palette.get("costume_primary"), panel_theme["accent"])
        outfit_alt = self._safe_hex(character_palette.get("costume_accent"), panel_theme["accent_alt"])
        eye = self._safe_hex(character_palette.get("eye"), panel_theme["accent"])
        return hair, skin, outfit, outfit_alt, eye

    def _draw_character_prop(self, draw, cx: int, top_y: int, descriptor: str, theme_slug: str, outfit_alt: str, scale: float) -> None:
        if "astronaut" in descriptor:
            draw.ellipse(
                (cx - int(42 * scale), top_y - int(76 * scale), cx + int(42 * scale), top_y + int(10 * scale)),
                outline="#ffffff",
                width=max(3, int(5 * scale)),
            )
            draw.arc(
                (cx - int(34 * scale), top_y - int(62 * scale), cx + int(34 * scale), top_y),
                start=190,
                end=350,
                fill=outfit_alt,
                width=max(2, int(4 * scale)),
            )
        elif "ninja-hero" in descriptor:
            draw.rounded_rectangle(
                (cx - int(38 * scale), top_y - int(70 * scale), cx + int(38 * scale), top_y - int(56 * scale)),
                radius=int(7 * scale),
                fill=outfit_alt,
                outline="#111111",
                width=2,
            )
        elif "web-hero" in descriptor:
            draw.line(
                [(cx + int(32 * scale), top_y + int(2 * scale)), (cx + int(78 * scale), top_y - int(44 * scale))],
                fill="#ffffff",
                width=max(2, int(4 * scale)),
            )
        elif "tech-armor" in descriptor:
            draw.ellipse(
                (cx - int(12 * scale), top_y + int(28 * scale), cx + int(12 * scale), top_y + int(52 * scale)),
                fill=outfit_alt,
                outline="#ffffff",
                width=2,
            )
        elif "dark-knight" in descriptor:
            draw.polygon(
                [
                    (cx - int(34 * scale), top_y - int(48 * scale)),
                    (cx - int(20 * scale), top_y - int(84 * scale)),
                    (cx - int(6 * scale), top_y - int(48 * scale)),
                ],
                fill="#111111",
                outline="#111111",
            )
            draw.polygon(
                [
                    (cx + int(6 * scale), top_y - int(48 * scale)),
                    (cx + int(20 * scale), top_y - int(84 * scale)),
                    (cx + int(34 * scale), top_y - int(48 * scale)),
                ],
                fill="#111111",
                outline="#111111",
            )
        elif "saiyan-warrior" in descriptor:
            for radius in (46, 62):
                draw.arc(
                    (cx - int(radius * scale), top_y - int(70 * scale), cx + int(radius * scale), top_y + int(90 * scale)),
                    start=205,
                    end=335,
                    fill=outfit_alt,
                    width=max(2, int(3 * scale)),
                )
        elif "pirate" in descriptor:
            draw.ellipse(
                (cx - int(42 * scale), top_y - int(68 * scale), cx + int(42 * scale), top_y - int(44 * scale)),
                fill=outfit_alt,
                outline="#111111",
                width=2,
            )
            draw.rounded_rectangle(
                (cx - int(30 * scale), top_y - int(58 * scale), cx + int(30 * scale), top_y - int(36 * scale)),
                radius=int(8 * scale),
                fill=outfit_alt,
                outline="#111111",
                width=2,
            )
        elif "book" in descriptor or "scroll" in descriptor or theme_slug == "fantasy":
            fill = "#fff3cf" if theme_slug != "manga" else "#f5f5f5"
            draw.rounded_rectangle((cx + int(18 * scale), top_y + int(28 * scale), cx + int(52 * scale), top_y + int(52 * scale)), radius=int(7 * scale), fill=fill, outline="#111111", width=2)
            draw.line([(cx + int(35 * scale), top_y + int(30 * scale)), (cx + int(35 * scale), top_y + int(50 * scale))], fill="#111111", width=1)
        elif "beaker" in descriptor or "flask" in descriptor or "lab" in descriptor:
            draw.rounded_rectangle((cx + int(22 * scale), top_y + int(26 * scale), cx + int(42 * scale), top_y + int(56 * scale)), radius=int(5 * scale), fill="#8be9fd", outline="#111111", width=2)
            draw.ellipse((cx + int(16 * scale), top_y + int(16 * scale), cx + int(48 * scale), top_y + int(36 * scale)), fill="#8be9fd88")
        elif "shield" in descriptor or "hero" in descriptor or theme_slug == "superhero":
            draw.polygon(
                [
                    (cx + int(24 * scale), top_y + int(20 * scale)),
                    (cx + int(44 * scale), top_y + int(20 * scale)),
                    (cx + int(50 * scale), top_y + int(34 * scale)),
                    (cx + int(34 * scale), top_y + int(56 * scale)),
                    (cx + int(18 * scale), top_y + int(34 * scale)),
                ],
                fill=outfit_alt,
                outline="#111111",
            )

    def _draw_anime_character(
        self,
        draw,
        x: int,
        y: int,
        w: int,
        h: int,
        descriptor: str,
        pose_name: str,
        expr: str,
        mouth: str,
        hair: str,
        skin: str,
        outfit: str,
        outfit_alt: str,
        eye: str,
        panel_theme: dict[str, str],
    ) -> None:
        scale = max(0.40, min(w / 190, max(h, 1) / 210))
        cx = x + w // 2
        ground_y = y + h + 2
        head_r = int(32 * scale)
        torso_w = int(46 * scale)
        torso_h = int(74 * scale)
        leg_h = int(64 * scale)
        body_top = ground_y - leg_h - torso_h
        head_cy = body_top - head_r + int(18 * scale)

        draw.ellipse((cx - int(40 * scale), ground_y - 6, cx + int(40 * scale), ground_y + 8), fill="#0000002a")
        draw.rounded_rectangle((cx - 16, ground_y - leg_h, cx - 3, ground_y), radius=8, fill=outfit_alt, outline="#111111", width=2)
        draw.rounded_rectangle((cx + 3, ground_y - leg_h, cx + 16, ground_y), radius=8, fill=outfit_alt, outline="#111111", width=2)
        draw.polygon(
            [
                (cx - torso_w // 2, body_top + 12),
                (cx, body_top),
                (cx + torso_w // 2, body_top + 12),
                (cx + torso_w // 2 - 6, body_top + torso_h),
                (cx - torso_w // 2 + 6, body_top + torso_h),
            ],
            fill=outfit,
            outline="#111111",
        )
        draw.polygon([(cx - 12, body_top + 6), (cx, body_top + 24), (cx + 12, body_top + 6)], fill="#ffffff", outline="#111111")
        draw.polygon([(cx - 4, body_top + 16), (cx, body_top + 30), (cx + 4, body_top + 16)], fill=outfit_alt, outline="#111111")
        draw.line([(cx - torso_w // 2 + 4, body_top + 26), (cx - 38, body_top + 44)], fill=outfit, width=max(4, int(8 * scale)))
        draw.line([(cx + torso_w // 2 - 4, body_top + 24), (cx + 40, body_top + (8 if pose_name in {"pointing", "explaining"} else 42))], fill=outfit, width=max(4, int(8 * scale)))
        draw.ellipse((cx - 48, body_top + 34, cx - 28, body_top + 54), fill=skin, outline="#111111", width=2)
        draw.ellipse((cx + 28, body_top + (0 if pose_name in {"pointing", "explaining"} else 32), cx + 48, body_top + (20 if pose_name in {"pointing", "explaining"} else 52)), fill=skin, outline="#111111", width=2)
        draw.rounded_rectangle((cx - 10, body_top - 2, cx + 10, body_top + 18), radius=5, fill=skin, outline="#111111", width=2)
        draw.ellipse((cx - head_r, head_cy - head_r, cx + head_r, head_cy + head_r), fill=skin, outline="#111111", width=3)
        self._draw_dataset_hair(draw, cx, head_cy, head_r + 2, hair, f"anime {descriptor}", scale)
        for blush_x in (cx - 22, cx + 10):
            draw.line([(blush_x, head_cy + 12), (blush_x + 8, head_cy + 16)], fill="#ff8fab", width=2)
            draw.line([(blush_x, head_cy + 17), (blush_x + 8, head_cy + 21)], fill="#ff8fab", width=2)
        self._draw_comic_eyes(draw, cx, head_cy + 4, max(28, int(head_r * 1.2)), expr, {"accent": eye}, max(1.0, scale * 1.08))
        self._draw_comic_mouth(draw, cx, head_cy + int(22 * scale), mouth, max(0.84, scale * 0.85))
        self._draw_emotion_marks(draw, cx + head_r + 12, head_cy - head_r + 6, ({"excited":"excited","proud":"proud"}.get(expr, "happy")), panel_theme["accent"], 0.9)
        self._draw_character_prop(draw, cx, body_top, descriptor, "anime", outfit_alt, scale)

    def _draw_pixar_character(
        self,
        draw,
        x: int,
        y: int,
        w: int,
        h: int,
        descriptor: str,
        pose_name: str,
        expr: str,
        mouth: str,
        hair: str,
        skin: str,
        outfit: str,
        outfit_alt: str,
        eye: str,
        panel_theme: dict[str, str],
        category: str,
    ) -> None:
        scale = max(0.40, min(w / 180, max(h, 1) / 205))
        cx = x + w // 2
        ground_y = y + h + 2
        if category == "mascot":
            blob_w = int(88 * scale)
            blob_h = int(92 * scale)
            draw.ellipse((cx - int(42 * scale), ground_y - 8, cx + int(42 * scale), ground_y + 10), fill="#00000022")
            draw.rounded_rectangle((cx - blob_w // 2, ground_y - blob_h, cx + blob_w // 2, ground_y), radius=int(28 * scale), fill=outfit, outline="#111111", width=3)
            draw.ellipse((cx - int(28 * scale), ground_y - blob_h + int(18 * scale), cx - int(4 * scale), ground_y - blob_h + int(48 * scale)), fill="#ffffff", outline="#111111", width=2)
            draw.ellipse((cx + int(4 * scale), ground_y - blob_h + int(18 * scale), cx + int(28 * scale), ground_y - blob_h + int(48 * scale)), fill="#ffffff", outline="#111111", width=2)
            draw.ellipse((cx - int(18 * scale), ground_y - blob_h + int(26 * scale), cx - int(8 * scale), ground_y - blob_h + int(40 * scale)), fill=eye)
            draw.ellipse((cx + int(8 * scale), ground_y - blob_h + int(26 * scale), cx + int(18 * scale), ground_y - blob_h + int(40 * scale)), fill=eye)
            self._draw_comic_mouth(draw, cx, ground_y - blob_h + int(58 * scale), mouth, max(0.75, scale * 0.8))
            draw.ellipse((cx - int(48 * scale), ground_y - blob_h + int(38 * scale), cx - int(26 * scale), ground_y - blob_h + int(60 * scale)), fill=outfit_alt, outline="#111111", width=2)
            draw.ellipse((cx + int(26 * scale), ground_y - blob_h + int(38 * scale), cx + int(48 * scale), ground_y - blob_h + int(60 * scale)), fill=outfit_alt, outline="#111111", width=2)
            return

        head_r = int(38 * scale)
        torso_w = int(58 * scale)
        torso_h = int(70 * scale)
        leg_h = int(46 * scale)
        body_top = ground_y - leg_h - torso_h
        head_cy = body_top - head_r + int(24 * scale)
        draw.ellipse((cx - int(46 * scale), ground_y - 8, cx + int(46 * scale), ground_y + 10), fill="#00000026")
        draw.rounded_rectangle((cx - 18, ground_y - leg_h, cx - 2, ground_y), radius=9, fill=outfit_alt, outline="#111111", width=2)
        draw.rounded_rectangle((cx + 2, ground_y - leg_h, cx + 18, ground_y), radius=9, fill=outfit_alt, outline="#111111", width=2)
        draw.rounded_rectangle((cx - torso_w // 2, body_top, cx + torso_w // 2, body_top + torso_h), radius=int(18 * scale), fill=outfit, outline="#111111", width=3)
        draw.ellipse((cx - int(16 * scale), body_top + int(18 * scale), cx + int(16 * scale), body_top + int(48 * scale)), fill="#ffffff99")
        draw.line([(cx - torso_w // 2 + 4, body_top + 26), (cx - 42, body_top + 44)], fill=outfit, width=max(6, int(10 * scale)))
        draw.line([(cx + torso_w // 2 - 4, body_top + 26), (cx + 42, body_top + (10 if pose_name in {"pointing", "explaining"} else 42))], fill=outfit, width=max(6, int(10 * scale)))
        draw.ellipse((cx - 52, body_top + 34, cx - 28, body_top + 58), fill=skin, outline="#111111", width=2)
        draw.ellipse((cx + 28, body_top + (0 if pose_name in {"pointing", "explaining"} else 32), cx + 52, body_top + (24 if pose_name in {"pointing", "explaining"} else 56)), fill=skin, outline="#111111", width=2)
        draw.ellipse((cx - head_r, head_cy - head_r, cx + head_r, head_cy + head_r), fill=skin, outline="#111111", width=3)
        self._draw_dataset_hair(draw, cx, head_cy, head_r + 1, hair, f"pixar {descriptor}", scale)
        draw.ellipse((cx - int(30 * scale), head_cy + int(10 * scale), cx - int(18 * scale), head_cy + int(20 * scale)), fill="#ffb5a7")
        draw.ellipse((cx + int(18 * scale), head_cy + int(10 * scale), cx + int(30 * scale), head_cy + int(20 * scale)), fill="#ffb5a7")
        self._draw_comic_eyes(draw, cx, head_cy + 6, max(30, int(head_r * 1.2)), expr, {"accent": eye}, max(1.04, scale * 1.1))
        draw.ellipse((cx - int(4 * scale), head_cy + int(14 * scale), cx + int(4 * scale), head_cy + int(20 * scale)), fill="#d2876c")
        self._draw_comic_mouth(draw, cx, head_cy + int(26 * scale), mouth, max(0.88, scale * 0.9))
        self._draw_character_prop(draw, cx, body_top, descriptor, "pixar", outfit_alt, scale)

    def _draw_superhero_dataset_character(
        self,
        draw,
        x: int,
        y: int,
        w: int,
        h: int,
        descriptor: str,
        expr: str,
        mouth: str,
        hair: str,
        skin: str,
        outfit: str,
        outfit_alt: str,
        eye: str,
        panel_theme: dict[str, str],
    ) -> None:
        scale = max(0.46, min(w / 190, max(h, 1) / 230))
        cx = x + w // 2
        ground_y = y + h + 2
        head_r = int(34 * scale)
        shoulder_w = int(82 * scale)
        waist_w = int(54 * scale)
        torso_h = int(86 * scale)
        leg_h = int(64 * scale)
        body_top = ground_y - leg_h - torso_h
        head_cy = body_top - head_r + int(22 * scale)
        draw.ellipse((cx - int(48 * scale), ground_y - 6, cx + int(48 * scale), ground_y + 10), fill="#0000002f")
        cape = [(cx - int(26 * scale), body_top + int(12 * scale)), (cx - int(62 * scale), ground_y - int(28 * scale)), (cx + int(2 * scale), ground_y - int(14 * scale)), (cx + int(18 * scale), body_top + int(34 * scale))]
        draw.polygon(cape, fill=outfit_alt, outline="#111111")
        draw.rounded_rectangle((cx - 20, ground_y - leg_h, cx - 4, ground_y), radius=8, fill=outfit_alt, outline="#111111", width=2)
        draw.rounded_rectangle((cx + 4, ground_y - leg_h, cx + 20, ground_y), radius=8, fill=outfit_alt, outline="#111111", width=2)
        draw.polygon(
            [
                (cx - shoulder_w // 2, body_top + 18),
                (cx + shoulder_w // 2, body_top + 18),
                (cx + waist_w // 2, body_top + torso_h),
                (cx - waist_w // 2, body_top + torso_h),
            ],
            fill=outfit,
            outline="#111111",
        )
        draw.ellipse((cx - int(18 * scale), body_top + int(24 * scale), cx + int(18 * scale), body_top + int(58 * scale)), fill="#ffffff", outline="#111111", width=2)
        draw.arc((cx - int(10 * scale), body_top + int(30 * scale), cx + int(10 * scale), body_top + int(50 * scale)), start=40, end=320, fill=panel_theme["accent"], width=4)
        draw.rectangle((cx - waist_w // 2, body_top + torso_h - 18, cx + waist_w // 2, body_top + torso_h - 6), fill="#111111")
        draw.line([(cx - shoulder_w // 2 + 6, body_top + 28), (cx - 56, body_top + 10)], fill=outfit, width=max(6, int(10 * scale)))
        draw.line([(cx + shoulder_w // 2 - 6, body_top + 28), (cx + 54, body_top - 18)], fill=outfit, width=max(6, int(10 * scale)))
        draw.ellipse((cx - 66, body_top, cx - 42, body_top + 24), fill=skin, outline="#111111", width=2)
        draw.ellipse((cx + 42, body_top - 28, cx + 68, body_top - 2), fill=skin, outline="#111111", width=2)
        draw.ellipse((cx - head_r, head_cy - head_r, cx + head_r, head_cy + head_r), fill=skin, outline="#111111", width=3)
        self._draw_dataset_hair(draw, cx, head_cy, head_r + 1, hair, f"hero {descriptor}", scale)
        draw.rounded_rectangle((cx - int(28 * scale), head_cy - int(12 * scale), cx + int(28 * scale), head_cy + int(10 * scale)), radius=int(8 * scale), outline=outfit_alt, width=max(2, int(3 * scale)))
        self._draw_comic_eyes(draw, cx, head_cy + 1, max(26, int(head_r * 0.95)), expr, {"accent": eye}, max(0.86, scale * 0.9))
        self._draw_comic_mouth(draw, cx, head_cy + int(22 * scale), mouth, max(0.82, scale * 0.84))
        self._draw_emotion_marks(draw, cx + head_r + 14, head_cy - head_r + 8, "excited" if expr == "excited" else "proud", panel_theme["accent"], 0.95)
        self._draw_character_prop(draw, cx, body_top, descriptor, "superhero", outfit_alt, scale)

    def _draw_manga_character(
        self,
        draw,
        x: int,
        y: int,
        w: int,
        h: int,
        descriptor: str,
        pose_name: str,
        expr: str,
        mouth: str,
        hair: str,
        skin: str,
        outfit: str,
        outfit_alt: str,
        eye: str,
        panel_theme: dict[str, str],
    ) -> None:
        hair = self._mono_hex(hair, -26)
        skin = self._mono_hex(skin, 32)
        outfit = self._mono_hex(outfit, -36)
        outfit_alt = self._mono_hex(outfit_alt, 16)
        eye = self._mono_hex(eye, -60)
        scale = max(0.40, min(w / 188, max(h, 1) / 210))
        cx = x + w // 2
        ground_y = y + h + 2
        head_r = int(31 * scale)
        torso_w = int(44 * scale)
        torso_h = int(76 * scale)
        leg_h = int(68 * scale)
        body_top = ground_y - leg_h - torso_h
        head_cy = body_top - head_r + int(18 * scale)
        draw.ellipse((cx - int(38 * scale), ground_y - 4, cx + int(38 * scale), ground_y + 8), fill="#00000022")
        for line_x in range(x + 4, x + w - 4, 18):
            draw.line([(line_x, y + 12), (line_x - 32, ground_y - 10)], fill="#11111112", width=1)
        draw.rounded_rectangle((cx - 14, ground_y - leg_h, cx - 2, ground_y), radius=7, fill=outfit_alt, outline="#111111", width=2)
        draw.rounded_rectangle((cx + 2, ground_y - leg_h, cx + 14, ground_y), radius=7, fill=outfit_alt, outline="#111111", width=2)
        draw.polygon([(cx - torso_w // 2, body_top + 12), (cx + torso_w // 2, body_top + 12), (cx + torso_w // 2 - 6, body_top + torso_h), (cx - torso_w // 2 + 6, body_top + torso_h)], fill=outfit, outline="#111111")
        draw.line([(cx - torso_w // 2 + 4, body_top + 24), (cx - 38, body_top + 44)], fill=outfit, width=max(4, int(8 * scale)))
        draw.line([(cx + torso_w // 2 - 4, body_top + 24), (cx + 42, body_top + (10 if pose_name in {"pointing", "explaining"} else 42))], fill=outfit, width=max(4, int(8 * scale)))
        draw.ellipse((cx - head_r, head_cy - head_r, cx + head_r, head_cy + head_r), fill=skin, outline="#111111", width=3)
        self._draw_dataset_hair(draw, cx, head_cy, head_r + 2, hair, f"manga {descriptor} spiky", scale)
        draw.line([(cx - head_r + 8, head_cy - head_r + 6), (cx + head_r - 8, head_cy - head_r + 2)], fill="#111111", width=2)
        self._draw_comic_eyes(draw, cx, head_cy + 3, max(26, int(head_r * 1.02)), "think" if expr == "think" else expr, {"accent": eye}, max(0.9, scale * 0.96))
        self._draw_comic_mouth(draw, cx, head_cy + int(22 * scale), mouth, max(0.8, scale * 0.82))
        self._draw_character_prop(draw, cx, body_top, descriptor, "manga", outfit_alt, scale)

    def _draw_kids_character(
        self,
        draw,
        x: int,
        y: int,
        w: int,
        h: int,
        descriptor: str,
        pose_name: str,
        expr: str,
        mouth: str,
        hair: str,
        skin: str,
        outfit: str,
        outfit_alt: str,
        eye: str,
        panel_theme: dict[str, str],
    ) -> None:
        scale = max(0.42, min(w / 170, max(h, 1) / 190))
        cx = x + w // 2
        ground_y = y + h + 2
        head_r = int(40 * scale)
        torso_w = int(42 * scale)
        torso_h = int(54 * scale)
        leg_h = int(42 * scale)
        body_top = ground_y - leg_h - torso_h
        head_cy = body_top - head_r + int(28 * scale)
        draw.ellipse((cx - int(44 * scale), ground_y - 8, cx + int(44 * scale), ground_y + 10), fill="#00000022")
        draw.rounded_rectangle((cx - 16, ground_y - leg_h, cx - 2, ground_y), radius=8, fill=outfit_alt, outline="#111111", width=2)
        draw.rounded_rectangle((cx + 2, ground_y - leg_h, cx + 16, ground_y), radius=8, fill=outfit_alt, outline="#111111", width=2)
        draw.rounded_rectangle((cx - torso_w // 2, body_top, cx + torso_w // 2, body_top + torso_h), radius=int(14 * scale), fill=outfit, outline="#111111", width=3)
        draw.rectangle((cx - int(10 * scale), body_top + int(10 * scale), cx + int(10 * scale), body_top + torso_h - int(8 * scale)), fill="#ffffff", outline="#111111", width=2)
        draw.line([(cx - torso_w // 2 + 4, body_top + 18), (cx - 34, body_top + 34)], fill=outfit, width=max(5, int(8 * scale)))
        draw.line([(cx + torso_w // 2 - 4, body_top + 18), (cx + 36, body_top + (8 if pose_name in {"pointing", "explaining"} else 34))], fill=outfit, width=max(5, int(8 * scale)))
        draw.ellipse((cx - head_r, head_cy - head_r, cx + head_r, head_cy + head_r), fill=skin, outline="#111111", width=3)
        self._draw_dataset_hair(draw, cx, head_cy, head_r + 1, hair, f"kids {descriptor}", scale)
        for freckle_x in (-14, -8, 8, 14):
            draw.ellipse((cx + freckle_x - 1, head_cy + 16, cx + freckle_x + 1, head_cy + 18), fill="#d08b6f")
        self._draw_comic_eyes(draw, cx, head_cy + 6, max(30, int(head_r * 1.12)), expr, {"accent": eye}, max(1.05, scale * 1.12))
        self._draw_comic_mouth(draw, cx, head_cy + int(24 * scale), mouth, max(0.86, scale * 0.88))
        self._draw_character_prop(draw, cx, body_top, descriptor, "cartoon_kids", outfit_alt, scale)

    def _draw_fantasy_character(
        self,
        draw,
        x: int,
        y: int,
        w: int,
        h: int,
        descriptor: str,
        pose_name: str,
        expr: str,
        mouth: str,
        hair: str,
        skin: str,
        outfit: str,
        outfit_alt: str,
        eye: str,
        panel_theme: dict[str, str],
    ) -> None:
        scale = max(0.40, min(w / 186, max(h, 1) / 205))
        cx = x + w // 2
        ground_y = y + h + 2
        head_r = int(33 * scale)
        robe_w = int(74 * scale)
        robe_h = int(96 * scale)
        body_top = ground_y - robe_h
        head_cy = body_top - head_r + int(18 * scale)
        draw.ellipse((cx - int(42 * scale), ground_y - 6, cx + int(42 * scale), ground_y + 10), fill="#00000029")
        cloak = [(cx - int(28 * scale), body_top + int(10 * scale)), (cx - int(62 * scale), ground_y), (cx + int(2 * scale), ground_y - int(12 * scale)), (cx + int(22 * scale), body_top + int(28 * scale))]
        draw.polygon(cloak, fill=outfit_alt, outline="#111111")
        draw.polygon([(cx - robe_w // 2, body_top + 12), (cx + robe_w // 2, body_top + 12), (cx + int(26 * scale), ground_y), (cx - int(26 * scale), ground_y)], fill=outfit, outline="#111111")
        draw.polygon([(cx - 14, body_top + 12), (cx, body_top + 36), (cx + 14, body_top + 12)], fill="#fff4d6", outline="#111111")
        draw.line([(cx - int(18 * scale), body_top + int(12 * scale)), (cx - int(10 * scale), ground_y - int(8 * scale))], fill=outfit_alt, width=3)
        draw.line([(cx + int(18 * scale), body_top + int(12 * scale)), (cx + int(10 * scale), ground_y - int(8 * scale))], fill=outfit_alt, width=3)
        draw.line([(cx - 20, body_top + 28), (cx - 48, body_top + 48)], fill=outfit, width=max(5, int(8 * scale)))
        draw.line([(cx + 20, body_top + 24), (cx + 44, body_top + (0 if pose_name == "spell_cast" else 42))], fill=outfit, width=max(5, int(8 * scale)))
        draw.ellipse((cx - 58, body_top + 38, cx - 34, body_top + 60), fill=skin, outline="#111111", width=2)
        draw.ellipse((cx + 34, body_top + (0 if pose_name == "spell_cast" else 32), cx + 58, body_top + (22 if pose_name == "spell_cast" else 54)), fill=skin, outline="#111111", width=2)
        draw.ellipse((cx - head_r, head_cy - head_r, cx + head_r, head_cy + head_r), fill=skin, outline="#111111", width=3)
        self._draw_dataset_hair(draw, cx, head_cy, head_r + 1, hair, f"fantasy {descriptor}", scale)
        self._draw_comic_eyes(draw, cx, head_cy + 3, max(26, int(head_r * 1.0)), expr, {"accent": eye}, max(0.9, scale * 0.95))
        self._draw_comic_mouth(draw, cx, head_cy + int(22 * scale), mouth, max(0.82, scale * 0.84))
        draw.line([(cx + 46, body_top + 4), (cx + 58, body_top - 36)], fill="#111111", width=3)
        self._draw_starburst(draw, cx + 60, body_top - 42, int(8 * scale), "#ffffff", panel_theme["accent"], spikes=6)
        self._draw_character_prop(draw, cx, body_top, descriptor, "fantasy", outfit_alt, scale)

    def _draw_dataset_character(
        self,
        draw,
        x: int,
        y: int,
        w: int,
        h: int,
        visual: dict[str, Any],
        style_record: dict[str, Any],
        panel_theme: dict[str, str],
    ) -> None:
        character = visual.get("character") or {}
        descriptor = self._record_blob(character)
        theme_profile = visual.get("theme_profile") or {}
        theme_slug = str(theme_profile.get("slug") or "").lower()
        category = str(character.get("category") or "").lower()
        variant = str(character.get("render_variant") or character.get("sub_type") or "").lower()
        is_archive = character.get("source") == "archive"
        if category == "mascot" or variant in {"electric-mascot", "lion", "alien", "mythic-monkey", "ai-spirit"}:
            self._draw_theme_mascot_character(draw, x, y, w, h, visual, panel_theme)
            return
        if "robot" in descriptor:
            self._draw_robot_character(draw, x, y, w, h, visual, panel_theme)
            return

        emotion_name = (visual.get("emotion") or {}).get("name", "happy")
        pose_name = (visual.get("pose") or {}).get("name", "standing")
        expr, mouth = self._emotion_face(emotion_name)
        hair, skin, outfit, outfit_alt, eye = self._character_palette_values(character, panel_theme)
        if theme_slug == "superhero" or is_archive:
            self._draw_superhero_dataset_character(draw, x, y, w, h, descriptor, expr, mouth, hair, skin, outfit, outfit_alt, eye, panel_theme)
        elif theme_slug == "anime":
            self._draw_anime_character(draw, x, y, w, h, descriptor, pose_name, expr, mouth, hair, skin, outfit, outfit_alt, eye, panel_theme)
        elif theme_slug == "pixar":
            self._draw_pixar_character(draw, x, y, w, h, descriptor, pose_name, expr, mouth, hair, skin, outfit, outfit_alt, eye, panel_theme, category)
        elif theme_slug == "manga":
            self._draw_manga_character(draw, x, y, w, h, descriptor, pose_name, expr, mouth, hair, skin, outfit, outfit_alt, eye, panel_theme)
        elif theme_slug == "fantasy":
            self._draw_fantasy_character(draw, x, y, w, h, descriptor, pose_name, expr, mouth, hair, skin, outfit, outfit_alt, eye, panel_theme)
        elif theme_slug == "horror":
            self._draw_manga_character(draw, x, y, w, h, f"horror mystery {descriptor}", pose_name, expr, mouth, hair, skin, outfit, outfit_alt, eye, panel_theme)
        else:
            self._draw_kids_character(draw, x, y, w, h, descriptor, pose_name, expr, mouth, hair, skin, outfit, outfit_alt, eye, panel_theme)

    def _draw_theme_mascot_character(self, draw, x: int, y: int, w: int, h: int, visual: dict[str, Any], panel_theme: dict[str, str]) -> None:
        character = visual.get("character") or {}
        variant = str(character.get("render_variant") or character.get("sub_type") or "").lower()
        hair, skin, outfit, outfit_alt, eye = self._character_palette_values(character, panel_theme)
        emotion_name = (visual.get("emotion") or {}).get("name", "happy")
        expr, mouth = self._emotion_face(emotion_name)
        scale = max(0.40, min(w / 180, max(h, 1) / 205))
        cx = x + w // 2
        ground_y = y + h + 2
        draw.ellipse((cx - int(48 * scale), ground_y - 8, cx + int(48 * scale), ground_y + 10), fill="#0000002a")

        if variant == "electric-mascot":
            tail = [
                (cx + int(44 * scale), ground_y - int(92 * scale)),
                (cx + int(82 * scale), ground_y - int(132 * scale)),
                (cx + int(66 * scale), ground_y - int(132 * scale)),
                (cx + int(98 * scale), ground_y - int(176 * scale)),
                (cx + int(58 * scale), ground_y - int(142 * scale)),
                (cx + int(72 * scale), ground_y - int(142 * scale)),
                (cx + int(36 * scale), ground_y - int(106 * scale)),
            ]
            draw.polygon(tail, fill=outfit_alt, outline="#111111")
            draw.ellipse((cx - int(40 * scale), ground_y - int(106 * scale), cx + int(40 * scale), ground_y - int(20 * scale)), fill=outfit, outline="#111111", width=3)
            draw.ellipse((cx - int(46 * scale), ground_y - int(174 * scale), cx + int(46 * scale), ground_y - int(84 * scale)), fill=outfit, outline="#111111", width=3)
            draw.polygon([(cx - int(32 * scale), ground_y - int(160 * scale)), (cx - int(48 * scale), ground_y - int(224 * scale)), (cx - int(10 * scale), ground_y - int(176 * scale))], fill=outfit, outline="#111111")
            draw.polygon([(cx + int(32 * scale), ground_y - int(160 * scale)), (cx + int(48 * scale), ground_y - int(224 * scale)), (cx + int(10 * scale), ground_y - int(176 * scale))], fill=outfit, outline="#111111")
            draw.ellipse((cx - int(30 * scale), ground_y - int(130 * scale), cx - int(14 * scale), ground_y - int(112 * scale)), fill=eye)
            draw.ellipse((cx + int(14 * scale), ground_y - int(130 * scale), cx + int(30 * scale), ground_y - int(112 * scale)), fill=eye)
            draw.ellipse((cx - int(42 * scale), ground_y - int(108 * scale), cx - int(22 * scale), ground_y - int(88 * scale)), fill="#ef4444", outline="#111111", width=2)
            draw.ellipse((cx + int(22 * scale), ground_y - int(108 * scale), cx + int(42 * scale), ground_y - int(88 * scale)), fill="#ef4444", outline="#111111", width=2)
            self._draw_comic_mouth(draw, cx, ground_y - int(102 * scale), mouth, max(0.7, scale * 0.75))
            return

        if variant == "lion":
            draw.ellipse((cx - int(60 * scale), ground_y - int(172 * scale), cx + int(60 * scale), ground_y - int(52 * scale)), fill=hair, outline="#111111", width=3)
            draw.ellipse((cx - int(42 * scale), ground_y - int(154 * scale), cx + int(42 * scale), ground_y - int(70 * scale)), fill=skin, outline="#111111", width=3)
            draw.ellipse((cx - int(44 * scale), ground_y - int(84 * scale), cx + int(44 * scale), ground_y), fill=outfit, outline="#111111", width=3)
            draw.ellipse((cx - int(24 * scale), ground_y - int(126 * scale), cx - int(10 * scale), ground_y - int(110 * scale)), fill=eye)
            draw.ellipse((cx + int(10 * scale), ground_y - int(126 * scale), cx + int(24 * scale), ground_y - int(110 * scale)), fill=eye)
            draw.polygon([(cx, ground_y - int(104 * scale)), (cx - int(8 * scale), ground_y - int(92 * scale)), (cx + int(8 * scale), ground_y - int(92 * scale))], fill="#111111")
            self._draw_comic_mouth(draw, cx, ground_y - int(86 * scale), mouth, max(0.7, scale * 0.72))
            return

        if variant == "alien":
            draw.line((cx - int(26 * scale), ground_y - int(160 * scale), cx - int(42 * scale), ground_y - int(202 * scale)), fill="#111111", width=3)
            draw.line((cx + int(26 * scale), ground_y - int(160 * scale), cx + int(42 * scale), ground_y - int(202 * scale)), fill="#111111", width=3)
            draw.ellipse((cx - int(50 * scale), ground_y - int(164 * scale), cx + int(50 * scale), ground_y - int(64 * scale)), fill=skin, outline="#111111", width=3)
            draw.rounded_rectangle((cx - int(42 * scale), ground_y - int(80 * scale), cx + int(42 * scale), ground_y), radius=int(28 * scale), fill=outfit, outline="#111111", width=3)
            draw.ellipse((cx - int(30 * scale), ground_y - int(128 * scale), cx - int(8 * scale), ground_y - int(104 * scale)), fill="#ffffff", outline="#111111", width=2)
            draw.ellipse((cx + int(8 * scale), ground_y - int(128 * scale), cx + int(30 * scale), ground_y - int(104 * scale)), fill="#ffffff", outline="#111111", width=2)
            draw.ellipse((cx - int(22 * scale), ground_y - int(120 * scale), cx - int(14 * scale), ground_y - int(112 * scale)), fill=eye)
            draw.ellipse((cx + int(14 * scale), ground_y - int(120 * scale), cx + int(22 * scale), ground_y - int(112 * scale)), fill=eye)
            self._draw_comic_mouth(draw, cx, ground_y - int(92 * scale), mouth, max(0.66, scale * 0.68))
            return

        if variant == "ai-spirit":
            draw.ellipse((cx - int(54 * scale), ground_y - int(152 * scale), cx + int(54 * scale), ground_y - int(44 * scale)), fill=f"{outfit}88", outline=outfit_alt, width=3)
            draw.ellipse((cx - int(32 * scale), ground_y - int(126 * scale), cx - int(14 * scale), ground_y - int(108 * scale)), fill="#ffffff")
            draw.ellipse((cx + int(14 * scale), ground_y - int(126 * scale), cx + int(32 * scale), ground_y - int(108 * scale)), fill="#ffffff")
            draw.ellipse((cx - int(25 * scale), ground_y - int(120 * scale), cx - int(20 * scale), ground_y - int(115 * scale)), fill=eye)
            draw.ellipse((cx + int(20 * scale), ground_y - int(120 * scale), cx + int(25 * scale), ground_y - int(115 * scale)), fill=eye)
            self._draw_comic_mouth(draw, cx, ground_y - int(96 * scale), mouth, max(0.7, scale * 0.72))
            return

        draw.rounded_rectangle((cx - int(44 * scale), ground_y - int(116 * scale), cx + int(44 * scale), ground_y - int(16 * scale)), radius=int(30 * scale), fill=outfit, outline="#111111", width=3)
        draw.ellipse((cx - int(42 * scale), ground_y - int(158 * scale), cx + int(42 * scale), ground_y - int(76 * scale)), fill=skin, outline="#111111", width=3)
        self._draw_comic_eyes(draw, cx, ground_y - int(118 * scale), max(24, int(34 * scale)), expr, {"accent": eye}, max(0.8, scale * 0.82))
        self._draw_comic_mouth(draw, cx, ground_y - int(94 * scale), mouth, max(0.72, scale * 0.74))

    def _draw_dataset_hair(self, draw, cx: int, cy: int, radius: int, hair_color: str, descriptor: str, scale: float) -> None:
        draw.pieslice((cx - radius - 2, cy - radius - 14, cx + radius + 2, cy + radius - 6), 180, 360, fill=hair_color, outline="#111111")
        if "pigtail" in descriptor or "twin" in descriptor:
            draw.ellipse((cx - radius - 22, cy - 10, cx - radius + 8, cy + 26), fill=hair_color, outline="#111111")
            draw.ellipse((cx + radius - 8, cy - 10, cx + radius + 22, cy + 26), fill=hair_color, outline="#111111")
        elif "braid" in descriptor or "ponytail" in descriptor:
            draw.ellipse((cx + radius - 4, cy + 2, cx + radius + 26, cy + 38), fill=hair_color, outline="#111111")
        elif "scientist" in descriptor or "white hair" in descriptor or "hero" in descriptor:
            points = [
                (cx - radius, cy - 8),
                (cx - radius - 14, cy - radius + 4),
                (cx - 10, cy - radius - 12),
                (cx, cy - radius - 20),
                (cx + 12, cy - radius - 12),
                (cx + radius + 14, cy - radius + 4),
                (cx + radius, cy - 8),
            ]
            draw.polygon(points, fill=hair_color, outline="#111111")
        else:
            draw.rounded_rectangle((cx - radius, cy - radius - 10, cx + radius, cy - 2), radius=10, fill=hair_color, outline="#111111")

    def _draw_robot_character(self, draw, x: int, y: int, w: int, h: int, visual: dict[str, Any], panel_theme: dict[str, str]) -> None:
        character = visual.get("character") or {}
        palette = character.get("palette") or {}
        body = self._safe_hex(palette.get("costume_primary"), "#4da3ff")
        accent = self._safe_hex(palette.get("costume_accent"), "#ffd54f")
        screen = self._safe_hex(palette.get("eye"), "#a8f0ff")
        cx = x + w // 2
        ground_y = y + h
        draw.ellipse((cx - 44, ground_y - 8, cx + 44, ground_y + 10), fill="#00000035")
        draw.rounded_rectangle((cx - 46, ground_y - 120, cx + 46, ground_y - 26), radius=18, fill=body, outline="#111111", width=3)
        draw.rounded_rectangle((cx - 28, ground_y - 104, cx + 28, ground_y - 62), radius=10, fill="#e7f7ff", outline="#111111", width=2)
        expr, mouth = self._emotion_face((visual.get("emotion") or {}).get("name", "happy"))
        self._draw_comic_eyes(draw, cx, ground_y - 86, 20, expr, {"accent": screen}, 0.7)
        self._draw_comic_mouth(draw, cx, ground_y - 62, mouth, 0.6)
        draw.line([(cx, ground_y - 120), (cx, ground_y - 144)], fill="#111111", width=3)
        draw.ellipse((cx - 8, ground_y - 154, cx + 8, ground_y - 138), fill=accent, outline="#111111", width=2)
        draw.line([(cx - 46, ground_y - 86), (cx - 74, ground_y - 64)], fill=body, width=9)
        draw.line([(cx + 46, ground_y - 86), (cx + 74, ground_y - 64)], fill=body, width=9)
        draw.ellipse((cx - 84, ground_y - 76, cx - 62, ground_y - 54), fill=accent, outline="#111111", width=2)
        draw.ellipse((cx + 62, ground_y - 76, cx + 84, ground_y - 54), fill=accent, outline="#111111", width=2)
        draw.rectangle((cx - 32, ground_y - 24, cx + 32, ground_y), fill=accent, outline="#111111", width=2)
        draw.rectangle((cx - 18, ground_y - 18, cx - 2, ground_y + 10), fill="#4d4d4d", outline="#111111", width=2)
        draw.rectangle((cx + 2, ground_y - 18, cx + 18, ground_y + 10), fill="#4d4d4d", outline="#111111", width=2)
        self._draw_emotion_marks(draw, cx + 52, ground_y - 122, (visual.get("emotion") or {}).get("name", "happy"), accent, 0.9)

    def _draw_emotion_marks(self, draw, x: int, y: int, emotion_name: str, accent: str, scale: float) -> None:
        if emotion_name in {"happy", "excited", "proud"}:
            self._draw_starburst(draw, x, y, int(8 * scale), "#ffffff", accent, spikes=6)
            self._draw_starburst(draw, x + int(20 * scale), y + int(12 * scale), int(6 * scale), "#ffffff", accent, spikes=6)
        elif emotion_name in {"surprised", "shocked"}:
            draw.text((x, y), "!", font=self._font(max(14, int(18 * scale)), bold=True), fill=accent)
            draw.text((x + int(14 * scale), y - int(8 * scale)), "!", font=self._font(max(14, int(18 * scale)), bold=True), fill="#111111")
        elif emotion_name == "thinking":
            draw.text((x, y), "...", font=self._font(max(12, int(14 * scale)), bold=True), fill=accent)

    # ─── SUPERHERO COMIC RENDERER ─────────────────────────────────────────────

    def _render(
        self,
        *,
        image_path: Path,
        title: str,
        subject: str,
        style: str,
        scenes: list[Scene],
        visual_plan: dict[str, Any],
    ) -> list[Path]:
        if Image is None or ImageDraw is None or ImageFont is None:
            raise DependencyError("Comic rendering needs Pillow installed in the Python environment.")

        comic_id = image_path.stem
        page_dir = self.output_dir / f"{comic_id}_pages"
        page_dir.mkdir(parents=True, exist_ok=True)
        page_indices = self._paginate_story_indices(len(scenes))
        scene_pages = [[scenes[index] for index in page] for page in page_indices] or [scenes]
        panel_pages = [[visual_plan["panels"][index] for index in page] for page in page_indices] or [visual_plan["panels"]]
        page_paths: list[Path] = []
        page_count = len(scene_pages)

        for page_index, page_scenes in enumerate(scene_pages):
            page_path = page_dir / f"{comic_id}-page-{page_index + 1:03d}.png"
            self._render_page(
                image_path=page_path,
                title=title,
                subject=subject,
                style=style,
                scenes=page_scenes,
                panel_visuals=panel_pages[page_index],
                visual_plan=visual_plan,
                page_number=page_index + 1,
                page_count=page_count,
                panel_start_index=page_indices[page_index][0] if page_indices and page_indices[page_index] else 0,
            )
            page_paths.append(page_path)

        if page_paths:
            with Image.open(page_paths[0]) as preview:
                preview.save(image_path, format="PNG", dpi=(300, 300))
        return page_paths

    def _render_image_comic(
        self,
        *,
        image_path: Path,
        title: str,
        subject: str,
        scenes: list[Scene],
        visual_plan: dict[str, Any],
        source_images: list[Any],
        analyses: list[ImageComicAnalysis],
    ) -> list[Path]:
        if Image is None or ImageDraw is None or ImageFont is None:
            raise DependencyError("Comic rendering needs Pillow installed in the Python environment.")

        comic_id = image_path.stem
        page_dir = self.output_dir / f"{comic_id}_pages"
        page_dir.mkdir(parents=True, exist_ok=True)
        page_indices = self._paginate_story_indices(len(scenes))
        page_paths: list[Path] = []
        page_count = len(page_indices)

        for page_index, page in enumerate(page_indices):
            page_path = page_dir / f"{comic_id}-page-{page_index + 1:03d}.png"
            self._render_image_page(
                image_path=page_path,
                title=title,
                subject=subject,
                scenes=[scenes[index] for index in page],
                panel_visuals=[visual_plan["panels"][index] for index in page],
                visual_plan=visual_plan,
                source_images=source_images,
                analyses=analyses,
                page_number=page_index + 1,
                page_count=page_count,
                panel_start_index=page[0] if page else 0,
            )
            page_paths.append(page_path)

        if page_paths:
            with Image.open(page_paths[0]) as preview:
                preview.save(image_path, format="PNG", dpi=(300, 300))
        return page_paths

    def _render_image_page(
        self,
        *,
        image_path: Path,
        title: str,
        subject: str,
        scenes: list[Scene],
        panel_visuals: list[dict[str, Any]],
        visual_plan: dict[str, Any],
        source_images: list[Any],
        analyses: list[ImageComicAnalysis],
        page_number: int,
        page_count: int,
        panel_start_index: int,
    ) -> None:
        palette = STYLE_LIBRARY[visual_plan["palette_style"]]
        style_record = visual_plan["style"]
        layout_style = str(style_record.get("layout_style") or visual_plan["theme_profile"].get("layout_style") or "balanced_grid")
        margin = 36
        gutter = 24
        layout = self._layout_boxes(len(scenes), margin, gutter, layout_style, page_number=page_number)
        max_x = max(lx + lw for lx, ly, lw, lh, _ in layout)
        max_y = max(ly + lh for lx, ly, lw, lh, _ in layout)
        show_title = page_number == 1
        title_h = 132 if show_title else 0
        footer_h = 34
        width = max_x + margin
        header_space = title_h + gutter if show_title else 0
        height = header_space + max_y + margin * 2 + footer_h

        img = Image.new("RGB", (width, height), _hex_to_rgb(palette["page"]))
        draw = ImageDraw.Draw(img)
        self._gradient_rect(
            img,
            0,
            0,
            width,
            height,
            self._blend_hex(palette["page"], "#ffffff", 0.14),
            self._blend_hex(palette["panel_bot"], "#ffffff", 0.32),
        )
        self._draw_page_halftone(draw, 0, 0, width, height, palette["accent"], density=26, dot_r=1.3, opacity=0.018)
        draw.rounded_rectangle((10, 10, width - 10, height - 10), radius=18, outline="#111111", width=4)
        draw.rounded_rectangle((18, 18, width - 18, height - 18), radius=12, outline="#ffffffbb", width=1)

        if show_title:
            self._draw_dataset_title_banner(
                img,
                draw,
                margin,
                margin,
                width - margin * 2,
                title_h,
                title,
                "image comic",
                palette,
                style_record,
                visual_plan.get("cast_sources", []),
            )

        panel_y_offset = margin + header_space
        for lx, ly, lw, lh, scene_idx in layout:
            global_index = panel_start_index + scene_idx
            source_index = global_index % max(len(source_images), 1)
            self._draw_image_comic_panel(
                img,
                draw,
                lx,
                ly + panel_y_offset,
                lw,
                lh,
                scenes[scene_idx],
                panel_visuals[scene_idx],
                source_images[source_index],
                analyses[source_index],
                palette,
                style_record,
                global_index,
            )

        img.save(image_path, format="PNG", dpi=(300, 300))

    def _render_page(
        self,
        *,
        image_path: Path,
        title: str,
        subject: str,
        style: str,
        scenes: list[Scene],
        panel_visuals: list[dict[str, Any]],
        visual_plan: dict[str, Any],
        page_number: int,
        page_count: int,
        panel_start_index: int,
    ) -> None:
        if Image is None or ImageDraw is None or ImageFont is None:
            raise DependencyError("Comic rendering needs Pillow installed in the Python environment.")

        palette = STYLE_LIBRARY[visual_plan["palette_style"]]
        style_record = visual_plan["style"]
        layout_style = str(style_record.get("layout_style") or visual_plan["theme_profile"].get("layout_style") or "balanced_grid")
        margin = 36  # Consistent outer margin
        gutter = 24  # Increased gutter for better breathing room
        layout = self._layout_boxes(len(scenes), margin, gutter, layout_style, page_number=page_number)
        max_x = max(lx + lw for lx, ly, lw, lh, _ in layout)
        max_y = max(ly + lh for lx, ly, lw, lh, _ in layout)
        show_title = page_number == 1
        title_h = 132 if show_title else 0
        footer_h = 34
        width = max_x + margin
        header_space = title_h + gutter if show_title else 0
        height = header_space + max_y + margin * 2 + footer_h

        img = Image.new("RGB", (width, height), _hex_to_rgb(palette["page"]))
        draw = ImageDraw.Draw(img)
        page_top = self._blend_hex(palette["page"], "#ffffff", 0.18)
        page_bottom = self._blend_hex(palette["panel_bot"], "#ffffff", 0.36)
        self._gradient_rect(img, 0, 0, width, height, page_top, page_bottom)
        self._draw_page_halftone(draw, 0, 0, width, height, palette["accent"], density=26, dot_r=1.3, opacity=0.018)
        draw.rounded_rectangle((10, 10, width - 10, height - 10), radius=18, outline="#111111", width=4)
        draw.rounded_rectangle((18, 18, width - 18, height - 18), radius=12, outline="#ffffffbb", width=1)
        if show_title:
            self._draw_dataset_title_banner(
                img,
                draw,
                margin,
                margin,
                width - margin * 2,
                title_h,
                title,
                subject,
                palette,
                style_record,
                visual_plan.get("cast_sources", []),
            )

        panel_y_offset = margin + header_space
        for lx, ly, lw, lh, scene_idx in layout:
            scene = scenes[scene_idx]
            visual = panel_visuals[scene_idx]
            px, py = lx, ly + panel_y_offset
            self._draw_dataset_panel(
                img,
                draw,
                px,
                py,
                lw,
                lh,
                scene,
                visual,
                palette,
                style_record,
                panel_start_index + scene_idx,
            )

        img.save(image_path, format="PNG", dpi=(300, 300))

    def _build_layout(self, n, margin, gutter, cw, ch, sw, sh):
        """Return list of (x, y, w, h, scene_index) for comic panel grid.
        First panel is always a wide splash; rest fill a 2-col grid."""
        layout = []
        # Panel 0: wide splash (full width)
        full_w = margin + sw + gutter + cw
        layout.append((margin, 0, full_w - margin * 2 + margin, sh, 0))
        # Remaining panels: 2-col grid
        remaining = list(range(1, n))
        row = 0
        for i, si in enumerate(remaining):
            col = i % 2
            lx = margin + col * (cw + gutter)
            ly = sh + gutter + row * (ch + gutter)
            layout.append((lx, ly, cw, ch, si))
            if col == 1:
                row += 1
        return layout

    def _draw_title_banner(self, img, draw, x, y, w, h, title, subject, palette):
        """Classic comic book title header with publisher-style design."""
        # Main banner fill
        self._gradient_rect(img, x, y, x + w, y + h, palette["title_bg"], palette["accent_alt"], vertical=False)

        # Diagonal speed lines through the banner (dynamic energy)
        for i in range(0, w, 30):
            draw.line([(x + i, y), (x + i + h, y + h)], fill="#ffffff15", width=8)

        # Stripe: comic-style coloured ribbon at top of banner
        stripe_h = 18
        self._gradient_rect(img, x, y, x + w, y + stripe_h, palette["title_stripe"], palette["panel_top"], vertical=False)
        draw.rectangle((x, y + stripe_h, x + w, y + stripe_h + 4), fill="#1a1a1a")

        # Outer border — thick black
        draw.rectangle((x, y, x + w, y + h), outline="#1a1a1a", width=6)
        draw.rectangle((x + 6, y + 6, x + w - 6, y + h - 6), outline="#ffffff", width=2)

        # Star-burst decorations in corners
        self._draw_starburst(draw, x + 70, y + h // 2 + 10, 44, palette["title_stripe"], "#1a1a1a", spikes=12)
        self._draw_starburst(draw, x + w - 70, y + h // 2 + 10, 44, palette["title_stripe"], "#1a1a1a", spikes=12)

        # TITLE text with knockout shadow + outline
        title_font = self._font(54, bold=True)
        badge_font = self._font(18, bold=True)
        num_font = self._font(13, bold=True)

        title_text = title.upper()[:38]
        # Multi-layer text: shadow → dark outline → white fill
        for dx, dy in [(5, 5), (3, 3), (1, 1)]:
            draw.text((x + 28 + dx, y + stripe_h + 14 + dy), title_text, font=title_font, fill="#00000060")
        draw.text((x + 28, y + stripe_h + 14), title_text, font=title_font, fill="#ffffff")

        # Subject badge pill
        badge_text = f"  ✦  {subject.upper()} STUDY COMIC  ✦  "
        bw = len(badge_text) * 9 + 10
        by = y + stripe_h + 76
        draw.rounded_rectangle((x + 26, by, x + 26 + bw, by + 26), radius=12,
                                fill=palette["title_stripe"], outline="#1a1a1a", width=2)
        draw.text((x + 30, by + 4), badge_text, font=badge_font, fill="#1a1a1a")

        # Page number and branding bottom-right
        draw.text((x + w - 180, y + h - 22), "COMINOTE  •  STUDY EDITION", font=num_font, fill="#ffffff90")

    def _draw_superhero_panel(self, img, draw, px, py, pw, ph, scene, palette,
                               char_type, expr, mouth, idx, is_splash=False):
        """Draw a full superhero comic book panel."""
        # ── Drop shadow (heavy, like printed comics)
        shadow_off = 9
        draw.rectangle((px + shadow_off, py + shadow_off, px + pw + shadow_off, py + ph + shadow_off),
                        fill="#1a1a1a")

        # ── Sky / background gradient
        self._gradient_rect(img, px, py, px + pw, py + ph, palette["sky_top"], palette["sky_bot"])

        # ── Halftone dot fill for vintage print feel
        self._draw_page_halftone(draw, px + 4, py + 40, px + pw - 4, py + ph - 4,
                                  palette["accent_alt"], density=16, dot_r=3, opacity=0.18)

        # ── Rich environment background
        self._draw_comic_bg(draw, img, px, py, pw, ph, scene.background, palette, is_splash)

        # ── Ground line
        ground_y = py + ph - 85
        draw.rectangle((px, ground_y, px + pw, ground_y + 5), fill=_hex_to_rgb(palette["ground"]))
        for xi in range(px, px + pw, 16):
            draw.rectangle((xi, ground_y + 5, xi + 10, ground_y + 8), fill=_lerp_color(
                _hex_to_rgb(palette["ground"]), (0, 0, 0), 0.3))

        # ── Thick black outer panel border (most important comic element)
        draw.rectangle((px, py, px + pw, py + ph), outline="#1a1a1a", width=7)

        # ── Secondary accent border inside
        draw.rectangle((px + 7, py + 7, px + pw - 7, py + ph - 7), outline=palette["accent"], width=2)
        draw.rectangle((px + 10, py + 10, px + pw - 10, py + ph - 10), outline="#1a1a1a", width=1)

        # ── Header bar (coloured ribbon at top of panel)
        header_h = 44
        self._gradient_rect(img, px + 1, py + 1, px + pw - 1, py + header_h,
                             palette["accent"], palette["accent_alt"], vertical=False)
        # speed lines in header
        for xi in range(0, pw, 25):
            draw.line([(px + xi, py), (px + xi + header_h, py + header_h)], fill="#ffffff18", width=6)
        # Re-draw borders over gradient
        draw.rectangle((px, py, px + pw, py + ph), outline="#1a1a1a", width=7)
        draw.line([(px + 1, py + header_h), (px + pw - 1, py + header_h)], fill="#1a1a1a", width=3)

        # Panel number circle (bold black circle, white number)
        nc_x, nc_y = px + 28, py + 22
        draw.ellipse((nc_x - 18, nc_y - 18, nc_x + 18, nc_y + 18), fill="#1a1a1a")
        draw.ellipse((nc_x - 14, nc_y - 14, nc_x + 14, nc_y + 14), fill="#ffffff")
        num_font = self._font(18, bold=True)
        draw.text((nc_x, nc_y), str(idx + 1), font=num_font, fill=palette["accent"], anchor="mm")

        # Panel title in header
        hdr_font = self._font(22, bold=True)
        draw.text((nc_x + 26, py + 12), scene.title.upper(), font=hdr_font, fill="#ffffff")

        # ── Effect word starburst in top-right corner — ALWAYS present, big and bold
        effect = EFFECT_WORDS[idx % len(EFFECT_WORDS)]
        eff_x = px + pw - (80 if is_splash else 65)
        eff_y = py + (80 if is_splash else 70)
        eff_r = 50 if is_splash else 40
        self._draw_big_effect(draw, eff_x, eff_y, eff_r, effect, palette)

        # ── Character — bigger for splash panel
        scale = 1.35 if is_splash else 1.0
        char_cx = px + int(140 * scale)
        char_cy = py + ph - 45
        self._draw_superhero_character(draw, char_cx, char_cy, char_type, expr, mouth, palette, scale=scale)

        # ── Speech bubble — bigger and to the right of character
        bubble_x1 = px + int(240 * (1.1 if is_splash else 1.0))
        bubble_y1 = py + header_h + 16
        bubble_x2 = px + pw - 16
        bubble_y2 = py + ph - 92
        self._draw_comic_speech_bubble(draw, bubble_x1, bubble_y1, bubble_x2, bubble_y2,
                                        scene.dialogue, scene.speaker, palette, is_splash)

        # ── Narrator caption box at bottom (yellow, like classic comics)
        cap_y1 = py + ph - 82
        cap_y2 = py + ph - 10
        self._draw_narrator_caption(draw, px + 8, cap_y1, px + pw - 8, cap_y2,
                                     scene.caption, palette)

        # ── Speaker name badge (angled, punchy)
        name_font = self._font(13, bold=True)
        name_text = scene.speaker.upper()
        nw = len(name_text) * 8 + 16
        nx1, ny1 = px + 12, py + ph - 118
        nx2, ny2 = nx1 + nw, ny1 + 20
        draw.rectangle((nx1 + 2, ny1 + 2, nx2 + 2, ny2 + 2), fill="#1a1a1a")
        draw.rectangle((nx1, ny1, nx2, ny2), fill=palette["accent"], outline="#1a1a1a", width=2)
        draw.text((nx1 + 8, ny1 + 3), name_text, font=name_font, fill="#ffffff")

    def _draw_comic_speech_bubble(self, draw, x1, y1, x2, y2, text, speaker, palette, large=False):
        """Classic comic book speech bubble with thick black border and pointer."""
        if x2 <= x1 + 40 or y2 <= y1 + 20:
            return
        bw = x2 - x1
        bh = y2 - y1
        # Cloud bumps along top edge
        bump_r = 14
        bumps_y = y1
        for bxi in range(x1 + bump_r + 6, x2 - bump_r - 6, bump_r * 2 - 2):
            draw.ellipse((bxi - bump_r - 2, bumps_y - bump_r + 4,
                          bxi + bump_r + 2, bumps_y + bump_r + 4), fill="#1a1a1a")
        for bxi in range(x1 + bump_r + 6, x2 - bump_r - 6, bump_r * 2 - 2):
            draw.ellipse((bxi - bump_r, bumps_y - bump_r + 6,
                          bxi + bump_r, bumps_y + bump_r + 2), fill="#ffffff")
        # Main bubble outline (thick black = comic book look)
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill="#1a1a1a")
        draw.rounded_rectangle((x1 + 4, y1 + 4, x2 - 4, y2 - 4), radius=17, fill="#ffffff")
        # Subtle inner tone
        draw.rounded_rectangle((x1 + 5, y1 + 5, x2 - 5, y1 + 28), radius=14,
                                fill=_lerp_color((255, 255, 255), _hex_to_rgb(palette["sky_bot"]), 0.25))
        # Tail (pointing toward character at lower-left)
        tx = x1 + 36
        tail = [(tx, y2 - 6), (tx + 30, y2 - 6), (tx - 22, y2 + 38)]
        tail_in = [(tx + 3, y2 - 4), (tx + 26, y2 - 4), (tx - 16, y2 + 28)]
        draw.polygon(tail, fill="#1a1a1a")
        draw.polygon(tail_in, fill="#ffffff")
        # Dialogue text
        font_sz = 19 if large else 16
        font = self._font(font_sz)
        max_w = (bw - 32) // max(font_sz // 2, 1)
        wrapped = textwrap.fill(text, width=max(max_w, 18))
        draw.multiline_text((x1 + 16, y1 + 14), wrapped, font=font, fill="#1a1a1a",
                             spacing=7, stroke_width=0)

    def _draw_narrator_caption(self, draw, x1, y1, x2, y2, caption, palette):
        """Classic yellow narrator box at the bottom of a panel."""
        if x2 <= x1 or y2 <= y1:
            return
        # Shadow
        draw.rectangle((x1 + 3, y1 + 3, x2 + 3, y2 + 3), fill="#00000050")
        # Yellow box — CLASSIC comic narrator colour
        draw.rectangle((x1, y1, x2, y2), fill=_lerp_color(_hex_to_rgb(palette["panel_top"]), (255, 255, 255), 0.35),
                        outline="#1a1a1a", width=3)
        # Left accent stripe
        draw.rectangle((x1, y1, x1 + 6, y2), fill=palette["accent"])
        # Narrator label
        lbl_font = self._font(11, bold=True)
        draw.text((x1 + 10, y1 + 3), "NARRATOR:", font=lbl_font, fill=palette["caption_bg"])
        # Caption text
        cap_font = self._font(13, bold=True)
        short = textwrap.shorten(caption, width=72, placeholder="...")
        draw.text((x1 + 10, y1 + 17), short, font=cap_font, fill="#1a1a1a")

    # ─── Superhero Character Drawing ─────────────────────────────────────────

    def _draw_superhero_character(self, draw, cx, cy, char_type: str,
                                   expr: str, mouth: str, palette, scale=1.0):
        """Draw a bold superhero-style character with dynamic pose."""
        style = CHARACTER_STYLES.get(char_type, CHARACTER_STYLES["student"])
        hair_col, skin_col, costume_col, costume_alt = style

        s = scale
        head_r = int(54 * s)
        body_w = int(64 * s)
        body_h = int(90 * s)     # taller, more heroic body
        leg_h = int(70 * s)
        arm_w = int(20 * s)

        bx = cx - body_w // 2
        by = cy - leg_h - body_h

        # ── Ground shadow
        draw.ellipse((cx - int(44 * s), cy - int(6 * s), cx + int(44 * s), cy + int(10 * s)),
                     fill="#00000030")

        # ── Legs — heroic stance (slightly apart)
        ll_x = bx + int(6 * s)
        rl_x = bx + body_w - int(26 * s)
        for lx in [ll_x, rl_x]:
            draw.rounded_rectangle((lx, by + body_h - int(4 * s), lx + int(20 * s), cy - int(4 * s)),
                                    radius=int(8 * s), fill=costume_alt, outline="#1a1a1a", width=int(2 * s))
        # Shoes
        draw.rounded_rectangle((ll_x - int(4 * s), cy - int(8 * s), ll_x + int(24 * s), cy + int(8 * s)),
                                radius=int(6 * s), fill="#1a1a1a")
        draw.rounded_rectangle((rl_x - int(4 * s), cy - int(8 * s), rl_x + int(24 * s), cy + int(8 * s)),
                                radius=int(6 * s), fill="#1a1a1a")

        # ── Cape / back element for teacher/narrator
        if char_type in ("teacher", "narrator"):
            cape_pts = [
                (cx - int(38 * s), by),
                (cx - int(60 * s), cy - int(30 * s)),
                (cx - int(44 * s), cy + int(10 * s)),
                (cx - int(20 * s), by + int(40 * s)),
            ]
            draw.polygon(cape_pts, fill=costume_alt, outline="#1a1a1a", width=int(2 * s))

        # ── Body — muscular, slightly trapezoidal
        body_pts = [
            (bx - int(6 * s), by + int(20 * s)),        # top-left with shoulder flare
            (bx + body_w + int(6 * s), by + int(20 * s)),  # top-right
            (bx + body_w, by + body_h),                 # bottom-right
            (bx, by + body_h),                          # bottom-left
        ]
        draw.polygon(body_pts, fill=costume_col, outline="#1a1a1a", width=int(3 * s))

        # Costume chest emblem — stylised 'C' badge
        badge_cx, badge_cy = cx, by + int(42 * s)
        badge_r = int(18 * s)
        draw.ellipse((badge_cx - badge_r, badge_cy - badge_r,
                      badge_cx + badge_r, badge_cy + badge_r),
                     fill="#ffffff", outline="#1a1a1a", width=int(2 * s))
        draw.arc((badge_cx - int(10 * s), badge_cy - int(10 * s),
                  badge_cx + int(10 * s), badge_cy + int(10 * s)),
                 start=40, end=320, fill=palette["accent"], width=int(4 * s))

        # Costume waist belt line
        belt_y = by + body_h - int(18 * s)
        draw.rectangle((bx, belt_y, bx + body_w, belt_y + int(10 * s)),
                        fill="#1a1a1a")
        draw.rectangle((cx - int(10 * s), belt_y - int(2 * s), cx + int(10 * s), belt_y + int(12 * s)),
                        fill=palette["title_stripe"], outline="#1a1a1a", width=1)

        # ── Arms — action poses based on expression
        if expr in ("excited", "happy"):
            # Right arm UP — power pose
            arm_pts_r = [
                (bx + body_w + int(4 * s), by + int(22 * s)),
                (bx + body_w + int(28 * s), by - int(30 * s)),
                (bx + body_w + int(14 * s), by + int(26 * s)),
            ]
            draw.polygon(arm_pts_r, fill=costume_col, outline="#1a1a1a", width=int(2 * s))
            # Fist / hand (raised)
            fist_cx = bx + body_w + int(30 * s)
            fist_cy = by - int(38 * s)
            self._draw_fist(draw, fist_cx, fist_cy, int(12 * s), skin_col)
        else:
            # Right arm out — pointing or reaching
            arm_pts_r = [
                (bx + body_w + int(4 * s), by + int(22 * s)),
                (bx + body_w + int(38 * s), by + int(40 * s)),
                (bx + body_w + int(16 * s), by + int(30 * s)),
            ]
            draw.polygon(arm_pts_r, fill=costume_col, outline="#1a1a1a", width=int(2 * s))
            draw.ellipse((bx + body_w + int(30 * s), by + int(32 * s),
                          bx + body_w + int(52 * s), by + int(52 * s)),
                         fill=skin_col, outline="#1a1a1a", width=int(2 * s))

        if expr == "shock":
            # Left arm out to side — surprise pose
            arm_pts_l = [
                (bx - int(4 * s), by + int(22 * s)),
                (bx - int(36 * s), by + int(10 * s)),
                (bx - int(14 * s), by + int(30 * s)),
            ]
            draw.polygon(arm_pts_l, fill=costume_col, outline="#1a1a1a", width=int(2 * s))
            draw.ellipse((bx - int(48 * s), by + int(2 * s),
                          bx - int(28 * s), by + int(22 * s)),
                         fill=skin_col, outline="#1a1a1a", width=int(2 * s))
        else:
            # Left arm — angled down
            arm_pts_l = [
                (bx - int(4 * s), by + int(22 * s)),
                (bx - int(22 * s), by + int(52 * s)),
                (bx - int(8 * s), by + int(28 * s)),
            ]
            draw.polygon(arm_pts_l, fill=costume_col, outline="#1a1a1a", width=int(2 * s))
            draw.ellipse((bx - int(28 * s), by + int(46 * s),
                          bx - int(8 * s), by + int(66 * s)),
                         fill=skin_col, outline="#1a1a1a", width=int(2 * s))

        # ── Head
        hx = cx - head_r
        hy = by - head_r * 2 + int(14 * s)
        head_cx = cx
        head_cy = hy + head_r
        # Head shadow
        draw.ellipse((hx + int(6 * s), hy + int(6 * s),
                      hx + head_r * 2 + int(6 * s), hy + head_r * 2 + int(6 * s)),
                     fill="#00000028")
        # Neck
        draw.rounded_rectangle((cx - int(14 * s), hy + head_r * 2 - int(4 * s),
                                 cx + int(14 * s), by + int(22 * s)),
                                radius=int(4 * s), fill=skin_col, outline="#1a1a1a", width=int(2 * s))
        # Head fill
        draw.ellipse((hx, hy, hx + head_r * 2, hy + head_r * 2),
                     fill=skin_col, outline="#1a1a1a", width=int(3 * s))

        # ── Hair
        self._draw_comic_hair(draw, head_cx, head_cy, head_r, hair_col, char_type, s)

        # ── Eyes
        self._draw_comic_eyes(draw, head_cx, head_cy + int(4 * s), head_r, expr, palette, s)

        # ── Nose
        draw.ellipse((head_cx - int(4 * s), head_cy + int(12 * s),
                      head_cx + int(4 * s), head_cy + int(19 * s)), fill="#e09070")

        # ── Mouth
        self._draw_comic_mouth(draw, head_cx, head_cy + int(28 * s), mouth, s)

        # ── Blush marks
        draw.ellipse((head_cx - int(38 * s), head_cy + int(10 * s),
                      head_cx - int(18 * s), head_cy + int(22 * s)), fill="#ffb3a780")
        draw.ellipse((head_cx + int(18 * s), head_cy + int(10 * s),
                      head_cx + int(38 * s), head_cy + int(22 * s)), fill="#ffb3a780")

        # ── Accessories
        if char_type == "scientist":
            # Glasses
            draw.arc((head_cx - int(28 * s), head_cy - int(10 * s),
                      head_cx - int(6 * s), head_cy + int(10 * s)),
                     start=0, end=360, fill="#1a1a1a", width=int(2 * s))
            draw.arc((head_cx + int(6 * s), head_cy - int(10 * s),
                      head_cx + int(28 * s), head_cy + int(10 * s)),
                     start=0, end=360, fill="#1a1a1a", width=int(2 * s))
            draw.line([(head_cx - int(6 * s), head_cy - int(2 * s)),
                       (head_cx + int(6 * s), head_cy - int(2 * s))],
                      fill="#1a1a1a", width=int(2 * s))
            # Lab coat lapels
            draw.line([(cx - int(12 * s), by + int(18 * s)),
                       (cx - int(4 * s), by + int(50 * s))],
                      fill="#ffffff60", width=int(3 * s))
        elif char_type == "teacher":
            # Graduation cap
            draw.rectangle((head_cx - int(44 * s), head_cy - int(50 * s),
                             head_cx + int(44 * s), head_cy - int(35 * s)),
                            fill="#1a1a1a")
            draw.rectangle((head_cx - int(36 * s), head_cy - int(58 * s),
                             head_cx + int(36 * s), head_cy - int(48 * s)),
                            fill="#1a1a1a")
            # Tassel
            draw.line([(head_cx + int(40 * s), head_cy - int(50 * s)),
                       (head_cx + int(48 * s), head_cy - int(18 * s))],
                      fill=palette["title_stripe"], width=int(3 * s))
            draw.ellipse((head_cx + int(43 * s), head_cy - int(22 * s),
                          head_cx + int(55 * s), head_cy - int(10 * s)),
                         fill=palette["title_stripe"])

    def _draw_fist(self, draw, cx, cy, r, skin_col):
        draw.rounded_rectangle((cx - r, cy - r // 2, cx + r, cy + r),
                                radius=r // 3, fill=skin_col, outline="#1a1a1a", width=2)
        for fi in range(3):
            draw.line([(cx - r + 4, cy - r // 2 + fi * (r // 3)),
                       (cx + r - 4, cy - r // 2 + fi * (r // 3))],
                      fill="#00000030", width=1)

    def _draw_comic_hair(self, draw, cx, cy, r, hair_col, char_type, s=1.0):
        if char_type == "student":
            pts = [
                (cx - r, cy - int(10 * s)), (cx - r - int(20 * s), cy - r - int(10 * s)),
                (cx - r + int(8 * s), cy - r - int(22 * s)), (cx, cy - r - int(28 * s)),
                (cx + r - int(8 * s), cy - r - int(22 * s)), (cx + r + int(20 * s), cy - r - int(10 * s)),
                (cx + r, cy - int(10 * s)),
            ]
        elif char_type == "teacher":
            pts = [
                (cx - r, cy - int(8 * s)), (cx - r - int(8 * s), cy - r),
                (cx - r + int(4 * s), cy - r - int(12 * s)), (cx, cy - r - int(16 * s)),
                (cx + r - int(4 * s), cy - r - int(12 * s)), (cx + r + int(8 * s), cy - r),
                (cx + r, cy - int(8 * s)),
            ]
        elif char_type == "scientist":
            pts = [
                (cx - r, cy - int(6 * s)),
                (cx - r - int(12 * s), cy - r + int(6 * s)), (cx - r + int(4 * s), cy - r - int(4 * s)),
                (cx - r + int(16 * s), cy - r - int(18 * s)), (cx, cy - r - int(24 * s)),
                (cx + r - int(16 * s), cy - r - int(18 * s)), (cx + r - int(4 * s), cy - r - int(4 * s)),
                (cx + r + int(12 * s), cy - r + int(6 * s)), (cx + r, cy - int(6 * s)),
            ]
        else:
            pts = [
                (cx - r, cy - int(6 * s)), (cx - r - int(22 * s), cy + r),
                (cx - r + int(8 * s), cy - r - int(14 * s)), (cx, cy - r - int(22 * s)),
                (cx + r - int(8 * s), cy - r - int(14 * s)), (cx + r + int(22 * s), cy + r),
                (cx + r, cy - int(6 * s)),
            ]
        draw.polygon(pts, fill=hair_col, outline="#1a1a1a", width=int(2 * s))

    def _draw_comic_eyes(self, draw, cx, cy, r, expr, palette, s=1.0):
        eye_off = int(20 * s)
        ew = int(16 * s)
        eh = int(20 * s)
        for ex in [cx - eye_off, cx + eye_off]:
            if expr == "shock":
                draw.ellipse((ex - ew, cy - eh, ex + ew, cy + eh),
                             fill="#ffffff", outline="#1a1a1a", width=int(2 * s))
                draw.ellipse((ex - int(9 * s), cy - int(9 * s), ex + int(9 * s), cy + int(9 * s)),
                             fill="#1a1a1a")
                draw.ellipse((ex - int(4 * s), cy - int(4 * s), ex + int(2 * s), cy + int(2 * s)),
                             fill="#ffffff")
            elif expr == "think":
                if ex < cx:
                    draw.arc((ex - ew, cy - eh // 2, ex + ew, cy + eh // 2),
                             start=0, end=180, fill="#1a1a1a", width=int(3 * s))
                else:
                    draw.ellipse((ex - ew, cy - eh, ex + ew, cy + eh),
                                 fill="#ffffff", outline="#1a1a1a", width=int(2 * s))
                    draw.ellipse((ex - int(6 * s), cy - int(6 * s), ex + int(6 * s), cy + int(6 * s)),
                                 fill="#1a1a1a")
            elif expr == "excited":
                draw.ellipse((ex - ew, cy - eh, ex + ew, cy + eh),
                             fill="#ffffff", outline="#1a1a1a", width=int(2 * s))
                for sr in range(4):
                    rad = math.radians(sr * 90 + 45)
                    sx = ex + int(7 * math.cos(rad) * s)
                    sy = cy + int(7 * math.sin(rad) * s)
                    draw.line([(ex, cy), (sx, sy)], fill=palette["accent"], width=int(2 * s))
                draw.ellipse((ex - int(5 * s), cy - int(5 * s), ex + int(5 * s), cy + int(5 * s)),
                             fill=palette["accent"])
            else:
                draw.ellipse((ex - ew, cy - eh, ex + ew, cy + eh),
                             fill="#ffffff", outline="#1a1a1a", width=int(2 * s))
                iris = _lerp_color(_hex_to_rgb(palette["accent"]), _hex_to_rgb("#1a1a1a"), 0.4)
                draw.ellipse((ex - int(10 * s), cy - int(10 * s), ex + int(10 * s), cy + int(10 * s)),
                             fill=iris)
                draw.ellipse((ex - int(6 * s), cy - int(6 * s), ex + int(6 * s), cy + int(6 * s)),
                             fill="#1a1a1a")
                draw.ellipse((ex - int(4 * s), cy - int(9 * s), ex + int(1 * s), cy - int(4 * s)),
                             fill="#ffffff")
            draw.arc((ex - ew - int(2 * s), cy - eh - int(4 * s),
                      ex + ew + int(2 * s), cy + int(2 * s)),
                     start=200, end=340, fill="#1a1a1a", width=int(3 * s))

    def _draw_comic_mouth(self, draw, cx, my, mouth_style, s=1.0):
        if mouth_style == "smile":
            draw.arc((cx - int(16 * s), my - int(10 * s), cx + int(16 * s), my + int(10 * s)),
                     start=0, end=180, fill="#c0554a", width=int(3 * s))
        elif mouth_style == "wide":
            draw.arc((cx - int(18 * s), my - int(14 * s), cx + int(18 * s), my + int(14 * s)),
                     start=0, end=180, fill="#c0554a", width=int(3 * s))
            draw.ellipse((cx - int(14 * s), my - int(4 * s), cx + int(14 * s), my + int(12 * s)),
                         fill="#ff8076")
        elif mouth_style == "grin":
            draw.arc((cx - int(18 * s), my - int(12 * s), cx + int(18 * s), my + int(12 * s)),
                     start=0, end=180, fill="#c0554a", width=int(4 * s))
            draw.rounded_rectangle((cx - int(12 * s), my - int(4 * s), cx + int(12 * s), my + int(4 * s)),
                                    radius=int(3 * s), fill="#ffffff")
        elif mouth_style == "hmm":
            draw.line([(cx - int(12 * s), my), (cx + int(12 * s), my)], fill="#c0554a", width=int(3 * s))
        elif mouth_style == "o":
            draw.ellipse((cx - int(10 * s), my - int(12 * s), cx + int(10 * s), my + int(8 * s)),
                         fill="#c0554a", outline="#1a1a1a", width=int(2 * s))
            draw.ellipse((cx - int(6 * s), my - int(8 * s), cx + int(6 * s), my + int(4 * s)),
                         fill="#ff8076")
        else:
            draw.arc((cx - int(14 * s), my - int(8 * s), cx + int(14 * s), my + int(8 * s)),
                     start=0, end=180, fill="#c0554a", width=int(3 * s))

    # ─── Drawing Helpers ─────────────────────────────────────────────────────

    def _draw_big_effect(self, draw, cx, cy, r, word, palette, family: str = "display"):
        """Large starburst effect word — POW! BAM! etc."""
        ray_len = r + max(8, int(r * 0.45))
        # Outer rays
        for angle in range(0, 360, 18):
            rad = math.radians(angle)
            ex = cx + int(ray_len * math.cos(rad))
            ey = cy + int(ray_len * math.sin(rad))
            draw.line([(cx, cy), (ex, ey)], fill=palette["accent_alt"], width=max(2, int(r * 0.18)))
        # Starburst shape (not a simple circle — jagged polygon)
        pts = []
        spikes = 12
        for i in range(spikes * 2):
            rad = math.radians(i * 180 / spikes - 90)
            radius = r if i % 2 == 0 else r * 0.65
            pts.append((cx + int(radius * math.cos(rad)), cy + int(radius * math.sin(rad))))
        # Shadow
        shadow_pts = [(x + 4, y + 4) for x, y in pts]
        draw.polygon(shadow_pts, fill="#00000050")
        # Main burst
        draw.polygon(pts, fill=palette["accent"], outline="#1a1a1a", width=3)
        # Inner highlight ring
        inner_pts = []
        for i in range(spikes * 2):
            rad = math.radians(i * 180 / spikes - 90)
            radius = (r - 8) if i % 2 == 0 else (r - 8) * 0.65
            inner_pts.append((cx + int(radius * math.cos(rad)), cy + int(radius * math.sin(rad))))
        draw.polygon(inner_pts, outline=palette["title_stripe"], width=2)
        # Word text — bold, white with black outline
        font = self._font(max(14, int(r * 0.52)), bold=True, family=family)
        for dx, dy in [(2, 2), (-2, 2), (2, -2), (-2, -2)]:
            draw.text((cx + dx, cy + dy), word, font=font, fill="#1a1a1a", anchor="mm")
        draw.text((cx, cy), word, font=font, fill="#ffffff", anchor="mm")

    def _draw_starburst(self, draw, cx, cy, r, fill_col, outline_col, spikes=8):
        pts = []
        for i in range(spikes * 2):
            rad = math.radians(i * 180 / spikes - 90)
            radius = r if i % 2 == 0 else r * 0.55
            pts.append((cx + int(radius * math.cos(rad)), cy + int(radius * math.sin(rad))))
        draw.polygon(pts, fill=fill_col, outline=outline_col)

    def _draw_comic_bg(self, draw, img, px, py, pw, ph, background: str, palette, is_splash=False):
        """Draw a rich comic-book environment background."""
        bg_lower = background.lower()
        sky_col = _hex_to_rgb(palette["sky_top"])
        sky_bot_col = _hex_to_rgb(palette["sky_bot"])
        ground_col = _hex_to_rgb(palette["ground"])
        horizon_y = py + ph - 130
        accent_col = _hex_to_rgb(palette["accent"])

        # Sky gradient above horizon
        self._gradient_rect(img, px + 1, py + 44, px + pw - 1, horizon_y, palette["sky_top"], palette["sky_bot"])
        # Ground below horizon
        self._gradient_rect(img, px + 1, horizon_y, px + pw - 1, py + ph - 85,
                             palette["ground"], palette["panel_top"])

        # Speed/action lines on sky for drama
        if is_splash:
            sc_x = px + pw - 80
            sc_y = py + 120
            for angle in range(160, 280, 12):
                rad = math.radians(angle)
                line_len = 180
                ex = sc_x + int(line_len * math.cos(rad))
                ey = sc_y + int(line_len * math.sin(rad))
                draw.line([(sc_x, sc_y), (ex, ey)], fill="#ffffff25", width=4)

        if "lab" in bg_lower or "science" in bg_lower or "green" in bg_lower:
            # Lab bench with equipment
            bench_y = horizon_y - 50
            draw.rectangle((px + pw - 220, bench_y, px + pw - 10, horizon_y),
                            fill=_lerp_color(ground_col, (255, 255, 255), 0.55), outline="#1a1a1a", width=2)
            # Test tubes
            tube_cols = ["#00ccff", "#ff6688", "#44ff88"]
            for ti, tx in enumerate([px + pw - 200, px + pw - 165, px + pw - 130]):
                tc = _lerp_color(_hex_to_rgb(tube_cols[ti % 3]), (255, 255, 255), 0.35)
                draw.rounded_rectangle((tx, bench_y - 52, tx + 16, bench_y),
                                        radius=5, fill=tc, outline="#1a1a1a", width=2)
                draw.rectangle((tx - 5, bench_y - 58, tx + 21, bench_y - 50),
                                fill=_lerp_color(ground_col, (0, 0, 0), 0.35))
            # Beaker
            draw.rounded_rectangle((px + pw - 108, bench_y - 42, px + pw - 68, bench_y),
                                    radius=8, fill=_lerp_color((100, 200, 255), (255, 255, 255), 0.4),
                                    outline="#1a1a1a", width=2)
            draw.ellipse((px + pw - 100, bench_y - 16, px + pw - 76, bench_y - 4),
                         fill="#00ccff88")
        elif "board" in bg_lower or "desk" in bg_lower or "classroom" in bg_lower or "studio" in bg_lower:
            # Chalkboard / whiteboard
            board_x1 = px + pw - 210
            board_y1 = py + 60
            board_x2 = px + pw - 18
            board_y2 = horizon_y - 20
            draw.rectangle((board_x1 + 5, board_y1 + 5, board_x2 + 5, board_y2 + 5), fill="#00000040")
            draw.rectangle((board_x1, board_y1, board_x2, board_y2),
                            fill="#1a5432", outline="#8B5E3C", width=4)
            # Chalk lines (formula/text)
            draw.line([(board_x1 + 12, board_y1 + 22), (board_x2 - 12, board_y1 + 22)],
                      fill="#ffffff70", width=2)
            draw.text((board_x1 + 14, board_y1 + 30),
                      "∑ E = mc²", font=self._font(14, bold=True), fill="#ffffffcc")
            draw.line([(board_x1 + 12, board_y1 + 58), (board_x2 - 40, board_y1 + 58)],
                      fill="#ffffff50", width=1)
            # Tray at bottom
            draw.rectangle((board_x1, board_y2, board_x2, board_y2 + 10),
                            fill="#8B5E3C", outline="#1a1a1a", width=1)
        elif "hall" in bg_lower or "archive" in bg_lower or "library" in bg_lower or "garden" in bg_lower:
            # Columns / arches
            col_col = _lerp_color(sky_col, (255, 255, 255), 0.55)
            for ci_x in [px + pw - 220, px + pw - 150, px + pw - 50]:
                draw.rounded_rectangle((ci_x, py + 60, ci_x + 28, horizon_y),
                                        radius=6, fill=col_col, outline="#1a1a1a", width=2)
                draw.ellipse((ci_x - 8, py + 48, ci_x + 36, py + 80),
                             fill=col_col, outline="#1a1a1a", width=2)
        else:
            # Bookshelf default
            shelf_x = px + pw - 200
            for row_i in range(2):
                shelf_y = py + 70 + row_i * 70
                draw.rectangle((shelf_x, shelf_y + 52, shelf_x + 168, shelf_y + 58),
                                fill="#8B5E3C", outline="#1a1a1a", width=1)
                bk_cols = [palette["accent"], palette["accent_alt"], palette["accent3"],
                           palette["panel_top"], palette["accent"]]
                bk_x = shelf_x + 4
                for bki, bkc in enumerate(bk_cols):
                    bw_v = 24 + (bki % 3) * 6
                    draw.rectangle((bk_x, shelf_y + 2, bk_x + bw_v, shelf_y + 52),
                                    fill=_lerp_color(_hex_to_rgb(bkc), (255, 255, 255), 0.3),
                                    outline="#1a1a1a", width=1)
                    bk_x += bw_v + 2

    def _draw_speed_lines(self, draw, cx, cy, width, height, hex_col, count=20, opacity=0.05):
        """Radial speed lines from a point — classic comic energy effect."""
        col = _lerp_color(_hex_to_rgb(hex_col), (255, 255, 255), 1 - opacity)
        for i in range(count):
            angle = math.radians(i * 360 / count)
            # Vary start radius so lines don't originate from exact centre
            sr = 80 + (i % 3) * 40
            far = max(width, height) * 1.4
            sx = cx + int(sr * math.cos(angle))
            sy = cy + int(sr * math.sin(angle))
            ex = cx + int(far * math.cos(angle))
            ey = cy + int(far * math.sin(angle))
            draw.line([(sx, sy), (ex, ey)], fill=col, width=2 + (i % 2))

    def _draw_page_halftone(self, draw, x1, y1, x2, y2, hex_col, density=20, dot_r=2, opacity=0.12):
        """Ben-Day halftone dot pattern — the vintage comic printing look."""
        base = _hex_to_rgb(hex_col)
        dot_col = _lerp_color(base, (255, 255, 255), 1 - opacity)
        row = 0
        for gy in range(int(y1), int(y2), density):
            offset = (density // 2) if (row % 2) else 0
            for gx in range(int(x1) + offset, int(x2), density):
                draw.ellipse((gx, gy, gx + dot_r, gy + dot_r), fill=dot_col)
            row += 1

    def _gradient_rect(self, img, x1, y1, x2, y2, color_start: str, color_end: str, vertical=True):
        draw = ImageDraw.Draw(img)
        c1 = _hex_to_rgb(color_start)
        c2 = _hex_to_rgb(color_end)
        span = (y2 - y1) if vertical else (x2 - x1)
        for i in range(span):
            t = i / max(span, 1)
            col = _lerp_color(c1, c2, t)
            if vertical:
                draw.line([(x1, y1 + i), (x2, y1 + i)], fill=col)
            else:
                draw.line([(x1 + i, y1), (x1 + i, y2)], fill=col)

    def _draw_halftone(self, draw, x1, y1, x2, y2, hex_col, density=24, opacity=0.15):
        base = _hex_to_rgb(hex_col)
        dot_col = _lerp_color(base, (255, 255, 255), 1 - opacity)
        for gy in range(y1, y2, density):
            for gx in range(x1, x2, density):
                r = density // 5
                draw.ellipse((gx, gy, gx + r, gy + r), fill=dot_col)

    def _draw_burst(self, draw, cx, cy, r, fill_col, outline_col):
        """Draw a starburst / explosion shape."""
        pts = []
        spikes = 8
        for i in range(spikes * 2):
            rad = math.radians(i * 180 / spikes - 90)
            radius = r if i % 2 == 0 else r * 0.62
            pts.append((cx + int(radius * math.cos(rad)), cy + int(radius * math.sin(rad))))
        draw.polygon(pts, fill=fill_col, outline=outline_col)

    def _font(self, size: int, bold: bool = False, family: str = "comic"):
        if ImageFont is None:
            raise DependencyError("Pillow is required for comic rendering.")

        family = (family or "comic").lower()
        font_map = {
            "rounded": [
                "/Library/Fonts/Fredoka-Bold.ttf" if bold else "/Library/Fonts/Fredoka-Regular.ttf",
                "/Library/Fonts/Nunito-Bold.ttf" if bold else "/Library/Fonts/Nunito-Regular.ttf",
                "/Library/Fonts/Poppins-Bold.ttf" if bold else "/Library/Fonts/Poppins-Regular.ttf",
                "/System/Library/Fonts/Supplemental/Arial Rounded Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial Rounded Bold.ttf",
                "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
                "/System/Library/Fonts/Helvetica.ttc",
                "/System/Library/Fonts/Supplemental/Trebuchet MS Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Trebuchet MS.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            ],
            "playful": [
                "/Library/Fonts/Fredoka-Bold.ttf" if bold else "/Library/Fonts/Fredoka-Regular.ttf",
                "/System/Library/Fonts/Supplemental/Marker Felt Wide.ttf",
                "/System/Library/Fonts/Supplemental/Chalkboard Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Chalkboard.ttf",
                "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
                "/System/Library/Fonts/Helvetica.ttc",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            ],
            "display": [
                "/Library/Fonts/Bangers-Regular.ttf",
                "/System/Library/Fonts/Supplemental/Impact.ttf",
                "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
            ],
            "condensed": [
                "/System/Library/Fonts/Supplemental/Arial Narrow Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial Narrow.ttf",
                "/System/Library/Fonts/Supplemental/Trebuchet MS Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Trebuchet MS.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
            ],
            "serif": [
                "/System/Library/Fonts/Supplemental/Georgia Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Georgia.ttf",
                "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
            ],
            "sans": [
                "/Library/Fonts/Fredoka-Bold.ttf" if bold else "/Library/Fonts/Fredoka-Regular.ttf",
                "/Library/Fonts/Nunito-Bold.ttf" if bold else "/Library/Fonts/Nunito-Regular.ttf",
                "/Library/Fonts/Poppins-Bold.ttf" if bold else "/Library/Fonts/Poppins-Regular.ttf",
                "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
                "/System/Library/Fonts/Supplemental/Helvetica.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            ],
            "comic": [
                "/Library/Fonts/ComicNeue-Bold.ttf" if bold else "/Library/Fonts/ComicNeue-Regular.ttf",
                "/Library/Fonts/Comic Neue Bold.ttf" if bold else "/Library/Fonts/Comic Neue Regular.ttf",
                "/System/Library/Fonts/Supplemental/Comic Sans MS Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Comic Sans MS.ttf",
                "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
                "/System/Library/Fonts/Supplemental/Helvetica.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            ],
        }
        candidates = font_map.get(family, font_map["comic"]) + [
            "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf" if bold else "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        ]
        for candidate in candidates:
            path = Path(candidate)
            if path.exists():
                try:
                    return ImageFont.truetype(str(path), size=size)
                except OSError:
                    continue
        return ImageFont.load_default()

    def _ensure_jpeg(self, comic_id: str, png_path: Path) -> Path:
        jpeg_path = self.output_dir / f"{comic_id}.jpeg"
        if jpeg_path.exists():
            return jpeg_path
        if Image is None:
            raise DependencyError("JPEG export needs Pillow installed in the Python environment.")
        with Image.open(png_path) as source:
            source.convert("RGB").save(jpeg_path, format="JPEG", quality=94)
        return jpeg_path

    def _ensure_pdf(self, comic_id: str, png_path: Path) -> Path:
        pdf_path = self.output_dir / f"{comic_id}.pdf"
        if pdf_path.exists():
            return pdf_path
        if Image is None:
            raise DependencyError("PDF export needs Pillow installed in the Python environment.")
        page_paths = self._comic_page_paths(comic_id)
        if page_paths:
            images = []
            try:
                for path in page_paths:
                    with Image.open(path) as source:
                        images.append(source.convert("RGB").copy())
                if images:
                    first, rest = images[0], images[1:]
                    first.save(pdf_path, format="PDF", resolution=150.0, save_all=True, append_images=rest)
                    return pdf_path
            finally:
                for image in images:
                    image.close()
        with Image.open(png_path) as source:
            source.convert("RGB").save(pdf_path, format="PDF", resolution=150.0)
        return pdf_path

    def _comic_page_paths(self, comic_id: str) -> list[Path]:
        metadata_path = self.output_dir / f"{comic_id}.json"
        if not metadata_path.exists():
            return []
        try:
            payload = json.loads(metadata_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return []
        page_files = ((payload.get("meta") or {}).get("page_image_files") or [])
        paths: list[Path] = []
        for value in page_files:
            path = self.output_dir / str(value)
            if path.exists() and path.suffix.lower() == ".png":
                paths.append(path)
        return paths

    def _write_metadata(self, comic_id: str, payload: dict) -> None:
        path = self.output_dir / f"{comic_id}.json"
        path.write_text(json.dumps(payload, indent=2), encoding="utf-8")

    @staticmethod
    def _notify(
        callback: ProgressCallback | None, stage: str, progress: int,
        message: str, meta: dict[str, Any] | None = None,
    ) -> None:
        if callback is None:
            return
        callback(stage, progress, message, meta)
